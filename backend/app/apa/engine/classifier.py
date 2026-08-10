"""Document region / heading classification with confidence and evidence."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

from app.apa.models import Region
from app.apa.registry.formatting import effective_bold, effective_italic


@dataclass(frozen=True)
class ClassifiedParagraph:
    index: int
    paragraph: Paragraph
    region: Region
    confidence: float
    text: str
    style_name: str
    heading_level: Optional[int] = None
    evidence: tuple[str, ...] = field(default_factory=tuple)
    title_element: Optional[str] = None  # paper_title|author|affiliation|course|instructor|due_date
    is_block_quote_candidate: bool = False
    block_quote_confidence: float = 0.0


def _safe_style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name if paragraph.style else ""
    except Exception:
        return ""


def _normalize(text: str) -> str:
    return " ".join((text or "").strip().lower().split())


def _outline_level(paragraph: Paragraph) -> Optional[int]:
    try:
        level = paragraph.paragraph_format.outline_level
        if level is None:
            return None
        # python-docx may expose WD_OUTLINE_LEVEL enum or int
        value = int(level) if not hasattr(level, "real") else int(level)
        if 0 <= value <= 8:
            return value + 1  # outline 0 ~= Heading 1
    except Exception:
        pass
    match = re.match(r"heading\s+(\d+)", _safe_style_name(paragraph), flags=re.I)
    if match:
        return int(match.group(1))
    return None


def _style_heading_level(style_name: str) -> Optional[int]:
    match = re.match(r"heading\s+(\d+)$", (style_name or "").strip(), flags=re.I)
    if match:
        return int(match.group(1))
    return None


def _is_centered(paragraph: Paragraph) -> bool:
    return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def _looks_like_keywords_label(text: str) -> bool:
    return bool(re.match(r"^keywords\s*:", _normalize(text)))


def _looks_like_date(text: str) -> bool:
    return bool(
        re.search(
            r"\b(january|february|march|april|may|june|july|august|september|"
            r"october|november|december|"
            r"jan\.?|feb\.?|mar\.?|apr\.?|jun\.?|jul\.?|aug\.?|sep\.?|sept\.?|"
            r"oct\.?|nov\.?|dec\.?)\b"
            r"|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
            r"|\b\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4}\b"
            r"|\b[A-Za-z]{3,9}\s+\d{1,2},\s*\d{4}\b",
            text,
            flags=re.I,
        )
    )


def _looks_like_course(text: str) -> bool:
    # Accept CS 240, CS240, PSYCH 101, PSYCH 101: Introduction to Psychology, etc.
    return bool(
        re.search(
            r"\b(course|psy(?:ch)?|engl|hist|bio|math|cs|stat|chem|phys|soc|anth)\b"
            r"|[A-Z]{2,6}\s*\d{2,4}(?:\s*:\s*.+)?",
            text,
            flags=re.I,
        )
    )


def _looks_like_instructor(text: str) -> bool:
    if re.search(r"\b(dr\.|prof\.|professor|instructor)\b", text, flags=re.I):
        return True
    # Bare personal name near end of title page is handled in refine pass.
    return False


def _looks_like_affiliation(text: str) -> bool:
    return bool(
        re.search(
            r"\b(university|college|department|school|institute|faculty|program|"
            r"centre|center|academy|polytechnic)\b",
            text,
            flags=re.I,
        )
    )


def _looks_like_person_name(text: str) -> bool:
    """Conservative person-name heuristic for title-page authors (not citations)."""
    t = (text or "").strip()
    if not t or len(t) > 80:
        return False
    if _looks_like_date(t) or _looks_like_course(t) or _looks_like_affiliation(t):
        return False
    if re.search(r"\b(dr\.|prof\.|professor|instructor|acknowledg|thanks)\b", t, re.I):
        return False
    # "Jane Student" / "Jane A. Student" / "Jane Student and John Student"
    if re.match(
        r"^[A-Z][A-Za-z'’\-]+(?:\s+[A-Z]\.)?(?:\s+[A-Z][A-Za-z'’\-]+)+"
        r"(?:\s+(?:and|&)\s+[A-Z][A-Za-z'’\-]+(?:\s+[A-Z]\.)?(?:\s+[A-Z][A-Za-z'’\-]+)+)*"
        r"$",
        t,
    ):
        return True
    # Comma form: "Student, Jane"
    if re.match(r"^[A-Z][A-Za-z'’\-]+,\s*[A-Z][A-Za-z'’\-\.\s]+$", t) and len(t.split()) <= 5:
        return True
    return False


def _find_section_markers(paragraphs: list[Paragraph]) -> dict:
    ref_start = None
    abstract_idx = None
    keywords_idx = None
    for index, paragraph in enumerate(paragraphs):
        norm = _normalize(paragraph.text)
        if ref_start is None and norm == "references":
            ref_start = index
        if abstract_idx is None and norm == "abstract":
            abstract_idx = index
        if keywords_idx is None and _looks_like_keywords_label(paragraph.text):
            keywords_idx = index
    return {
        "ref_start": ref_start,
        "abstract_idx": abstract_idx,
        "keywords_idx": keywords_idx,
    }


def _detect_body_start(
    paragraphs: list[Paragraph],
    abstract_idx: Optional[int],
    ref_start: Optional[int],
) -> tuple[int, tuple[str, ...]]:
    """
    Body starts after title-page front matter.

    Prefer: Abstract → repeated paper title → first non-References heading
    → known body section labels. Never require Introduction.
    """
    evidence: list[str] = []
    n = len(paragraphs)
    body_start = n

    # Collect early Title-style / centered title candidates on page 1 block.
    title_block_end = 0
    first_title_text = ""
    for index, paragraph in enumerate(paragraphs[:12]):
        style = _safe_style_name(paragraph).lower()
        text = paragraph.text.strip()
        if not text:
            continue
        # A page break marks the end of the title-page block.
        try:
            if index > 0 and paragraph.paragraph_format.page_break_before is True:
                break
        except Exception:
            pass
        if style.startswith("heading") and _normalize(text) != "abstract":
            break
        if style in {"title", "subtitle"} or (
            index < 8 and _is_centered(paragraph) and len(text.split()) >= 2
        ):
            title_block_end = max(title_block_end, index + 1)
            if not first_title_text and style in {"title", "subtitle"}:
                first_title_text = text.lower()
            elif not first_title_text and index == 0:
                first_title_text = text.lower()
        elif text and index > 0 and not _is_centered(paragraph):
            break

    if abstract_idx is not None:
        body_start = abstract_idx
        evidence.append("abstract_heading")
    else:
        # Repeated paper title after title-page metadata is a strong body marker.
        if first_title_text and title_block_end > 0:
            for index in range(title_block_end, min(title_block_end + 8, n)):
                text = paragraphs[index].text.strip()
                if text and text.lower() == first_title_text:
                    body_start = index
                    evidence.append("repeated_paper_title")
                    break

        # First content heading (not References) after title block.
        for index, paragraph in enumerate(paragraphs):
            style = _safe_style_name(paragraph).lower()
            text = paragraph.text.strip()
            if not text:
                continue
            if _normalize(text) == "references":
                continue
            style_level = _style_heading_level(style)
            if style_level is not None:
                if index == 0 and _is_centered(paragraph) and style_level == 1:
                    continue
                # Prefer the earliest content heading unless repeated title is earlier.
                if index < body_start:
                    body_start = index
                    evidence.append(f"heading_{style_level}_style")
                break
            if _normalize(text) in {
                "method",
                "methods",
                "results",
                "discussion",
                "literature review",
            }:
                if index < body_start:
                    body_start = index
                    evidence.append("known_body_label")
                break
            # Optional: Introduction as a *signal*, not a requirement.
            if _normalize(text) == "introduction" and index > 0:
                if index < body_start:
                    body_start = index
                    evidence.append("introduction_label_optional")
                break

    if ref_start is not None:
        body_start = min(body_start, ref_start)
        evidence.append("references_bound")

    return body_start, tuple(evidence)


def _classify_heading_level(
    paragraph: Paragraph,
    style_name: str,
) -> tuple[Optional[int], float, tuple[str, ...]]:
    evidence: list[str] = []
    style_level = _style_heading_level(style_name)
    outline = _outline_level(paragraph)
    bold = effective_bold(paragraph)
    italic = effective_italic(paragraph)
    centered = _is_centered(paragraph)
    leftish = paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)
    indent = paragraph.paragraph_format.first_line_indent
    left_indent = paragraph.paragraph_format.left_indent

    if style_level is not None:
        evidence.append(f"style:Heading {style_level}")
        conf = 0.97
        # Formatting can refine confidence but style wins for level identity.
        if style_level == 1 and centered:
            evidence.append("centered")
            conf = 0.99
        if style_level in {2, 3} and leftish:
            evidence.append("flush_left")
            conf = max(conf, 0.97)
        if style_level in {4, 5}:
            evidence.append("style_level_4_or_5")
        return style_level, conf, tuple(evidence)

    if outline is not None and 1 <= outline <= 5:
        evidence.append(f"outline_level:{outline}")
        return outline, 0.9, tuple(evidence)

    # Formatting-only heuristics — lower confidence.
    text = paragraph.text.strip()
    if not text or len(text) > 120:
        return None, 0.0, tuple(evidence)

    ends_with_period = text.endswith(".")
    indented = False
    try:
        if indent is not None and abs(indent.inches - 0.5) < 0.05:
            indented = True
        if left_indent is not None and abs(left_indent.inches - 0.5) < 0.05:
            indented = True
    except Exception:
        pass

    if indented and bold is True and ends_with_period:
        evidence.extend(["indent_0.5", "bold", "ends_period"])
        if italic is True:
            evidence.append("italic")
            return 5, 0.78, tuple(evidence)
        return 4, 0.78, tuple(evidence)

    if centered and bold is True:
        evidence.extend(["centered", "bold"])
        return 1, 0.72, tuple(evidence)

    if leftish and bold is True and italic is True:
        evidence.extend(["flush_left", "bold", "italic"])
        return 3, 0.7, tuple(evidence)

    if leftish and bold is True and italic is False:
        evidence.extend(["flush_left", "bold"])
        return 2, 0.68, tuple(evidence)

    return None, 0.0, tuple(evidence)


def _region_for_heading_level(level: int) -> Region:
    return {
        1: Region.LEVEL_1_HEADING,
        2: Region.LEVEL_2_HEADING,
        3: Region.LEVEL_3_HEADING,
        4: Region.LEVEL_4_HEADING,
        5: Region.LEVEL_5_HEADING,
    }.get(level, Region.HEADING)


def classify_document(doc: Document) -> list[ClassifiedParagraph]:
    paragraphs = list(doc.paragraphs)
    markers = _find_section_markers(paragraphs)
    ref_start = markers["ref_start"]
    abstract_idx = markers["abstract_idx"]
    keywords_idx = markers["keywords_idx"]
    body_start, body_evidence = _detect_body_start(
        paragraphs, abstract_idx, ref_start
    )

    # Title-page paper title candidate: first substantial title-styled or
    # first centered longish line before body.
    paper_title_idx: Optional[int] = None
    for index, paragraph in enumerate(paragraphs[: body_start if body_start < len(paragraphs) else 8]):
        style = _safe_style_name(paragraph).lower()
        text = paragraph.text.strip()
        if not text:
            continue
        if style in {"title", "subtitle"}:
            paper_title_idx = index
            break
        if paper_title_idx is None and len(text.split()) >= 3:
            paper_title_idx = index
            break

    # Body repeated title: first body paragraph matching paper title text.
    body_title_idx: Optional[int] = None
    paper_title_text = (
        paragraphs[paper_title_idx].text.strip().lower()
        if paper_title_idx is not None
        else ""
    )
    if paper_title_text and body_start < len(paragraphs):
        for index in range(body_start, min(body_start + 3, len(paragraphs))):
            if paragraphs[index].text.strip().lower() == paper_title_text:
                body_title_idx = index
                break

    classified: list[ClassifiedParagraph] = []
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text
        stripped = text.strip()
        style = _safe_style_name(paragraph)
        region = Region.UNKNOWN
        confidence = 0.5
        heading_level: Optional[int] = None
        evidence: list[str] = []
        title_element: Optional[str] = None

        if not stripped:
            region = Region.UNKNOWN
            confidence = 0.9
            evidence.append("empty")
        elif ref_start is not None and index == ref_start:
            region = Region.REFERENCES_HEADING
            confidence = 0.99
            heading_level = 1
            evidence.append("references_label")
        elif ref_start is not None and index > ref_start:
            if _normalize(stripped).startswith("appendix"):
                region = Region.APPENDIX
                confidence = 0.95
                evidence.append("appendix_after_references")
            elif _looks_like_reference_entry(stripped):
                region = Region.REFERENCE_ENTRY
                confidence = 0.97
                evidence.append("after_references_pattern")
            else:
                region = Region.REFERENCE_ENTRY
                confidence = 0.85
                evidence.append("after_references_weak")
        elif _normalize(stripped).startswith("appendix") and (
            ref_start is None or index > ref_start
        ):
            region = Region.APPENDIX
            confidence = 0.9
            evidence.append("appendix_label")
        elif _looks_like_table_note(stripped):
            region = Region.TABLE_NOTE
            confidence = 0.88
            evidence.append("table_note")
        elif _looks_like_figure_note(stripped):
            region = Region.FIGURE_NOTE
            confidence = 0.88
            evidence.append("figure_note")
        elif abstract_idx is not None and index == abstract_idx:
            region = Region.ABSTRACT_HEADING
            confidence = 0.98
            heading_level = 1
            evidence.append("abstract_label")
        elif keywords_idx is not None and index == keywords_idx:
            region = Region.KEYWORDS
            confidence = 0.95
            evidence.append("keywords_label")
        elif (
            abstract_idx is not None
            and index > abstract_idx
            and (keywords_idx is None or index < keywords_idx)
            and (body_title_idx is None or index < body_title_idx)
            and index < (body_start if body_start != abstract_idx else len(paragraphs))
        ):
            # Abstract body text between Abstract heading and keywords/body.
            region = Region.ABSTRACT
            confidence = 0.9
            evidence.append("abstract_body")
        elif body_title_idx is not None and index == body_title_idx:
            region = Region.BODY_TITLE
            confidence = 0.95
            evidence.append("repeated_title")
        elif index == paper_title_idx and index < body_start:
            region = Region.PAPER_TITLE
            confidence = 0.93 if style.lower() in {"title", "subtitle"} else 0.85
            title_element = "paper_title"
            evidence.append("paper_title_candidate")
        elif _looks_like_table_caption(stripped):
            region = Region.TABLE_CAPTION
            confidence = 0.9
        elif _looks_like_figure_caption(stripped):
            region = Region.FIGURE_CAPTION
            confidence = 0.9
        else:
            level, level_conf, level_evidence = _classify_heading_level(
                paragraph, style
            )
            style_level = _style_heading_level(style)
            if style_level is not None and index >= body_start:
                region = _region_for_heading_level(style_level)
                heading_level = style_level
                confidence = level_conf
                evidence.extend(level_evidence)
            elif (
                level is not None
                and level_conf >= 0.9
                and index >= body_start
                and style_level is not None
            ):
                region = _region_for_heading_level(level)
                heading_level = level
                confidence = level_conf
                evidence.extend(level_evidence)
            elif index < body_start:
                region = Region.TITLE_PAGE
                confidence = 0.8 if body_start < len(paragraphs) else 0.55
                evidence.extend(["before_body", *body_evidence])
                if _looks_like_date(stripped):
                    title_element = "due_date"
                    confidence = max(confidence, 0.86)
                elif _looks_like_instructor(stripped):
                    title_element = "instructor"
                    confidence = max(confidence, 0.86)
                elif _looks_like_course(stripped):
                    title_element = "course"
                    confidence = max(confidence, 0.84)
                elif _looks_like_affiliation(stripped):
                    title_element = "affiliation"
                    confidence = max(confidence, 0.84)
                elif paper_title_idx is not None and index > paper_title_idx and _looks_like_person_name(
                    stripped
                ):
                    # May be one of multiple authors; refined below.
                    title_element = "author"
                    confidence = max(confidence, 0.82)
                    evidence.append("author_name_heuristic")
                elif index == (paper_title_idx or -1) + 1:
                    # Fallback single-author slot when name heuristic is weak.
                    title_element = "author"
                    confidence = max(confidence, 0.78)
                    evidence.append("author_position_fallback")
            elif style.lower().startswith("heading"):
                # Ambiguous heading style after body start without clear level.
                region = Region.HEADING
                heading_level = level
                confidence = max(level_conf, 0.75)
                evidence.extend(level_evidence or ["heading_style"])
            else:
                region = Region.BODY
                confidence = 0.92 if body_start < len(paragraphs) else 0.7
                evidence.append("body_default")

        classified.append(
            ClassifiedParagraph(
                index=index,
                paragraph=paragraph,
                region=region,
                confidence=confidence,
                text=text,
                style_name=style,
                heading_level=heading_level,
                evidence=tuple(evidence),
                title_element=title_element,
            )
        )

    classified = _apply_block_quote_pass(classified)
    classified = _refine_title_page_authors(classified)
    return _apply_phase2d_region_pass(doc, classified)


def _refine_title_page_authors(
    classified: list[ClassifiedParagraph],
) -> list[ClassifiedParagraph]:
    """Mark consecutive title-page person-name lines as authors; avoid metadata FPs."""
    title_idx = next(
        (c.index for c in classified if c.region == Region.PAPER_TITLE),
        None,
    )
    if title_idx is None:
        return classified

    by_index = {c.index: c for c in classified}
    # Title-page block ends at first body/abstract/etc.
    title_page_idxs = sorted(
        c.index
        for c in classified
        if c.region in {Region.TITLE_PAGE, Region.PAPER_TITLE} and c.index >= title_idx
    )
    # After affiliations begin, a bare name near the end may be the instructor.
    seen_affiliation = False
    seen_course = False
    for idx in title_page_idxs:
        item = by_index[idx]
        text = (item.text or "").strip()
        if not text or item.region == Region.PAPER_TITLE:
            continue
        if item.title_element == "affiliation" or _looks_like_affiliation(text):
            seen_affiliation = True
            continue
        if item.title_element == "course" or _looks_like_course(text):
            seen_course = True
            continue
        if item.title_element in {"instructor", "due_date"}:
            continue
        if item.title_element == "author":
            # Demote false author after course/affiliation if it looks like instructor name.
            if (seen_course or seen_affiliation) and _looks_like_person_name(text):
                # Keep as author only if still in author block (before course).
                if seen_course:
                    by_index[idx] = ClassifiedParagraph(
                        index=item.index,
                        paragraph=item.paragraph,
                        region=item.region,
                        confidence=max(item.confidence, 0.84),
                        text=item.text,
                        style_name=item.style_name,
                        heading_level=item.heading_level,
                        evidence=tuple([*item.evidence, "instructor_position_refine"]),
                        title_element="instructor",
                        is_block_quote_candidate=item.is_block_quote_candidate,
                        block_quote_confidence=item.block_quote_confidence,
                    )
            continue
        # Promote untagged person names before course/affiliation to author.
        if (
            not seen_course
            and not item.title_element
            and _looks_like_person_name(text)
        ):
            by_index[idx] = ClassifiedParagraph(
                index=item.index,
                paragraph=item.paragraph,
                region=item.region,
                confidence=max(item.confidence, 0.83),
                text=item.text,
                style_name=item.style_name,
                heading_level=item.heading_level,
                evidence=tuple([*item.evidence, "author_refine"]),
                title_element="author",
                is_block_quote_candidate=item.is_block_quote_candidate,
                block_quote_confidence=item.block_quote_confidence,
            )

    return [by_index[c.index] for c in classified]


def _apply_phase2d_region_pass(
    doc: Document,
    classified: list[ClassifiedParagraph],
) -> list[ClassifiedParagraph]:
    """Overlay table/figure/appendix/list regions without rewriting text."""
    from app.apa.parsing.appendices import analyze_appendices
    from app.apa.parsing.figures import analyze_figures
    from app.apa.parsing.lists import analyze_lists
    from app.apa.parsing.tables import analyze_tables

    by_index = {item.index: item for item in classified}

    def _replace(index: int, region: Region, confidence: float, *extra: str) -> None:
        item = by_index.get(index)
        if item is None:
            return
        # Do not demote high-confidence References / Abstract labels.
        if item.region in {
            Region.REFERENCES_HEADING,
            Region.ABSTRACT_HEADING,
            Region.PAPER_TITLE,
            Region.BODY_TITLE,
        }:
            return
        by_index[index] = ClassifiedParagraph(
            index=item.index,
            paragraph=item.paragraph,
            region=region,
            confidence=max(item.confidence, confidence),
            text=item.text,
            style_name=item.style_name,
            heading_level=item.heading_level,
            evidence=tuple([*item.evidence, *extra]),
            title_element=item.title_element,
            is_block_quote_candidate=item.is_block_quote_candidate,
            block_quote_confidence=item.block_quote_confidence,
        )

    for table in analyze_tables(doc):
        if table.number_paragraph_index is not None and table.confidence >= 0.9:
            _replace(
                table.number_paragraph_index,
                Region.TABLE_NUMBER,
                table.confidence,
                "table_number_region",
            )
        if (
            table.title_paragraph_index is not None
            and table.confidence >= 0.9
            and table.number_paragraph_index is not None
        ):
            _replace(
                table.title_paragraph_index,
                Region.TABLE_TITLE,
                table.confidence,
                "table_title_region",
            )
        for n_idx in table.note_paragraph_indices:
            text = (doc.paragraphs[n_idx].text or "").strip().lower()
            if text.startswith("note."):
                region = Region.TABLE_GENERAL_NOTE
            elif text.startswith("*") or text.lower().startswith("p "):
                region = Region.TABLE_PROBABILITY_NOTE
            else:
                region = Region.TABLE_NOTE
            _replace(n_idx, region, table.confidence, "table_note_region")

    for fig in analyze_figures(doc):
        if fig.number_paragraph_index is not None and fig.confidence >= 0.9:
            _replace(
                fig.number_paragraph_index,
                Region.FIGURE_NUMBER,
                fig.confidence,
                "figure_number_region",
            )
        if (
            fig.title_paragraph_index is not None
            and fig.confidence >= 0.9
            and fig.number_paragraph_index is not None
        ):
            _replace(
                fig.title_paragraph_index,
                Region.FIGURE_TITLE,
                fig.confidence,
                "figure_title_region",
            )
        if fig.image_paragraph_index is not None and fig.is_apa_candidate:
            _replace(
                fig.image_paragraph_index,
                Region.FIGURE,
                fig.confidence,
                "figure_image_region",
            )
        for n_idx in fig.note_paragraph_indices:
            _replace(n_idx, Region.FIGURE_NOTE, fig.confidence, "figure_note_region")

    for apx in analyze_appendices(doc):
        if apx.label_paragraph_index is not None and apx.confidence >= 0.9:
            _replace(
                apx.label_paragraph_index,
                Region.APPENDIX_LABEL,
                apx.confidence,
                "appendix_label_region",
            )
        if apx.title_paragraph_index is not None and apx.confidence >= 0.9:
            _replace(
                apx.title_paragraph_index,
                Region.APPENDIX_TITLE,
                apx.confidence,
                "appendix_title_region",
            )
        for b_idx in apx.body_paragraph_indices:
            item = by_index.get(b_idx)
            if item is None:
                continue
            if item.region in {Region.REFERENCE_ENTRY, Region.REFERENCES_HEADING}:
                continue
            _replace(b_idx, Region.APPENDIX_BODY, apx.confidence, "appendix_body_region")

    for lst in analyze_lists(doc):
        if lst.list_type in {"NUMBERED_LIST", "BULLETED_LIST"} and lst.confidence >= 0.85:
            _replace(lst.paragraph_index, Region.LIST_ITEM, lst.confidence, "list_item_region")

    return [by_index[i] for i in sorted(by_index)]


def _looks_like_reference_entry(text: str) -> bool:
    if re.search(r"\((\d{4}[a-z]?|n\.d\.)\)", text):
        return True
    if re.search(r"10\.\d{4,9}/", text):
        return True
    if re.search(r"https?://", text, re.I):
        return True
    return bool(re.match(r"^[A-Z][^.]{1,80},\s*[A-Z]\.", text))


def _looks_like_table_note(text: str) -> bool:
    return bool(re.match(r"^note\.\s+", _normalize(text))) and "table" in _normalize(text)


def _looks_like_figure_note(text: str) -> bool:
    return bool(re.match(r"^note\.\s+", _normalize(text))) and "figure" in _normalize(text)


def _word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9']+", text))


def _has_quote_marks(text: str) -> bool:
    return any(q in text for q in ('"', "“", "”", "'", "‘", "’"))


def _looks_like_trailing_citation(text: str) -> bool:
    return bool(re.search(r"\([^)]*\d{4}[a-z]?[^)]*\)\s*$", text))


def _left_indent_inches(paragraph: Paragraph) -> Optional[float]:
    try:
        left = paragraph.paragraph_format.left_indent
        if left is None:
            return None
        return left.inches
    except Exception:
        return None


def _score_block_quote(item: ClassifiedParagraph, prev: Optional[ClassifiedParagraph]) -> tuple[bool, float, tuple[str, ...]]:
    if item.region not in {Region.BODY, Region.BLOCK_QUOTE, Region.UNKNOWN}:
        return False, 0.0, ()
    text = item.text.strip()
    if not text:
        return False, 0.0, ()

    evidence: list[str] = []
    score = 0.0
    words = _word_count(text)
    indent = _left_indent_inches(item.paragraph)

    if words >= 40:
        score += 0.35
        evidence.append("word_count_40+")
    if indent is not None and indent >= 0.35:
        score += 0.35
        evidence.append("left_indent")
    if _looks_like_trailing_citation(text):
        score += 0.2
        evidence.append("trailing_citation")
    if _has_quote_marks(text):
        # Long quotes should typically omit surrounding quote marks in APA block form.
        score -= 0.15
        evidence.append("has_quote_marks")
    if prev is not None and prev.region == Region.BODY and words >= 40:
        score += 0.1
        evidence.append("follows_body")
    style = (item.style_name or "").lower()
    if "quote" in style or "blockquote" in style:
        score += 0.25
        evidence.append("quote_style")

    # Extremely high confidence only with multiple converging signals.
    if (
        words >= 40
        and _looks_like_trailing_citation(text)
        and indent is not None
        and indent >= 0.35
        and not _has_quote_marks(text)
    ):
        score = max(score, 0.99)
        evidence.append("convergent_block_signals")

    confidence = max(0.0, min(0.995, score))
    is_candidate = confidence >= 0.7 and words >= 40
    return is_candidate, confidence, tuple(evidence)


def _apply_block_quote_pass(
    classified: list[ClassifiedParagraph],
) -> list[ClassifiedParagraph]:
    updated: list[ClassifiedParagraph] = []
    for i, item in enumerate(classified):
        prev = classified[i - 1] if i > 0 else None
        is_cand, conf, ev = _score_block_quote(item, prev)
        if is_cand and conf >= 0.99 and item.region == Region.BODY:
            updated.append(
                ClassifiedParagraph(
                    index=item.index,
                    paragraph=item.paragraph,
                    region=Region.BLOCK_QUOTE,
                    confidence=conf,
                    text=item.text,
                    style_name=item.style_name,
                    heading_level=item.heading_level,
                    evidence=tuple([*item.evidence, *ev, "block_quote"]),
                    title_element=item.title_element,
                    is_block_quote_candidate=True,
                    block_quote_confidence=conf,
                )
            )
        elif is_cand:
            updated.append(
                ClassifiedParagraph(
                    index=item.index,
                    paragraph=item.paragraph,
                    region=item.region,
                    confidence=item.confidence,
                    text=item.text,
                    style_name=item.style_name,
                    heading_level=item.heading_level,
                    evidence=tuple([*item.evidence, *ev]),
                    title_element=item.title_element,
                    is_block_quote_candidate=True,
                    block_quote_confidence=conf,
                )
            )
        else:
            updated.append(item)
    return updated


def _looks_like_table_caption(text: str) -> bool:
    return bool(re.match(r"^table\s+\d+", _normalize(text)))


def _looks_like_figure_caption(text: str) -> bool:
    return bool(re.match(r"^figure\s+\d+", _normalize(text)))
