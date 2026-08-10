from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from fastapi import HTTPException

from app.apa.engine.analyzer import analyze_document, analyze_document_path
from app.apa.engine.classifier import classify_document
from app.apa.engine.verifier import verify_fix_result
from app.apa.models import SPECIALIZED_REGIONS, Fixability, Region
from app.apa.parsing.docx_objects import (
    compare_document_preservation,
    snapshot_image_binaries,
)
from app.apa.profile import PROFILE_ID
from app.apa.registry import get_rule
from app.apa.registry.base import RuleContext
from app.apa.text_integrity import (
    TextIntegrityError,
    assert_text_integrity,
    snapshot_user_text,
)

logger = logging.getLogger(__name__)

# Generic body formatting rules must not apply to specialized regions.
_BODY_GENERIC_RULE_IDS = frozenset(
    {
        "APA7-GLOBAL-004",
        "APA7-GLOBAL-005",
        "APA7-GLOBAL-007",
        "APA7-GLOBAL-009",
        "APA7-GLOBAL-010",
    }
)


def _issue_conflicts_with_region(issue, classified) -> bool:
    """Drop generic BODY fixes when the target paragraph is specialized."""
    if issue.rule_id not in _BODY_GENERIC_RULE_IDS:
        return False
    if issue.paragraph_index is None:
        return False
    for item in classified:
        if item.index == issue.paragraph_index:
            if item.region in SPECIALIZED_REGIONS or item.region != Region.BODY:
                return True
            return False
    return False


def fix_document_path(
    input_path: str,
    output_path: str,
    template_id: str = PROFILE_ID,
) -> dict:
    """
    Production Fix pipeline (Phase 2A–2D).

    - Analyzes first
    - Applies ONLY SAFE_AUTO_FIX rules for detected issues
    - Never rewrites user-authored characters
    - Rejects output on text-integrity or package-preservation failure
    - Verifies remaining SAFE issues == 0 for production-supported rules
    """
    doc = Document(input_path)
    before_snapshot = snapshot_user_text(doc)
    before_images = snapshot_image_binaries(doc)

    analysis_before = analyze_document(doc, template_id=template_id)
    classified = classify_document(doc)
    context = RuleContext(doc=doc, classified=classified, template_id=template_id)

    fixes_applied = 0
    applied_rule_ids: list[str] = []

    for issue in analysis_before.safe_auto_fix:
        if issue.fixability != Fixability.SAFE_AUTO_FIX or not issue.can_fix:
            continue
        if _issue_conflicts_with_region(issue, classified):
            logger.info(
                "skipped_conflicting_body_rule rule=%s paragraph=%s",
                issue.rule_id,
                issue.paragraph_index,
            )
            continue

        rule = get_rule(issue.rule_id, template_id)
        if rule is None:
            continue
        if rule.fixability != Fixability.SAFE_AUTO_FIX:
            continue
        if not rule.production_supported or not rule.source.source_verified:
            continue
        if issue.confidence < rule.minimum_confidence:
            continue

        changed = rule.apply_fix(context, issue)
        if changed:
            fixes_applied += 1
            applied_rule_ids.append(issue.rule_id)

    after_snapshot = snapshot_user_text(doc)
    text_ok = True
    try:
        assert_text_integrity(before_snapshot, after_snapshot)
    except TextIntegrityError as exc:
        text_ok = False
        logger.error(
            "fix_rejected_text_integrity regions=%s",
            getattr(exc, "changed_regions", [])[:50],
        )
        raise HTTPException(
            status_code=500,
            detail=(
                "Formatting aborted to protect your document text. "
                "No modified file was saved."
            ),
        ) from exc

    # Save candidate only after text integrity; then verify package preservation.
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    after_doc = Document(output_path)
    after_images = snapshot_image_binaries(after_doc)
    preservation = compare_document_preservation(
        before_path=input_path,
        after_path=output_path,
        before_images=before_images,
        after_images=after_images,
        before_doc=Document(input_path),
        after_doc=after_doc,
        text_integrity_ok=text_ok,
    )

    if not preservation.document_preservation_ok:
        try:
            Path(output_path).unlink(missing_ok=True)
        except Exception:
            pass
        logger.error(
            "fix_rejected_preservation warnings=%s",
            preservation.warnings[:50],
        )
        raise HTTPException(
            status_code=500,
            detail=(
                "Formatting aborted to protect embedded document objects "
                "(images, links, notes, equations, or numbering). "
                "No modified file was saved."
            ),
        )

    verification = verify_fix_result(
        output_path=output_path,
        template_id=template_id,
        safe_issues_before=analysis_before.safe_fix_count,
        fixes_applied=fixes_applied,
        author_action_issues=analysis_before.author_action_count,
        text_integrity_ok=text_ok,
        document_preservation_ok=preservation.document_preservation_ok,
        preservation_warnings=preservation.warnings,
    )

    if not verification.verified:
        try:
            Path(output_path).unlink(missing_ok=True)
        except Exception:
            pass
        raise HTTPException(
            status_code=500,
            detail=(
                "Formatting verification failed. Remaining auto-fixable issues "
                "were detected after fix; no downloadable file was produced."
            ),
        )

    fixed_counts = {
        "safe_rules_applied": fixes_applied,
        "applied_rule_ids": len(set(applied_rule_ids)),
        "margins": sum(1 for r in applied_rule_ids if r == "APA7-GLOBAL-001"),
        "line_spacing": sum(
            1
            for r in applied_rule_ids
            if r in {"APA7-GLOBAL-004", "APA7-REFERENCE-005", "APA7-TITLE-004"}
        ),
        "paragraph_indent": sum(1 for r in applied_rule_ids if r == "APA7-GLOBAL-010"),
        "space_before": sum(
            1 for r in applied_rule_ids if r in {"APA7-GLOBAL-005", "APA7-REFERENCE-006"}
        ),
        "space_after": sum(
            1 for r in applied_rule_ids if r in {"APA7-GLOBAL-005", "APA7-REFERENCE-006"}
        ),
        "reference_hanging_indent": sum(
            1 for r in applied_rule_ids if r == "APA7-REFERENCE-008"
        ),
        "page_numbers": sum(1 for r in applied_rule_ids if r == "APA7-GLOBAL-011"),
        "headings": sum(1 for r in applied_rule_ids if "HEADING" in r),
        "tables": sum(1 for r in applied_rule_ids if r.startswith("APA7-TABLE")),
        "figures": sum(1 for r in applied_rule_ids if r.startswith("APA7-FIGURE")),
        "appendices": sum(1 for r in applied_rule_ids if r.startswith("APA7-APPENDIX")),
    }

    analysis_after = analyze_document_path(output_path, template_id=template_id)

    return {
        "fixed_counts": fixed_counts,
        "verification": verification.to_dict(),
        "preservation": preservation.to_dict(),
        "safe_auto_fix_before": [i.to_dict() for i in analysis_before.safe_auto_fix],
        "author_action_required": [
            i.to_dict() for i in analysis_before.author_action_required
        ],
        "safe_auto_fix_after": [i.to_dict() for i in analysis_after.safe_auto_fix],
        "engine": "apa_phase_2g",
    }
