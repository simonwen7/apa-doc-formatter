"""Deterministic user-text snapshot and integrity checks.

Compares user-authored characters only. Formatting metadata may change.
Field instructions and PAGE/NUMPAGES field results are excluded so safe
page-number formatting does not fail integrity.
Does not clear/rebuild paragraphs.
"""

from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

# Documented python-docx limitations for Phase 2A/2B:
# - Text boxes / floating shapes: not reliably enumerable.
# - Footnotes/endnotes: partial; we attempt XML walk when present.
# - True pagination / title vertical position: not available without layout engine.
# - Field codes: PAGE/NUMPAGES results excluded from snapshot intentionally.


@dataclass(frozen=True)
class TextSnapshot:
    segments: tuple[tuple[str, str], ...]  # (region_label, text)
    fingerprint: str

    @property
    def joined_text(self) -> str:
        return "\n".join(text for _, text in self.segments)


def _ancestor_fldsimple_is_page(node) -> bool:
    parent = node.getparent()
    while parent is not None:
        if parent.tag == qn("w:fldSimple"):
            instr = (parent.get(qn("w:instr")) or "").upper()
            return "PAGE" in instr or "NUMPAGES" in instr
        parent = parent.getparent()
    return False


def paragraph_user_text(paragraph: Paragraph) -> str:
    """
    Extract user-authored characters from a paragraph.

    Excludes:
    - w:instrText (field instructions)
    - text runs that are PAGE / NUMPAGES field results
    """
    try:
        p_elm = paragraph._p
    except Exception:
        return paragraph.text or ""

    parts: list[str] = []
    in_field = False
    field_is_page = False

    # Document-order walk of direct logical runs via iterative scan of children tree
    # using iter() but skipping PAGE field results.
    for node in p_elm.iter():
        tag = node.tag
        if tag == qn("w:fldChar"):
            fld_type = node.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_field = True
                field_is_page = False
            elif fld_type == "end":
                in_field = False
                field_is_page = False
            continue

        if tag == qn("w:instrText"):
            instr = (node.text or "").strip().upper()
            if in_field and ("PAGE" in instr or "NUMPAGES" in instr):
                field_is_page = True
            continue

        if tag == qn("w:t"):
            if in_field and field_is_page:
                continue
            if _ancestor_fldsimple_is_page(node):
                continue
            if node.text:
                parts.append(node.text)

    return "".join(parts)

def _header_footer_texts(doc: Document) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = []
    for index, section in enumerate(doc.sections):
        try:
            for paragraph in section.header.paragraphs:
                segments.append((f"header[{index}]", paragraph_user_text(paragraph)))
        except Exception:
            logger.info("text_integrity: unable to read header for section %s", index)
        try:
            for paragraph in section.footer.paragraphs:
                segments.append((f"footer[{index}]", paragraph_user_text(paragraph)))
        except Exception:
            logger.info("text_integrity: unable to read footer for section %s", index)
    return segments


def _table_texts(doc: Document) -> list[tuple[str, str]]:
    segments: list[tuple[str, str]] = []
    for t_index, table in enumerate(doc.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for p_index, paragraph in enumerate(cell.paragraphs):
                    segments.append(
                        (
                            f"table[{t_index}].r{r_index}.c{c_index}.p{p_index}",
                            paragraph_user_text(paragraph),
                        )
                    )
    return segments


def _iter_note_roots(doc: Document, reltype_substring: str):
    """Yield note XML roots via relationships (footnotes_part may be absent)."""
    from lxml import etree

    try:
        for rel in doc.part.rels.values():
            if reltype_substring not in (rel.reltype or "").lower():
                continue
            part = rel.target_part
            element = getattr(part, "element", None)
            if element is not None:
                yield element
                return
            blob = getattr(part, "blob", None)
            if blob:
                yield etree.fromstring(blob)
                return
    except Exception:
        pass
    try:
        if reltype_substring == "footnotes":
            part = getattr(doc.part, "footnotes_part", None)
            if part is not None:
                yield part.element
        elif reltype_substring == "endnotes":
            part = getattr(doc.part, "endnotes_part", None)
            if part is not None:
                yield part.element
    except Exception:
        return


def _footnote_endnote_texts(doc: Document) -> list[tuple[str, str]]:
    """Best-effort extraction; python-docx has limited footnote APIs.

    Labels prefer stable w:id values when present so Fix elsewhere cannot
    silently drop note identity without failing integrity.
    """
    segments: list[tuple[str, str]] = []
    try:
        _ = doc.element.body
    except Exception:
        return segments

    try:
        for root in _iter_note_roots(doc, "footnotes"):
            for index, footnote in enumerate(root.findall(qn("w:footnote"))):
                note_id = footnote.get(qn("w:id")) or str(index)
                texts = [
                    node.text or ""
                    for node in footnote.iter(qn("w:t"))
                    if node.text
                ]
                segments.append((f"footnote[id={note_id}]", "".join(texts)))
    except Exception:
        logger.info(
            "text_integrity: footnotes not fully supported by python-docx in this environment"
        )

    try:
        for root in _iter_note_roots(doc, "endnotes"):
            for index, endnote in enumerate(root.findall(qn("w:endnote"))):
                note_id = endnote.get(qn("w:id")) or str(index)
                texts = [
                    node.text or ""
                    for node in endnote.iter(qn("w:t"))
                    if node.text
                ]
                segments.append((f"endnote[id={note_id}]", "".join(texts)))
    except Exception:
        logger.info(
            "text_integrity: endnotes not fully supported by python-docx in this environment"
        )

    return segments


def snapshot_user_text(doc: Document) -> TextSnapshot:
    segments: list[tuple[str, str]] = []

    for index, paragraph in enumerate(doc.paragraphs):
        segments.append((f"paragraph[{index}]", paragraph_user_text(paragraph)))

    segments.extend(_table_texts(doc))
    segments.extend(_header_footer_texts(doc))
    segments.extend(_footnote_endnote_texts(doc))

    frozen = tuple(segments)
    digest = hashlib.sha256(
        "\n".join(f"{label}\t{text}" for label, text in frozen).encode("utf-8")
    ).hexdigest()
    return TextSnapshot(segments=frozen, fingerprint=digest)


def assert_text_integrity(
    before: TextSnapshot,
    after: TextSnapshot,
) -> None:
    if before.fingerprint == after.fingerprint and before.segments == after.segments:
        return

    changed_regions: list[str] = []
    before_map = {label: text for label, text in before.segments}
    after_map = {label: text for label, text in after.segments}

    for label in sorted(set(before_map) | set(after_map)):
        if before_map.get(label) != after_map.get(label):
            changed_regions.append(label)

    logger.error(
        "text_integrity_failed regions=%s before_fp=%s after_fp=%s",
        changed_regions[:50],
        before.fingerprint,
        after.fingerprint,
    )
    raise TextIntegrityError(
        "User-authored text changed during formatting.",
        changed_regions=changed_regions,
    )


class TextIntegrityError(Exception):
    def __init__(self, message: str, changed_regions: Optional[list[str]] = None):
        super().__init__(message)
        self.changed_regions = changed_regions or []
