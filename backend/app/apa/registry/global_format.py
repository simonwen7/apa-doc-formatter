from __future__ import annotations

from docx.shared import Inches

from app.apa.models import Fixability, Issue, Region, Severity, SourceRef
from app.apa.profile import MARGIN_INCHES
from app.apa.registry.base import BaseRule, RuleContext
from app.apa.registry.formatting import (
    clear_extra_paragraph_spacing,
    first_line_indent_inches,
    is_apa_valid_typography,
    is_fully_justified,
    is_left_aligned_or_default,
    line_spacing_is_double,
    set_double_spacing,
    set_first_line_indent,
    set_left_aligned,
    set_runs_apa_font,
    space_after_pt,
    space_before_pt,
)

SOURCE_MARGINS = SourceRef(
    official_resource="APA Style Student Paper Setup Guide; Publication Manual (7th ed.), Paper Format",
    publication_manual_section=None,
    notes="1-inch margins on all sides for student papers.",
    source_verified=True,
)

SOURCE_SPACING = SourceRef(
    official_resource="APA Style Student Paper Setup Guide; Publication Manual (7th ed.), Line Spacing",
    publication_manual_section=None,
    notes="Double-space the entire paper; do not add extra space before/after paragraphs.",
    source_verified=True,
)

SOURCE_ALIGN = SourceRef(
    official_resource="APA Style Student Paper Setup Guide; Publication Manual (7th ed.), Paragraph Alignment",
    publication_manual_section=None,
    notes="Align text to the left; leave the right edge ragged; do not fully justify.",
    source_verified=True,
)

SOURCE_INDENT = SourceRef(
    official_resource="APA Style Student Paper Setup Guide; Publication Manual (7th ed.), Paragraph Indentation",
    publication_manual_section=None,
    notes="Indent the first line of every paragraph 0.5 in.",
    source_verified=True,
)

SOURCE_FONT = SourceRef(
    official_resource="APA Style Student Paper Setup Guide; Publication Manual (7th ed.), Font",
    publication_manual_section=None,
    notes=(
        "Use an APA-accessible font (e.g., Times New Roman 12, Arial 11, Calibri 11). "
        "Font family/size changes are formatting-only."
    ),
    source_verified=True,
)


def _detect_margins(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for index, section in enumerate(context.doc.sections):
        for side, value in (
            ("top", section.top_margin),
            ("bottom", section.bottom_margin),
            ("left", section.left_margin),
            ("right", section.right_margin),
        ):
            inches = value.inches if value else None
            if inches is None or abs(inches - MARGIN_INCHES) > 0.01:
                issues.append(
                    Issue(
                        rule_id="APA7-GLOBAL-001",
                        category="Page Format",
                        message=f"Section {index} {side} margin should be 1 inch.",
                        expected="1 in",
                        actual="unknown" if inches is None else f"{inches:.3f} in",
                        location=f"section[{index}].{side}_margin",
                        severity=Severity.ERROR,
                        fixability=Fixability.SAFE_AUTO_FIX,
                        can_fix=True,
                        confidence=0.99,
                        source=SOURCE_MARGINS,
                        section_index=index,
                        region=Region.UNKNOWN,
                    )
                )
    return issues


def _fix_margins(context: RuleContext, issue: Issue) -> bool:
    if issue.section_index is None:
        return False
    section = context.doc.sections[issue.section_index]
    section.top_margin = Inches(MARGIN_INCHES)
    section.bottom_margin = Inches(MARGIN_INCHES)
    section.left_margin = Inches(MARGIN_INCHES)
    section.right_margin = Inches(MARGIN_INCHES)
    return True


def _body_paras(context: RuleContext, min_confidence: float = 0.85):
    for item in context.classified:
        if item.region == Region.BODY and item.confidence >= min_confidence and item.text.strip():
            yield item


def _detect_double_spacing(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for item in _body_paras(context):
        result = line_spacing_is_double(item.paragraph)
        if result is False:
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-004",
                    category="Body",
                    message=f"Paragraph {item.index} should be double-spaced.",
                    expected="double (2.0)",
                    actual="not double (effective)",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.96),
                    source=SOURCE_SPACING,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
        elif result is None:
            # Unresolved inherited spacing is NOT compliant — do not silent-pass.
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-004-UNRESOLVED",
                    category="Body",
                    message=(
                        f"Paragraph {item.index} line spacing could not be "
                        "verified from direct or style inheritance."
                    ),
                    expected="verified double (2.0)",
                    actual="unresolved effective spacing",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.WARNING,
                    fixability=Fixability.CONDITIONAL,
                    can_fix=False,
                    confidence=min(item.confidence, 0.7),
                    source=SOURCE_SPACING,
                    reason_not_fixable=(
                        "Effective line spacing could not be resolved confidently."
                    ),
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
    return issues


def _fix_double_spacing(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_double_spacing(context.doc.paragraphs[issue.paragraph_index])
    return True


def _detect_extra_spacing(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for item in _body_paras(context):
        before = space_before_pt(item.paragraph)
        after = space_after_pt(item.paragraph)
        if before > 0.5 or after > 0.5:
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-005",
                    category="Body",
                    message=(
                        f"Paragraph {item.index} should not include extra space "
                        "before or after the paragraph."
                    ),
                    expected="0 pt before/after",
                    actual=f"{before:.1f}/{after:.1f} pt",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.96),
                    source=SOURCE_SPACING,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
    return issues


def _fix_extra_spacing(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    clear_extra_paragraph_spacing(context.doc.paragraphs[issue.paragraph_index])
    return True


def _detect_left_align(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for item in _body_paras(context):
        if not is_left_aligned_or_default(item.paragraph):
            if is_fully_justified(item.paragraph):
                continue  # handled by GLOBAL-009
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-007",
                    category="Body",
                    message=f"Paragraph {item.index} should be left-aligned.",
                    expected="left",
                    actual=str(item.paragraph.alignment),
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.96),
                    source=SOURCE_ALIGN,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
    return issues


def _fix_left_align(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_left_aligned(context.doc.paragraphs[issue.paragraph_index])
    return True


def _detect_not_justified(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for item in _body_paras(context):
        if is_fully_justified(item.paragraph):
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-009",
                    category="Body",
                    message=f"Paragraph {item.index} should not be fully justified.",
                    expected="left / ragged right",
                    actual="justify",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.97),
                    source=SOURCE_ALIGN,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
    return issues


def _fix_not_justified(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_left_aligned(context.doc.paragraphs[issue.paragraph_index])
    return True


def _detect_first_line_indent(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for item in _body_paras(context):
        indent = first_line_indent_inches(item.paragraph)
        if indent is None:
            # Uncertain inheritance — do not definite-error in Phase 2A
            # unless alignment suggests flush body text with no indent set.
            # Still report with slightly lower confidence when text is body.
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-010",
                    category="Body",
                    message=(
                        f"Paragraph {item.index} first-line indent should be 0.5 inch."
                    ),
                    expected="0.5 in",
                    actual="not set",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.95),
                    source=SOURCE_INDENT,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
        elif abs(indent - 0.5) > 0.01:
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-010",
                    category="Body",
                    message=(
                        f"Paragraph {item.index} first-line indent should be 0.5 inch."
                    ),
                    expected="0.5 in",
                    actual=f"{indent:.3f} in",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.97),
                    source=SOURCE_INDENT,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
    return issues


def _fix_first_line_indent(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_first_line_indent(context.doc.paragraphs[issue.paragraph_index])
    return True


def _detect_body_font(context: RuleContext) -> list[Issue]:
    issues: list[Issue] = []
    for item in _body_paras(context):
        validity = is_apa_valid_typography(item.paragraph)
        # Only report confident violations. Unresolved theme/default fonts are
        # common in python-docx and must not flood CONDITIONAL noise; spacing
        # unresolved rules already prevent false score-100 for inheritance gaps.
        if validity is False:
            issues.append(
                Issue(
                    rule_id="APA7-GLOBAL-FONT",
                    category="Body",
                    message=(
                        f"Paragraph {item.index} uses a font that is not on the "
                        "APA-accessible allowlist."
                    ),
                    expected="APA-accessible font (e.g., Times New Roman 12)",
                    actual="invalid effective font",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=min(item.confidence, 0.95),
                    source=SOURCE_FONT,
                    paragraph_index=item.index,
                    region=Region.BODY,
                )
            )
    return issues


def _fix_body_font(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_runs_apa_font(context.doc.paragraphs[issue.paragraph_index])
    return True


SOURCE_PAGE_NUMBERS = SourceRef(
    official_resource=(
        "APA Style Student Paper Setup Guide; Publication Manual (7th ed.) "
        "page headers / page numbers guidance"
    ),
    publication_manual_section=None,
    notes=(
        "Put a page number in the top right corner of every page, including the "
        "title page (page 1). Student papers do not require a running head."
    ),
    source_verified=True,
)


def _detect_page_numbers(context: RuleContext) -> list[Issue]:
    from app.apa.registry.page_numbers import effective_section_has_page_field

    issues: list[Issue] = []
    for index, _section in enumerate(context.doc.sections):
        if effective_section_has_page_field(context.doc, index):
            continue
        issues.append(
            Issue(
                rule_id="APA7-GLOBAL-011",
                category="Page Format",
                message=(
                    f"Section {index} is missing a page number in the upper-right header."
                ),
                expected="PAGE field, top-right header",
                actual="no PAGE field detected",
                location=f"section[{index}].header",
                severity=Severity.ERROR,
                fixability=Fixability.SAFE_AUTO_FIX,
                can_fix=True,
                confidence=0.98,
                source=SOURCE_PAGE_NUMBERS,
                section_index=index,
                region=Region.HEADER,
            )
        )
    return issues


def _detect_running_head(context: RuleContext) -> list[Issue]:
    from app.apa.registry.page_numbers import detect_non_page_header_text

    issues: list[Issue] = []
    for index, section in enumerate(context.doc.sections):
        extra = detect_non_page_header_text(section)
        if not extra:
            continue
        issues.append(
            Issue(
                rule_id="APA7-GLOBAL-RUNNING-HEAD",
                category="Page Format",
                message=(
                    "Header contains text in addition to a page number. Student "
                    "papers do not require a running head unless an instructor "
                    "requests one."
                ),
                expected="page number only (student default)",
                actual="additional header text present",
                location=f"section[{index}].header",
                severity=Severity.WARNING,
                fixability=Fixability.AUTHOR_ACTION_REQUIRED,
                can_fix=False,
                confidence=0.8,
                source=SOURCE_PAGE_NUMBERS,
                reason_not_fixable=(
                    "Removing or altering header text could destroy instructor-"
                    "required running head content."
                ),
                section_index=index,
                region=Region.HEADER,
            )
        )
    return issues


def _fix_page_numbers(context: RuleContext, issue: Issue) -> bool:
    from app.apa.registry.page_numbers import ensure_page_number_in_section

    if issue.section_index is None:
        return False
    section = context.doc.sections[issue.section_index]
    return ensure_page_number_in_section(section, issue.section_index)


def _detect_title_page_number_one(context: RuleContext) -> list[Issue]:
    from docx.oxml.ns import qn

    issues: list[Issue] = []
    if not context.doc.sections:
        return issues
    section = context.doc.sections[0]
    try:
        pg_num = section._sectPr.find(qn("w:pgNumType"))
        start = pg_num.get(qn("w:start")) if pg_num is not None else None
        if start is None or start == "1":
            return issues
        issues.append(
            Issue(
                rule_id="APA7-GLOBAL-012",
                category="Page Format",
                message="Title page should be numbered as page 1.",
                expected="start=1",
                actual=f"start={start}",
                location="section[0].pgNumType",
                severity=Severity.ERROR,
                fixability=Fixability.SAFE_AUTO_FIX,
                can_fix=True,
                confidence=0.97,
                source=SOURCE_PAGE_NUMBERS,
                section_index=0,
                region=Region.HEADER,
            )
        )
    except Exception:
        pass
    return issues


def _fix_title_page_number_one(context: RuleContext, issue: Issue) -> bool:
    from app.apa.registry.page_numbers import ensure_title_page_is_page_one

    if issue.section_index is None:
        return False
    return ensure_title_page_is_page_one(context.doc.sections[issue.section_index])


RULES = [
    BaseRule(
        rule_id="APA7-GLOBAL-001",
        category="Page Format",
        description="All page margins must be 1 inch.",
        official_expectation="1-inch margins on all sides",
        source=SOURCE_MARGINS,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("DOCUMENT",),
        detector=_detect_margins,
        fixer=_fix_margins,
        minimum_confidence=0.99,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-004",
        category="Body",
        description="Ordinary paper text should be double-spaced.",
        official_expectation="Double-space the entire student paper",
        source=SOURCE_SPACING,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BODY",),
        detector=_detect_double_spacing,
        fixer=_fix_double_spacing,
        minimum_confidence=0.85,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-005",
        category="Body",
        description="Do not add extra space before or after paragraphs.",
        official_expectation="No extra before/after paragraph spacing",
        source=SOURCE_SPACING,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BODY",),
        detector=_detect_extra_spacing,
        fixer=_fix_extra_spacing,
        minimum_confidence=0.85,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-007",
        category="Body",
        description="Body text should be left-aligned.",
        official_expectation="Left align text; ragged right",
        source=SOURCE_ALIGN,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BODY",),
        detector=_detect_left_align,
        fixer=_fix_left_align,
        minimum_confidence=0.85,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-009",
        category="Body",
        description="Body text should not be fully justified.",
        official_expectation="Do not fully justify text",
        source=SOURCE_ALIGN,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BODY",),
        detector=_detect_not_justified,
        fixer=_fix_not_justified,
        minimum_confidence=0.85,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-010",
        category="Body",
        description="Body paragraph first-line indentation should be 0.5 inch.",
        official_expectation="0.5-inch first-line indent",
        source=SOURCE_INDENT,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BODY",),
        detector=_detect_first_line_indent,
        fixer=_fix_first_line_indent,
        minimum_confidence=0.85,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-011",
        category="Page Format",
        description="Student paper pages have page numbers in the upper-right header.",
        official_expectation="Top-right PAGE field on every page",
        source=SOURCE_PAGE_NUMBERS,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("HEADER",),
        detector=_detect_page_numbers,
        fixer=_fix_page_numbers,
        minimum_confidence=0.95,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-012",
        category="Page Format",
        description="Title page is page 1.",
        official_expectation="Section page numbering starts at 1",
        source=SOURCE_PAGE_NUMBERS,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("HEADER",),
        detector=_detect_title_page_number_one,
        fixer=_fix_title_page_number_one,
        minimum_confidence=0.95,
        production_supported=True,
        fixer_implemented=True,
        passing_fixture=True,
        failing_fixture=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-RUNNING-HEAD",
        category="Page Format",
        description="Student papers do not require a running head by default.",
        official_expectation="No required running head for student papers",
        source=SOURCE_PAGE_NUMBERS,
        fixability=Fixability.AUTHOR_ACTION_REQUIRED,
        applicable_regions=("HEADER",),
        detector=_detect_running_head,
        fixer=None,
        production_supported=True,
        fixer_implemented=False,
        detector_test=True,
    ),
    BaseRule(
        rule_id="APA7-GLOBAL-FONT",
        category="Body",
        description="Body text uses an APA-accessible font and size.",
        official_expectation="APA-accessible font (e.g., Times New Roman 12)",
        source=SOURCE_FONT,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BODY",),
        detector=_detect_body_font,
        fixer=_fix_body_font,
        minimum_confidence=0.85,
        production_supported=True,
        fixer_implemented=True,
        detector_test=True,
        fixer_test=True,
        post_fix_test=True,
        text_integrity_test=True,
        idempotency_test=True,
    ),
]
