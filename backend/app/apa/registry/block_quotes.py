"""Block quotation rules (Phase 2C). Formatting-only SAFE at confidence >= 0.99."""

from __future__ import annotations

import re

from docx.shared import Inches

from app.apa.models import Fixability, Issue, Region, Severity, SourceRef
from app.apa.registry.base import BaseRule, RuleContext
from app.apa.registry.formatting import (
    is_left_aligned_or_default,
    line_spacing_is_double,
    set_double_spacing,
    set_left_aligned,
    clear_extra_paragraph_spacing,
    space_after_pt,
    space_before_pt,
)

SOURCE_QUOTES = SourceRef(
    official_resource=(
        "APA Style guidance on quotations; Publication Manual (7th ed.) "
        "Works Credited in the Text / quotations"
    ),
    publication_manual_section=None,
    notes=(
        "Quotations of 40 words or more are formatted as block quotations: "
        "indented 0.5 in from the left margin, double-spaced, without quotation marks."
    ),
    source_verified=True,
)

BLOCK_SAFE_CONF = 0.99


def _word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9']+", text or ""))


def _detect_block_indent(context: RuleContext) -> list[Issue]:
    issues = []
    for item in context.classified:
        if not item.is_block_quote_candidate:
            continue
        if item.block_quote_confidence < BLOCK_SAFE_CONF:
            continue
        if item.region != Region.BLOCK_QUOTE:
            continue
        left = item.paragraph.paragraph_format.left_indent
        inches = left.inches if left is not None else None
        if inches is None or abs(inches - 0.5) > 0.05:
            issues.append(
                Issue(
                    rule_id="APA7-QUOTE-BLOCK-INDENT",
                    category="Quotations",
                    message=(
                        f"Paragraph {item.index} appears to be a block quotation and "
                        "should be indented 0.5 inch from the left margin."
                    ),
                    expected="0.5 in left indent",
                    actual="missing/incorrect indent",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=item.block_quote_confidence,
                    source=SOURCE_QUOTES,
                    paragraph_index=item.index,
                    region=Region.BLOCK_QUOTE,
                )
            )
    return issues


def _fix_block_indent(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    paragraph = context.doc.paragraphs[issue.paragraph_index]
    paragraph.paragraph_format.left_indent = Inches(0.5)
    return True


def _detect_block_spacing(context: RuleContext) -> list[Issue]:
    issues = []
    for item in context.classified:
        if item.region != Region.BLOCK_QUOTE or item.block_quote_confidence < BLOCK_SAFE_CONF:
            continue
        result = line_spacing_is_double(item.paragraph)
        if result is False:
            issues.append(
                Issue(
                    rule_id="APA7-QUOTE-BLOCK-SPACING",
                    category="Quotations",
                    message=f"Block quotation paragraph {item.index} should be double-spaced.",
                    expected="double (2.0)",
                    actual="not double",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=item.block_quote_confidence,
                    source=SOURCE_QUOTES,
                    paragraph_index=item.index,
                    region=Region.BLOCK_QUOTE,
                )
            )
        if space_before_pt(item.paragraph) > 0.5 or space_after_pt(item.paragraph) > 0.5:
            issues.append(
                Issue(
                    rule_id="APA7-QUOTE-BLOCK-PARA-SPACING",
                    category="Quotations",
                    message=(
                        f"Block quotation paragraph {item.index} should not include "
                        "extra paragraph spacing."
                    ),
                    expected="0 pt before/after",
                    actual="extra spacing",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=item.block_quote_confidence,
                    source=SOURCE_QUOTES,
                    paragraph_index=item.index,
                    region=Region.BLOCK_QUOTE,
                )
            )
        if not is_left_aligned_or_default(item.paragraph):
            issues.append(
                Issue(
                    rule_id="APA7-QUOTE-BLOCK-ALIGN",
                    category="Quotations",
                    message=f"Block quotation paragraph {item.index} should be left-aligned.",
                    expected="left",
                    actual=str(item.paragraph.alignment),
                    location=f"paragraph[{item.index}]",
                    severity=Severity.ERROR,
                    fixability=Fixability.SAFE_AUTO_FIX,
                    can_fix=True,
                    confidence=item.block_quote_confidence,
                    source=SOURCE_QUOTES,
                    paragraph_index=item.index,
                    region=Region.BLOCK_QUOTE,
                )
            )
    return issues


def _fix_double(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_double_spacing(context.doc.paragraphs[issue.paragraph_index])
    return True


def _fix_para_spacing(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    clear_extra_paragraph_spacing(context.doc.paragraphs[issue.paragraph_index])
    return True


def _fix_align(context: RuleContext, issue: Issue) -> bool:
    if issue.paragraph_index is None:
        return False
    set_left_aligned(context.doc.paragraphs[issue.paragraph_index])
    return True


def _detect_quote_marks_on_long(context: RuleContext) -> list[Issue]:
    issues = []
    for item in context.classified:
        if not item.is_block_quote_candidate:
            continue
        if _word_count(item.text) < 40:
            continue
        if any(q in item.text for q in ('"', "“", "”")):
            issues.append(
                Issue(
                    rule_id="APA7-QUOTE-MARKS-LONG",
                    category="Quotations",
                    message=(
                        f"Paragraph {item.index} may be a long quotation that still uses "
                        "quotation marks. Block quotations typically omit surrounding "
                        "quotation marks."
                    ),
                    expected="block quotation without surrounding quotation marks",
                    actual="quotation marks present on long quotation candidate",
                    location=f"paragraph[{item.index}]",
                    severity=Severity.WARNING,
                    fixability=Fixability.AUTHOR_ACTION_REQUIRED,
                    can_fix=False,
                    confidence=min(item.block_quote_confidence, 0.9),
                    source=SOURCE_QUOTES,
                    reason_not_fixable="Removing quotation marks would modify author text.",
                    paragraph_index=item.index,
                    region=item.region,
                )
            )
    return issues


def _detect_low_confidence_block(context: RuleContext) -> list[Issue]:
    issues = []
    for item in context.classified:
        if not item.is_block_quote_candidate:
            continue
        if item.block_quote_confidence >= BLOCK_SAFE_CONF:
            continue
        issues.append(
            Issue(
                rule_id="APA7-QUOTE-BLOCK-UNCERTAIN",
                category="Quotations",
                message=(
                    f"Paragraph {item.index} may be a block quotation, but confidence "
                    "is not high enough to change formatting automatically."
                ),
                expected="confident block-quotation classification",
                actual=f"confidence={item.block_quote_confidence:.2f}",
                location=f"paragraph[{item.index}]",
                severity=Severity.INFO,
                fixability=Fixability.CONDITIONAL,
                can_fix=False,
                confidence=item.block_quote_confidence,
                source=SOURCE_QUOTES,
                reason_not_fixable="Low-confidence classification must not auto-fix.",
                paragraph_index=item.index,
                region=item.region,
            )
        )
    return issues


def _filter(rule_id: str, detector):
    def _detect(context: RuleContext) -> list[Issue]:
        return [i for i in detector(context) if i.rule_id == rule_id]

    return _detect


_SAFE = dict(
    production_supported=True,
    fixer_implemented=True,
    passing_fixture=True,
    failing_fixture=True,
    detector_test=True,
    fixer_test=True,
    post_fix_test=True,
    text_integrity_test=True,
    idempotency_test=True,
)

RULES = [
    BaseRule(
        rule_id="APA7-QUOTE-BLOCK-INDENT",
        category="Quotations",
        description="High-confidence block quotations use 0.5-inch left indent.",
        official_expectation="0.5 in left indent for block quotations",
        source=SOURCE_QUOTES,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BLOCK_QUOTE",),
        detector=_detect_block_indent,
        fixer=_fix_block_indent,
        minimum_confidence=BLOCK_SAFE_CONF,
        **_SAFE,
    ),
    BaseRule(
        rule_id="APA7-QUOTE-BLOCK-SPACING",
        category="Quotations",
        description="High-confidence block quotations are double-spaced.",
        official_expectation="Double-space block quotations",
        source=SOURCE_QUOTES,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BLOCK_QUOTE",),
        detector=_filter("APA7-QUOTE-BLOCK-SPACING", _detect_block_spacing),
        fixer=_fix_double,
        minimum_confidence=BLOCK_SAFE_CONF,
        **_SAFE,
    ),
    BaseRule(
        rule_id="APA7-QUOTE-BLOCK-PARA-SPACING",
        category="Quotations",
        description="High-confidence block quotations have no extra paragraph spacing.",
        official_expectation="No extra before/after spacing",
        source=SOURCE_QUOTES,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BLOCK_QUOTE",),
        detector=_filter("APA7-QUOTE-BLOCK-PARA-SPACING", _detect_block_spacing),
        fixer=_fix_para_spacing,
        minimum_confidence=BLOCK_SAFE_CONF,
        **_SAFE,
    ),
    BaseRule(
        rule_id="APA7-QUOTE-BLOCK-ALIGN",
        category="Quotations",
        description="High-confidence block quotations are left-aligned.",
        official_expectation="Left-aligned block quotation",
        source=SOURCE_QUOTES,
        fixability=Fixability.SAFE_AUTO_FIX,
        applicable_regions=("BLOCK_QUOTE",),
        detector=_filter("APA7-QUOTE-BLOCK-ALIGN", _detect_block_spacing),
        fixer=_fix_align,
        minimum_confidence=BLOCK_SAFE_CONF,
        **_SAFE,
    ),
    BaseRule(
        rule_id="APA7-QUOTE-MARKS-LONG",
        category="Quotations",
        description="Long quotations typically omit surrounding quotation marks.",
        official_expectation="No surrounding quotation marks on block quotes",
        source=SOURCE_QUOTES,
        fixability=Fixability.AUTHOR_ACTION_REQUIRED,
        applicable_regions=("BODY", "BLOCK_QUOTE"),
        detector=_detect_quote_marks_on_long,
        fixer=None,
        production_supported=True,
        detector_test=True,
    ),
    BaseRule(
        rule_id="APA7-QUOTE-BLOCK-UNCERTAIN",
        category="Quotations",
        description="Low-confidence block-quote candidates are not auto-fixed.",
        official_expectation="Confident classification before formatting",
        source=SOURCE_QUOTES,
        fixability=Fixability.CONDITIONAL,
        applicable_regions=("BODY", "BLOCK_QUOTE"),
        detector=_detect_low_confidence_block,
        fixer=None,
        production_supported=False,
        detector_test=True,
    ),
]
