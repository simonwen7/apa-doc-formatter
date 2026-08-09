"""Safe PAGE field helpers for student-paper headers.

Never clears headers, never rewrites existing user-authored header text,
never deletes footer content, never duplicates PAGE fields.
"""

from __future__ import annotations

from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Twips
from docx.text.paragraph import Paragraph

from app.apa.registry.formatting import ensure_right_tab_stop
from app.apa.text_integrity import paragraph_user_text


def _paragraph_has_page_field(paragraph: Paragraph) -> bool:
    try:
        p_elm = paragraph._p
    except Exception:
        return False
    import re

    for node in p_elm.iter():
        if node.tag == qn("w:instrText"):
            instr = (node.text or "").upper()
            if re.search(r"\bPAGE\b", instr):
                return True
        if node.tag == qn("w:fldSimple"):
            instr = (node.get(qn("w:instr")) or "").upper()
            if re.search(r"\bPAGE\b", instr):
                return True
    return False


def section_has_page_field(section) -> bool:
    try:
        header = section.header
    except Exception:
        return False
    for paragraph in header.paragraphs:
        if _paragraph_has_page_field(paragraph):
            return True
    return False


def count_page_fields_in_header(section) -> int:
    count = 0
    try:
        header = section.header
    except Exception:
        return 0
    for paragraph in header.paragraphs:
        if _paragraph_has_page_field(paragraph):
            count += 1
    return count


def effective_section_has_page_field(doc: Document, section_index: int) -> bool:
    """Walk linked-to-previous chain to find an inherited PAGE field."""
    sections = list(doc.sections)
    index = section_index
    seen = set()
    while index >= 0 and index not in seen:
        seen.add(index)
        section = sections[index]
        if section_has_page_field(section):
            return True
        try:
            if index > 0 and section.header.is_linked_to_previous:
                index -= 1
                continue
        except Exception:
            pass
        break
    return False


def _append_page_field_runs(paragraph: Paragraph) -> None:
    """Append a PAGE field to an existing paragraph without clearing runs."""

    def _fld_char(char_type: str) -> None:
        run = paragraph.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), char_type)
        run._r.append(fld)

    _fld_char("begin")
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    _fld_char("separate")
    paragraph.add_run("")  # empty result; display filled by Word
    _fld_char("end")


def ensure_page_number_in_section(section, section_index: int = 0) -> bool:
    """
    Ensure a top-right PAGE field exists for this section's header.

    Returns True if a change was made.
    Preserves all existing header paragraph user text.
    Does nothing if a PAGE field already exists (here or via link-to-previous).
    """
    header = section.header

    # If linked to previous and previous provides PAGE, do not add another.
    try:
        if section_index > 0 and header.is_linked_to_previous:
            return False
    except Exception:
        pass

    if section_has_page_field(section):
        return False

    # Prefer first header paragraph; create one only if header has none.
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        paragraph = header.add_paragraph()

    user_text = paragraph_user_text(paragraph).strip()
    if not user_text:
        # Empty header paragraph: right-align and add PAGE field.
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _append_page_field_runs(paragraph)
        return True

    # Existing user text: keep it, add a right-tab + PAGE field on same paragraph.
    ensure_right_tab_stop(paragraph)
    paragraph.add_run("\t")
    _append_page_field_runs(paragraph)
    return True


def ensure_title_page_is_page_one(section) -> bool:
    """Set section start page number to 1 when safe. Returns True if changed."""
    try:
        # python-docx: section.start_type / page_number related via sectPr
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        start = pg_num.get(qn("w:start"))
        if start == "1":
            return False
        pg_num.set(qn("w:start"), "1")
        return True
    except Exception:
        return False


def detect_non_page_header_text(section) -> Optional[str]:
    """Return truncated label if header has user text beyond PAGE fields."""
    try:
        header = section.header
    except Exception:
        return None
    for paragraph in header.paragraphs:
        text = paragraph_user_text(paragraph).strip()
        if text:
            return text[:40]
    return None
