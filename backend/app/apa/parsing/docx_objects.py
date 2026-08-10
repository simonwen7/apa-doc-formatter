"""Complex DOCX package preservation checks (Phase 2D).

Compares critical OPC package objects before/after Fix without requiring
byte-identical whole documents. Formatting XML may change; semantic objects
must not disappear.
"""

from __future__ import annotations

import hashlib
import logging
import re
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


@dataclass
class DocumentPreservationReport:
    text_integrity_ok: bool = True
    images_preserved: bool = True
    hyperlinks_preserved: bool = True
    bookmarks_preserved: bool = True
    footnotes_preserved: bool = True
    endnotes_preserved: bool = True
    equations_preserved: bool = True
    relationships_preserved: bool = True
    embedded_objects_preserved: bool = True
    numbering_preserved: bool = True
    document_preservation_ok: bool = True
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return asdict(self)


def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def snapshot_image_binaries(doc: Document) -> dict[str, str]:
    """Map image partname -> sha256 of binary."""
    images: dict[str, str] = {}
    try:
        for rel in doc.part.rels.values():
            if "image" in (rel.reltype or "").lower():
                try:
                    blob = rel.target_part.blob
                    images[str(rel.rId)] = _hash_bytes(blob)
                except Exception:
                    images[str(rel.rId)] = "unreadable"
    except Exception:
        logger.info("preservation: unable to enumerate image relationships")
    return images


def snapshot_package_parts(path: str) -> dict[str, str]:
    """Hash selected zip parts for package-level regression checks."""
    interesting_prefixes = (
        "word/media/",
        "word/embeddings/",
        "word/charts/",
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/numbering.xml",
        "word/embeddings",
        "customXml/",
    )
    hashes: dict[str, str] = {}
    try:
        with zipfile.ZipFile(path, "r") as zf:
            for name in zf.namelist():
                if name.startswith(interesting_prefixes) or name in {
                    "word/footnotes.xml",
                    "word/endnotes.xml",
                    "word/numbering.xml",
                }:
                    hashes[name] = _hash_bytes(zf.read(name))
    except Exception:
        logger.info("preservation: unable to read package %s", Path(path).name)
    return hashes


def count_hyperlinks(doc: Document) -> int:
    count = 0
    try:
        for paragraph in doc.paragraphs:
            count += len(paragraph._p.findall(".//" + qn("w:hyperlink")))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count += len(paragraph._p.findall(".//" + qn("w:hyperlink")))
    except Exception:
        return -1
    return count


def count_bookmarks(doc: Document) -> int:
    try:
        return len(doc.element.body.findall(".//" + qn("w:bookmarkStart")))
    except Exception:
        return -1


def count_equations(doc: Document) -> int:
    """Count OMML equation roots best-effort."""
    try:
        # Office Math ML namespace varies; match local-name via string search.
        xml = doc.element.body.xml
        return xml.count("oMath") + xml.count("oMathPara")
    except Exception:
        return -1


def snapshot_note_ids(path: str) -> dict[str, list[str]]:
    """Extract footnote/endnote w:id values from package XML (empty if absent)."""
    result = {"footnotes": [], "endnotes": []}
    try:
        with zipfile.ZipFile(path, "r") as zf:
            for part, key, tag in (
                ("word/footnotes.xml", "footnotes", "w:footnote"),
                ("word/endnotes.xml", "endnotes", "w:endnote"),
            ):
                if part not in zf.namelist():
                    continue
                data = zf.read(part).decode("utf-8", errors="ignore")
                # Capture id attributes without logging note text.
                ids = re.findall(rf"<{tag}[^>]*\bw:id=\"(\d+)\"", data)
                if not ids:
                    ids = re.findall(rf"<{tag}[^>]*\bid=\"(\d+)\"", data)
                result[key] = ids
    except Exception:
        logger.info("preservation: unable to snapshot note ids for %s", Path(path).name)
    return result


def count_footnote_references(doc: Document) -> int:
    """Count footnote reference marks in the document body."""
    try:
        return len(doc.element.body.findall(".//" + qn("w:footnoteReference")))
    except Exception:
        return -1


def count_endnote_references(doc: Document) -> int:
    try:
        return len(doc.element.body.findall(".//" + qn("w:endnoteReference")))
    except Exception:
        return -1


def count_footnotes_endnotes(doc: Document) -> tuple[int, int]:
    from lxml import etree

    footnotes = 0
    endnotes = 0
    try:
        found = False
        for rel in doc.part.rels.values():
            if "footnotes" not in (rel.reltype or "").lower():
                continue
            part = rel.target_part
            element = getattr(part, "element", None)
            if element is None and getattr(part, "blob", None):
                element = etree.fromstring(part.blob)
            if element is not None:
                footnotes = len(element.findall(qn("w:footnote")))
                found = True
                break
        if not found:
            part = getattr(doc.part, "footnotes_part", None)
            if part is not None:
                footnotes = len(part.element.findall(qn("w:footnote")))
    except Exception:
        footnotes = -1
    try:
        found = False
        for rel in doc.part.rels.values():
            if "endnotes" not in (rel.reltype or "").lower():
                continue
            part = rel.target_part
            element = getattr(part, "element", None)
            if element is None and getattr(part, "blob", None):
                element = etree.fromstring(part.blob)
            if element is not None:
                endnotes = len(element.findall(qn("w:endnote")))
                found = True
                break
        if not found:
            part = getattr(doc.part, "endnotes_part", None)
            if part is not None:
                endnotes = len(part.element.findall(qn("w:endnote")))
    except Exception:
        endnotes = -1
    return footnotes, endnotes


def assert_images_preserved(before: dict[str, str], after: dict[str, str]) -> None:
    if before != after:
        lost = sorted(set(before) - set(after))
        changed = sorted(
            k for k in set(before) & set(after) if before[k] != after[k]
        )
        raise DocumentPreservationError(
            "Embedded image binaries changed or were removed.",
            warnings=[f"lost={lost[:20]}", f"changed={changed[:20]}"],
        )


def _count_xml_tags_in_part(path: str, part_name: str, local_tag: str) -> int:
    try:
        with zipfile.ZipFile(path, "r") as zf:
            if part_name not in zf.namelist():
                return -1
            data = zf.read(part_name).decode("utf-8", errors="ignore")
            return data.count(f"<{local_tag}") + data.count(f":{local_tag}")
    except Exception:
        return -1


def compare_document_preservation(
    *,
    before_path: str,
    after_path: str,
    before_images: dict[str, str],
    after_images: dict[str, str],
    before_doc: Document,
    after_doc: Document,
    text_integrity_ok: bool,
) -> DocumentPreservationReport:
    report = DocumentPreservationReport(text_integrity_ok=text_integrity_ok)
    warnings: list[str] = []

    try:
        assert_images_preserved(before_images, after_images)
        report.images_preserved = True
    except DocumentPreservationError as exc:
        report.images_preserved = False
        warnings.extend(exc.warnings)

    before_parts = snapshot_package_parts(before_path)
    after_parts = snapshot_package_parts(after_path)
    # Media/embeddings must not disappear or change binary content.
    for name, digest in before_parts.items():
        if name.startswith("word/media/") or name.startswith("word/embeddings/"):
            if name not in after_parts:
                report.embedded_objects_preserved = False
                warnings.append(f"missing_part:{name}")
            elif after_parts[name] != digest:
                report.images_preserved = False
                report.embedded_objects_preserved = False
                warnings.append(f"changed_part:{name}")

    # Numbering: require part survival + stable num counts (not byte-identical XML).
    if "word/numbering.xml" in before_parts:
        if "word/numbering.xml" not in after_parts:
            report.numbering_preserved = False
            warnings.append("numbering_xml_missing")
        else:
            before_nums = _count_xml_tags_in_part(before_path, "word/numbering.xml", "w:num")
            after_nums = _count_xml_tags_in_part(after_path, "word/numbering.xml", "w:num")
            if before_nums >= 0 and after_nums >= 0 and after_nums < before_nums:
                report.numbering_preserved = False
                warnings.append("numbering_definitions_decreased")

    # Footnotes/endnotes: part survival + count stability (XML pretty-print may change).
    for part_name, flag, tag in (
        ("word/footnotes.xml", "footnotes_preserved", "w:footnote"),
        ("word/endnotes.xml", "endnotes_preserved", "w:endnote"),
    ):
        if part_name in before_parts:
            if part_name not in after_parts:
                setattr(report, flag, False)
                warnings.append(f"{part_name}_missing")
            else:
                b = _count_xml_tags_in_part(before_path, part_name, tag)
                a = _count_xml_tags_in_part(after_path, part_name, tag)
                if b >= 0 and a >= 0 and a < b:
                    setattr(report, flag, False)
                    warnings.append(f"{part_name}_count_decreased")

    before_note_ids = snapshot_note_ids(before_path)
    after_note_ids = snapshot_note_ids(after_path)
    if before_note_ids["footnotes"] != after_note_ids["footnotes"]:
        report.footnotes_preserved = False
        warnings.append("footnote_ids_changed")
    if before_note_ids["endnotes"] != after_note_ids["endnotes"]:
        report.endnotes_preserved = False
        warnings.append("endnote_ids_changed")

    before_fn_refs = count_footnote_references(before_doc)
    after_fn_refs = count_footnote_references(after_doc)
    if before_fn_refs >= 0 and after_fn_refs >= 0 and before_fn_refs != after_fn_refs:
        report.footnotes_preserved = False
        warnings.append("footnote_reference_count_changed")
    before_en_refs = count_endnote_references(before_doc)
    after_en_refs = count_endnote_references(after_doc)
    if before_en_refs >= 0 and after_en_refs >= 0 and before_en_refs != after_en_refs:
        report.endnotes_preserved = False
        warnings.append("endnote_reference_count_changed")

    if count_hyperlinks(before_doc) != count_hyperlinks(after_doc):
        report.hyperlinks_preserved = False
        warnings.append("hyperlink_count_changed")

    if count_bookmarks(before_doc) != count_bookmarks(after_doc):
        report.bookmarks_preserved = False
        warnings.append("bookmark_count_changed")

    if count_equations(before_doc) != count_equations(after_doc):
        report.equations_preserved = False
        warnings.append("equation_count_changed")

    bf, be = count_footnotes_endnotes(before_doc)
    af, ae = count_footnotes_endnotes(after_doc)
    if bf >= 0 and af >= 0 and bf != af:
        report.footnotes_preserved = False
        warnings.append("footnote_count_changed")
    if be >= 0 and ae >= 0 and be != ae:
        report.endnotes_preserved = False
        warnings.append("endnote_count_changed")

    # Relationship count should not drop.
    try:
        before_rel_count = len(before_doc.part.rels)
        after_rel_count = len(after_doc.part.rels)
        if after_rel_count < before_rel_count:
            report.relationships_preserved = False
            warnings.append("relationship_count_decreased")
    except Exception:
        warnings.append("relationship_count_unreadable")

    report.warnings = warnings
    report.document_preservation_ok = all(
        [
            report.text_integrity_ok,
            report.images_preserved,
            report.hyperlinks_preserved,
            report.bookmarks_preserved,
            report.footnotes_preserved,
            report.endnotes_preserved,
            report.equations_preserved,
            report.relationships_preserved,
            report.embedded_objects_preserved,
            report.numbering_preserved,
        ]
    )
    return report


class DocumentPreservationError(Exception):
    def __init__(self, message: str, warnings: Optional[list[str]] = None):
        super().__init__(message)
        self.warnings = warnings or []
