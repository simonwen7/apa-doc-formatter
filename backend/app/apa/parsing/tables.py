"""Table region modeling for APA Phase 2D (read-only analysis)."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

_TABLE_NUMBER_RE = re.compile(r"^table\s+(\d+)\s*$", re.I)
_TABLE_NUMBER_INLINE_RE = re.compile(r"^table\s+(\d+)\b", re.I)
_NOTE_RE = re.compile(r"^note\.\s*", re.I)
_SPECIFIC_NOTE_RE = re.compile(r"^[a-z]\s+", re.I)
_PROB_NOTE_RE = re.compile(r"^\*+\s*p\s*[<≤=]", re.I)
_CALLOUT_RE = re.compile(r"\b(?:see\s+)?table\s+(\d+)\b", re.I)


@dataclass
class TableAnalysis:
    table_index: int
    confidence: float
    is_apa_candidate: bool
    number_paragraph_index: Optional[int] = None
    title_paragraph_index: Optional[int] = None
    note_paragraph_indices: list[int] = field(default_factory=list)
    callout_locations: list[int] = field(default_factory=list)
    row_count: int = 0
    column_count: int = 0
    merged_cells: bool = False
    detected_header_rows: int = 0
    border_state: str = "unknown"  # minimal|full_grid|mixed|unknown
    spacing_state: str = "unknown"
    alignment_state: str = "unknown"
    warnings: list[str] = field(default_factory=list)
    number_text: Optional[str] = None
    number_value: Optional[int] = None
    evidence: list[str] = field(default_factory=list)


def _paragraph_index_map(doc: Document) -> dict[int, int]:
    mapping: dict[int, int] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        mapping[id(paragraph._p)] = index
    return mapping


def _iter_body_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", child)
        elif child.tag == qn("w:tbl"):
            yield ("tbl", child)


def _has_merged_cells(table: Table) -> bool:
    try:
        for row in table._tbl.tr_lst:
            for cell in row.tc_lst:
                tc_pr = cell.tcPr
                if tc_pr is None:
                    continue
                if tc_pr.gridSpan is not None and int(tc_pr.gridSpan.val) > 1:
                    return True
                v_merge = tc_pr.find(qn("w:vMerge"))
                if v_merge is not None:
                    return True
    except Exception:
        return False
    return False


def _border_state(table: Table) -> str:
    """Best-effort border classification without rewriting the table."""
    try:
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            return "unknown"
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            return "unknown"
        present = []
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                continue
            val = (node.get(qn("w:val")) or "").lower()
            if val and val not in {"nil", "none"}:
                present.append(edge)
        if not present:
            return "minimal"
        if {"top", "bottom", "insideH"}.issubset(set(present)) and "insideV" not in present and "left" not in present:
            return "minimal"
        if "insideV" in present or "left" in present or "right" in present:
            return "full_grid"
        return "mixed"
    except Exception:
        return "unknown"


def _looks_like_layout_table(table: Table) -> bool:
    rows = len(table.rows)
    cols = len(table.columns) if table.rows else 0
    if rows <= 1 and cols <= 1:
        return True
    # Single-column multi-row is often a layout scaffold.
    if cols == 1 and rows >= 2:
        return True
    return False


def _paragraph_text_for_elm(doc: Document, elm, mapping: dict[int, int]) -> tuple[Optional[int], str]:
    index = mapping.get(id(elm))
    if index is None:
        return None, ""
    return index, doc.paragraphs[index].text.strip()


def analyze_tables(doc: Document) -> list[TableAnalysis]:
    mapping = _paragraph_index_map(doc)
    # Collect callouts across body paragraphs.
    callouts_by_number: dict[int, list[int]] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        if _TABLE_NUMBER_RE.match(text) or (
            _TABLE_NUMBER_INLINE_RE.match(text) and len(text.split()) <= 4
        ):
            continue
        for match in _CALLOUT_RE.finditer(paragraph.text or ""):
            num = int(match.group(1))
            callouts_by_number.setdefault(num, []).append(index)

    analyses: list[TableAnalysis] = []
    body_blocks = list(_iter_body_blocks(doc))
    table_counter = 0

    for i, (kind, elm) in enumerate(body_blocks):
        if kind != "tbl":
            continue
        table = Table(elm, doc)
        row_count = len(table.rows)
        col_count = len(table.columns) if table.rows else 0
        merged = _has_merged_cells(table)
        borders = _border_state(table)
        warnings: list[str] = []
        evidence: list[str] = []

        number_idx = None
        title_idx = None
        number_text = None
        number_value = None
        note_indices: list[int] = []

        # Look at up to two preceding body paragraphs for number/title.
        prev_paras: list[tuple[Optional[int], str]] = []
        for j in range(i - 1, max(-1, i - 4), -1):
            if body_blocks[j][0] != "p":
                break
            prev_paras.append(_paragraph_text_for_elm(doc, body_blocks[j][1], mapping))
        prev_paras.reverse()

        if prev_paras:
            last_idx, last_text = prev_paras[-1]
            if last_idx is not None and _TABLE_NUMBER_RE.match(last_text):
                number_idx = last_idx
                number_text = last_text
                number_value = int(_TABLE_NUMBER_RE.match(last_text).group(1))
                evidence.append("number_immediately_above")
            elif last_idx is not None and _TABLE_NUMBER_INLINE_RE.match(last_text) and len(last_text.split()) <= 4:
                number_idx = last_idx
                number_text = last_text
                number_value = int(_TABLE_NUMBER_INLINE_RE.match(last_text).group(1))
                evidence.append("number_inline_above")
            elif (
                len(prev_paras) >= 2
                and prev_paras[-2][0] is not None
                and _TABLE_NUMBER_RE.match(prev_paras[-2][1])
            ):
                number_idx = prev_paras[-2][0]
                number_text = prev_paras[-2][1]
                number_value = int(_TABLE_NUMBER_RE.match(number_text).group(1))
                title_idx = last_idx
                evidence.extend(["number_two_above", "title_immediately_above"])
            elif last_idx is not None and last_text and not _NOTE_RE.match(last_text):
                # Uncertain title-only association — keep low confidence later.
                title_idx = last_idx
                evidence.append("possible_title_above_no_number")
                warnings.append("title_without_confident_number")

            if (
                title_idx is None
                and number_idx is not None
                and last_idx is not None
                and last_idx != number_idx
                and last_text
                and not _NOTE_RE.match(last_text)
            ):
                title_idx = last_idx
                evidence.append("title_after_number")

        # Notes immediately after table.
        for j in range(i + 1, min(len(body_blocks), i + 5)):
            if body_blocks[j][0] != "p":
                break
            n_idx, n_text = _paragraph_text_for_elm(doc, body_blocks[j][1], mapping)
            if n_idx is None or not n_text:
                continue
            if _NOTE_RE.match(n_text) or _PROB_NOTE_RE.match(n_text) or (
                note_indices and _SPECIFIC_NOTE_RE.match(n_text)
            ):
                note_indices.append(n_idx)
                evidence.append("note_below")
            else:
                break

        layout = _looks_like_layout_table(table)
        if layout:
            evidence.append("layout_like_structure")
            warnings.append("possible_layout_table")

        is_apa = False
        confidence = 0.4
        if number_idx is not None and not layout:
            is_apa = True
            confidence = 0.95
            evidence.append("apa_number_signal")
        elif number_idx is not None and layout:
            is_apa = True
            confidence = 0.7
            warnings.append("number_on_layout_like_table")
        elif not layout and row_count >= 2 and col_count >= 2:
            # Multi-cell table without caption — uncertain scholarly table.
            is_apa = False
            confidence = 0.55
            warnings.append("uncaptioned_grid_table")
            evidence.append("grid_without_number")
        else:
            is_apa = False
            confidence = 0.35

        callouts = callouts_by_number.get(number_value or -1, []) if number_value else []

        analyses.append(
            TableAnalysis(
                table_index=table_counter,
                confidence=confidence,
                is_apa_candidate=is_apa,
                number_paragraph_index=number_idx,
                title_paragraph_index=title_idx,
                note_paragraph_indices=note_indices,
                callout_locations=callouts,
                row_count=row_count,
                column_count=col_count,
                merged_cells=merged,
                detected_header_rows=1 if row_count >= 2 and col_count >= 2 else 0,
                border_state=borders,
                spacing_state="unknown",
                alignment_state="unknown",
                warnings=warnings,
                number_text=number_text,
                number_value=number_value,
                evidence=evidence,
            )
        )
        table_counter += 1

    return analyses
