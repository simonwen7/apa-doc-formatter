from docx import Document
from app.core.config import TEMPLATES_DIR
from app.utils.json_helpers import load_json_file


def load_template(template_id: str) -> dict:
    template_path = TEMPLATES_DIR / f"{template_id.replace('-', '_')}.json"
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_id}")
    return load_json_file(template_path)


def _paragraph_has_page_break_before(paragraph) -> bool:
    try:
        return bool(paragraph.paragraph_format.page_break_before)
    except Exception:
        return False


def _paragraph_contains_manual_page_break(paragraph) -> bool:
    try:
        for run in paragraph.runs:
            xml = run._element.xml
            if 'w:br w:type="page"' in xml or ("w:br" in xml and 'w:type="page"' in xml):
                return True
    except Exception:
        pass
    return False


def _section_vertical_alignment_name(section):
    try:
        va = section.vertical_alignment
        return str(va) if va is not None else None
    except Exception:
        return None


def parse_docx(file_path: str) -> dict:
    doc = Document(file_path)

    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        runs = []
        for r in p.runs:
            runs.append(
                {
                    "text": r.text,
                    "bold": r.bold,
                    "italic": r.italic,
                    "underline": r.underline,
                    "font_name": r.font.name,
                    "font_size_pt": r.font.size.pt if r.font.size else None,
                }
            )

        para_format = p.paragraph_format
        paragraphs.append(
            {
                "index": i,
                "text": p.text.strip(),
                "style": p.style.name if p.style else None,
                "alignment": p.alignment,
                "line_spacing": para_format.line_spacing,
                "space_before_pt": para_format.space_before.pt if para_format.space_before else None,
                "space_after_pt": para_format.space_after.pt if para_format.space_after else None,
                "first_line_indent_in": para_format.first_line_indent.inches if para_format.first_line_indent else None,
                "left_indent_in": para_format.left_indent.inches if para_format.left_indent else None,
                "right_indent_in": para_format.right_indent.inches if para_format.right_indent else None,
                "page_break_before": _paragraph_has_page_break_before(p),
                "contains_manual_page_break": _paragraph_contains_manual_page_break(p),
                "runs": runs,
            }
        )

    sections = []
    for i, section in enumerate(doc.sections):
        sections.append(
            {
                "index": i,
                "top_margin": section.top_margin.inches if section.top_margin else None,
                "bottom_margin": section.bottom_margin.inches if section.bottom_margin else None,
                "left_margin": section.left_margin.inches if section.left_margin else None,
                "right_margin": section.right_margin.inches if section.right_margin else None,
                "vertical_alignment": _section_vertical_alignment_name(section),
            }
        )

    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                row_data.append(
                    {
                        "row": r_idx,
                        "col": c_idx,
                        "text": "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()),
                    }
                )
            rows.append(row_data)

        tables.append(
            {
                "index": t_idx,
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "rows": rows,
            }
        )

    return {
        "paragraphs": paragraphs,
        "sections": sections,
        "tables": tables,
        "paragraph_count": len(paragraphs),
        "section_count": len(sections),
        "table_count": len(tables),
    }
