import re
from typing import Dict, Optional, List, Set

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from docx.enum.section import WD_ALIGN_VERTICAL
except Exception:
    WD_ALIGN_VERTICAL = None


APA_FONT_NAME = "Times New Roman"
APA_FONT_SIZE_PT = 12
APA_LINE_SPACING = 2
APA_FIRST_LINE_INDENT_IN = 0.5
MAJOR_HEADING_SPACE_PT = 24


def _safe_style_name(paragraph) -> str:
    try:
        return paragraph.style.name if paragraph.style else ""
    except Exception:
        return ""


def _has_text(paragraph) -> bool:
    return bool(paragraph.text and paragraph.text.strip())


def _non_empty_runs(paragraph):
    return [run for run in paragraph.runs if run.text and run.text.strip()]


def _find_heading_index(doc: Document, heading_text: str) -> Optional[int]:
    target = heading_text.strip().lower()
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == target:
            return idx
    return None


def _find_references_start(doc: Document) -> Optional[int]:
    return _find_heading_index(doc, "references")


def _reference_indices(doc: Document, ref_start_idx: Optional[int]) -> List[int]:
    if ref_start_idx is None:
        return []

    indices = []
    for idx in range(ref_start_idx + 1, len(doc.paragraphs)):
        if _has_text(doc.paragraphs[idx]):
            indices.append(idx)
    return indices


def _is_major_section_heading_text(text: str) -> bool:
    t = text.strip().lower()
    if not t:
        return False

    if t in {"introduction", "conclusion", "references", "appendix", "appendices"}:
        return True

    if t.startswith("section "):
        return True

    if t.startswith("appendix "):
        return True

    return False


def _is_heading_like(paragraph) -> bool:
    text = paragraph.text.strip()
    style_name = _safe_style_name(paragraph)

    if not text:
        return False

    if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
        return True

    if text.startswith("Question "):
        return True

    if text.startswith("Prompt "):
        return True

    if text.startswith("[Prompt:"):
        return True

    if _is_major_section_heading_text(text):
        return True

    return False


def _is_instruction_like(paragraph) -> bool:
    text = paragraph.text.strip()
    style_name = _safe_style_name(paragraph).lower()

    if not text:
        return False

    upper = text.upper()
    if upper.startswith("WORD COUNT"):
        return True
    if upper.startswith("INCLUDE YOUR ANSWER"):
        return True
    if upper.startswith("WHICH PROMPT ARE YOU ANSWERING"):
        return True

    if style_name in {"paragraph", "list paragraph"}:
        return True

    return False


def _is_body_like(paragraph) -> bool:
    text = paragraph.text.strip()

    if not text:
        return False

    if _is_heading_like(paragraph):
        return False

    if _is_instruction_like(paragraph):
        return False

    if text.lower() == "references":
        return False

    return True


def _apply_font_to_runs(paragraph, fixed_counts):
    for run in paragraph.runs:
        if run.text is None:
            continue

        if run.font.name != APA_FONT_NAME:
            run.font.name = APA_FONT_NAME
            fixed_counts["font_name"] += 1

        size = run.font.size.pt if run.font.size else None
        if size != APA_FONT_SIZE_PT:
            run.font.size = Pt(APA_FONT_SIZE_PT)
            fixed_counts["font_size"] += 1


def _clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def _set_paragraph_text_single_run(paragraph, text: str):
    _clear_paragraph(paragraph)
    paragraph.add_run(text)


def _insert_paragraph_before(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _normalize_general_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\u00A0", " ")
    text = text.replace("\t", " ")
    text = text.replace("\r", " ")
    text = text.replace("\n", " ")

    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"([(\[])\s+", r"\1", text)
    text = re.sub(r"\s+([)\]])", r"\1", text)
    text = re.sub(r"\s*/\s*", "/", text)
    text = re.sub(r"\s*&\s*", " & ", text)
    text = re.sub(r"\s{2,}", " ", text)

    return text.strip()


def _normalize_reference_text(text: str) -> str:
    text = _normalize_general_text(text)
    text = re.sub(r"\.\s*\(\s*(\d{4}[a-z]?)\s*\)\s*\.", r". (\1).", text)
    text = re.sub(r"\s*,\s*", ", ", text)
    text = re.sub(r"\bdoi\s*:\s*", "https://doi.org/", text, flags=re.IGNORECASE)
    text = re.sub(r"(\d)\s+\(\s*([^)]+?)\s*\)", r"\1(\2)", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _strip_leading_bullet_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"^\s*[•▪■◦·]+\s*", "", text)
    text = re.sub(r"^\s*o\s+(?=[A-Z])", "", text)
    return text.strip()


def _sanitize_heading_text_in_place(paragraph, fixed_counts):
    original = paragraph.text
    normalized = _normalize_general_text(original)
    normalized = _strip_leading_bullet_text(normalized)

    if normalized != original:
        _set_paragraph_text_single_run(paragraph, normalized)
        fixed_counts["text_normalized"] += 1


def _sentence_case_title(title: str) -> str:
    if not title:
        return title

    title = _normalize_general_text(title)
    words = title.split(" ")
    if not words:
        return title

    def keep_token(token: str) -> bool:
        stripped = re.sub(r"^[\"'(\[]+|[\"')\].,:;!?]+$", "", token)
        if not stripped:
            return True
        if stripped.isupper() and len(stripped) > 1:
            return True
        if re.search(r"[A-Z].*[A-Z]", stripped):
            return True
        return False

    new_words = []
    for i, word in enumerate(words):
        if i == 0:
            if keep_token(word):
                new_words.append(word)
            else:
                new_words.append(word[:1].upper() + word[1:].lower())
        else:
            if keep_token(word):
                new_words.append(word)
            else:
                new_words.append(word.lower())

    result = " ".join(new_words)
    result = re.sub(r"(:\s+)([a-z])", lambda m: m.group(1) + m.group(2).upper(), result)
    return result


def _parse_journal_reference(text: str) -> Optional[Dict[str, str]]:
    text = _normalize_reference_text(text)

    pattern = re.compile(
        r"""
        ^
        (?P<authors>.+?)
        \s*\(\s*(?P<year>\d{4}[a-z]?)\s*\)\.\s*
        (?P<title>.+?)\.\s*
        (?P<journal>[^,]+?)\s*,\s*
        (?P<volume>\d+)
        (?:\(\s*(?P<issue>[^)]+?)\s*\))?
        (?:\s*,\s*(?P<pages>[^.]+?))?
        \.\s*
        (?P<tail>.*)
        $
        """,
        re.VERBOSE
    )

    m = pattern.match(text)
    if not m:
        return None

    groups = {k: (v.strip() if v else "") for k, v in m.groupdict().items()}
    groups["title"] = _sentence_case_title(groups["title"])
    return groups


def _rebuild_journal_reference(paragraph, parsed: Dict[str, str], fixed_counts):
    _clear_paragraph(paragraph)

    def add(text, italic=False, bold=False):
        if not text:
            return
        run = paragraph.add_run(text)
        run.italic = italic
        run.bold = bold
        run.font.name = APA_FONT_NAME
        run.font.size = Pt(APA_FONT_SIZE_PT)

    add(f"{parsed['authors']} ({parsed['year']}). ")
    add(f"{parsed['title']}. ")
    add(parsed["journal"], italic=True)
    add(", ")
    add(parsed["volume"], italic=True)

    if parsed.get("issue"):
        add(f"({parsed['issue']})", italic=False)

    if parsed.get("pages"):
        add(f", {parsed['pages']}")

    add(".")

    if parsed.get("tail"):
        tail = parsed["tail"].strip()
        if tail:
            add(f" {tail}")

    fixed_counts["reference_entries_rebuilt"] += 1


def _normalize_paragraph_text_in_place(paragraph, fixed_counts, reference_mode=False):
    original = paragraph.text
    normalized = _normalize_reference_text(original) if reference_mode else _normalize_general_text(original)

    if normalized != original:
        _set_paragraph_text_single_run(paragraph, normalized)
        fixed_counts["text_normalized"] += 1


def _apply_major_heading_spacing(paragraph, fixed_counts):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(MAJOR_HEADING_SPACE_PT)
    fmt.space_after = Pt(MAJOR_HEADING_SPACE_PT)
    fixed_counts["major_heading_spacing"] += 1


def _fix_body_paragraph(paragraph, fixed_counts):
    _normalize_paragraph_text_in_place(paragraph, fixed_counts)

    fmt = paragraph.paragraph_format
    fmt.line_spacing = APA_LINE_SPACING
    fmt.first_line_indent = Inches(APA_FIRST_LINE_INDENT_IN)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.page_break_before = False

    fixed_counts["line_spacing"] += 1
    fixed_counts["paragraph_indent"] += 1
    fixed_counts["space_before"] += 1
    fixed_counts["space_after"] += 1

    _apply_font_to_runs(paragraph, fixed_counts)
    fixed_counts["body_paragraphs_touched"] += 1


def _remove_bullets_numbering(paragraph):
    try:
        pPr = paragraph._p.get_or_add_pPr()
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)
    except Exception:
        pass


def _fix_heading_paragraph(paragraph, fixed_counts):
    _sanitize_heading_text_in_place(paragraph, fixed_counts)
    _remove_bullets_numbering(paragraph)

    text = paragraph.text.strip()
    text_lower = text.lower()
    style_name = _safe_style_name(paragraph)
    fmt = paragraph.paragraph_format

    fmt.first_line_indent = Inches(0)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.line_spacing = APA_LINE_SPACING

    if _is_major_section_heading_text(text):
        _apply_major_heading_spacing(paragraph, fixed_counts)
    else:
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)

    if text_lower in {"introduction", "conclusion", "references"} or text_lower.startswith("section ") or text_lower.startswith("appendix"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in _non_empty_runs(paragraph):
            run.bold = True
            run.italic = False
    elif style_name == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in _non_empty_runs(paragraph):
            run.bold = True
            run.italic = False
    elif style_name == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in _non_empty_runs(paragraph):
            run.bold = True
            run.italic = False
    elif style_name == "Heading 3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in _non_empty_runs(paragraph):
            run.bold = True
            run.italic = True
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if _is_major_section_heading_text(text) else WD_ALIGN_PARAGRAPH.LEFT

    if text_lower == "references":
        fmt.page_break_before = True
        fixed_counts["references_page_break"] += 1
    elif text_lower == "introduction":
        fmt.page_break_before = True
        fixed_counts["introduction_page_break"] += 1
    else:
        fmt.page_break_before = False

    _apply_font_to_runs(paragraph, fixed_counts)


def _fix_references_heading(paragraph, fixed_counts):
    normalized = _strip_leading_bullet_text(_normalize_general_text(paragraph.text))
    if normalized.lower() != "references":
        normalized = "References"

    _set_paragraph_text_single_run(paragraph, normalized)
    _remove_bullets_numbering(paragraph)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(0)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.line_spacing = APA_LINE_SPACING
    fmt.page_break_before = True
    _apply_major_heading_spacing(paragraph, fixed_counts)

    for run in _non_empty_runs(paragraph):
        run.bold = True
        run.italic = False

    _apply_font_to_runs(paragraph, fixed_counts)
    fixed_counts["references_page_break"] += 1


def _fix_reference_entry(paragraph, fixed_counts):
    raw_text = paragraph.text
    normalized = _normalize_reference_text(raw_text)

    parsed = _parse_journal_reference(normalized)
    if parsed:
        _rebuild_journal_reference(paragraph, parsed, fixed_counts)
    else:
        if normalized != raw_text:
            _set_paragraph_text_single_run(paragraph, normalized)
            fixed_counts["text_normalized"] += 1

    fmt = paragraph.paragraph_format
    fmt.line_spacing = APA_LINE_SPACING
    fmt.left_indent = Inches(0.5)
    fmt.first_line_indent = Inches(-0.5)
    fmt.right_indent = None
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.page_break_before = False

    fixed_counts["reference_line_spacing"] += 1
    fixed_counts["reference_hanging_indent"] += 1
    fixed_counts["reference_space_before"] += 1
    fixed_counts["reference_space_after"] += 1

    _apply_font_to_runs(paragraph, fixed_counts)
    fixed_counts["reference_entries_touched"] += 1


def _delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _is_bullet_only_paragraph(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False

    if re.fullmatch(r"[•▪■◦·]+", text):
        return True

    if len(text) <= 3 and not re.search(r"[A-Za-z0-9]", text):
        return True

    style_name = _safe_style_name(paragraph).lower()
    if style_name == "list paragraph" and len(text) <= 3:
        return True

    return False


def _remove_intro_leading_junk(doc: Document, fixed_counts):
    """
    专门删除 Introduction 前面紧挨着的垃圾段落：
    - 空段
    - 纯 bullet/symbol 段
    """
    intro_idx = _find_heading_index(doc, "introduction")
    if intro_idx is None:
        return

    changed = True
    while changed:
        changed = False
        intro_idx = _find_heading_index(doc, "introduction")
        if intro_idx is None or intro_idx == 0:
            return

        prev_para = doc.paragraphs[intro_idx - 1]
        prev_text = prev_para.text.strip()

        if not prev_text or _is_bullet_only_paragraph(prev_para):
            _delete_paragraph(prev_para)
            fixed_counts["intro_leading_junk_removed"] += 1
            changed = True


def _remove_extra_empty_paragraphs(doc: Document, fixed_counts):
    changed = True
    while changed:
        changed = False
        paragraphs = list(doc.paragraphs)

        for i, para in enumerate(paragraphs):
            text = para.text.strip()

            if _is_bullet_only_paragraph(para):
                _delete_paragraph(para)
                fixed_counts["bullet_only_paragraphs_removed"] += 1
                changed = True
                break

            if not text:
                prev_text = paragraphs[i - 1].text.strip() if i > 0 else ""
                next_text = paragraphs[i + 1].text.strip() if i + 1 < len(paragraphs) else ""

                if not prev_text or not next_text:
                    _delete_paragraph(para)
                    fixed_counts["empty_paragraphs_removed"] += 1
                    changed = True
                    break

                if _is_major_section_heading_text(prev_text) or _is_major_section_heading_text(next_text):
                    _delete_paragraph(para)
                    fixed_counts["empty_paragraphs_removed"] += 1
                    changed = True
                    break

                if prev_text.lower() == "references" or next_text.lower() == "references":
                    _delete_paragraph(para)
                    fixed_counts["empty_paragraphs_removed"] += 1
                    changed = True
                    break


def _get_title_page_indices(doc: Document) -> Set[int]:
    intro_idx = _find_heading_index(doc, "introduction")
    if intro_idx is None:
        intro_idx = len(doc.paragraphs)

    result = set()
    for idx in range(intro_idx):
        if doc.paragraphs[idx].text.strip():
            result.add(idx)
    return result


def _configure_title_page_vertical_center(doc: Document, fixed_counts):
    if not doc.sections:
        return

    # 不碰 header / footer，只设正文 section 垂直居中
    try:
        if WD_ALIGN_VERTICAL is not None:
            doc.sections[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            fixed_counts["title_page_vertical_center"] += 1
    except Exception:
        pass


def _fix_title_page_paragraph(paragraph, fixed_counts):
    _normalize_paragraph_text_in_place(paragraph, fixed_counts)
    _remove_bullets_numbering(paragraph)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(0)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.line_spacing = APA_LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.page_break_before = False

    _apply_font_to_runs(paragraph, fixed_counts)
    fixed_counts["title_page_paragraphs_centered"] += 1


def _force_cover_page_group_to_middle(doc: Document, fixed_counts, title_page_indices: Set[int]):
    """
    不靠 Word 的 vertical alignment 玄学，直接在 cover page 第一条正文前插入若干空段，
    把整组内容推到页面中部附近。
    不动页眉页脚。
    """
    ordered = sorted(title_page_indices)
    if not ordered:
        return

    first_idx = ordered[0]
    first_para = doc.paragraphs[first_idx]

    # 先清掉第一段自己的段前空白
    fmt = first_para.paragraph_format
    fmt.space_before = Pt(0)

    # 粗暴但稳定：插入若干空段，把正文组推到中间
    # 这个值是按你截图效果调的，通常 6~8 段比较接近整页中间
    blank_count = 7

    for _ in range(blank_count):
        new_para = _insert_paragraph_before(first_para, "")
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_fmt = new_para.paragraph_format
        new_fmt.first_line_indent = Inches(0)
        new_fmt.left_indent = None
        new_fmt.right_indent = None
        new_fmt.line_spacing = APA_LINE_SPACING
        new_fmt.space_before = Pt(0)
        new_fmt.space_after = Pt(0)

    fixed_counts["cover_page_leading_blanks_inserted"] += blank_count


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = qn(f"w:{edge}")
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)

            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def _format_table_borders_apa(table):
    rows = table.rows
    if not rows:
        return

    last_row_idx = len(rows) - 1

    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            _set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )

            if r_idx == 0:
                _set_cell_border(cell, top={"val": "single", "sz": "8", "color": "000000"})
                _set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"})

            if r_idx == last_row_idx:
                _set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"})


def _set_repeat_table_header(row):
    try:
        trPr = row._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "true")
        trPr.append(tblHeader)
    except Exception:
        pass


def _fix_table_paragraph(paragraph, fixed_counts, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    if not paragraph.text.strip():
        return

    _normalize_paragraph_text_in_place(paragraph, fixed_counts)

    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Inches(0)
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.line_spacing = APA_LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)

    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic

    _apply_font_to_runs(paragraph, fixed_counts)


def _fix_tables(doc: Document, fixed_counts):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        fixed_counts["tables_touched"] += 1

        rows = table.rows
        if rows:
            _set_repeat_table_header(rows[0])

        for r_idx, row in enumerate(rows):
            for cell in row.cells:
                try:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                except Exception:
                    pass

                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    fmt = para.paragraph_format
                    fmt.first_line_indent = Inches(0)
                    fmt.left_indent = None
                    fmt.right_indent = None
                    fmt.line_spacing = 1.0
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)

                    for run in para.runs:
                        if r_idx == 0:
                            run.bold = True
                        run.font.name = APA_FONT_NAME
                        run.font.size = Pt(APA_FONT_SIZE_PT)

        _format_table_borders_apa(table)

    paragraphs = list(doc.paragraphs)
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        if re.match(r"^Table\s+\d+\.?$", text, flags=re.IGNORECASE):
            _fix_table_paragraph(
                para,
                fixed_counts,
                bold=True,
                italic=False,
                align=WD_ALIGN_PARAGRAPH.LEFT
            )
            fixed_counts["table_caption_paragraphs_touched"] += 1

            if i + 1 < len(paragraphs):
                next_para = paragraphs[i + 1]
                if next_para.text.strip():
                    _fix_table_paragraph(
                        next_para,
                        fixed_counts,
                        bold=False,
                        italic=True,
                        align=WD_ALIGN_PARAGRAPH.LEFT
                    )
                    fixed_counts["table_title_paragraphs_touched"] += 1


def fix_apa_format(input_path: str, output_path: str) -> dict:
    doc = Document(input_path)

    fixed_counts = {
        "margins": 0,
        "title_page_vertical_center": 0,
        "title_page_paragraphs_centered": 0,
        "cover_page_leading_blanks_inserted": 0,
        "introduction_page_break": 0,
        "references_page_break": 0,
        "major_heading_spacing": 0,
        "empty_paragraphs_removed": 0,
        "bullet_only_paragraphs_removed": 0,
        "intro_leading_junk_removed": 0,
        "tables_touched": 0,
        "table_caption_paragraphs_touched": 0,
        "table_title_paragraphs_touched": 0,
        "body_paragraphs_touched": 0,
        "paragraph_indent": 0,
        "line_spacing": 0,
        "space_before": 0,
        "space_after": 0,
        "font_name": 0,
        "font_size": 0,
        "reference_entries_touched": 0,
        "reference_entries_rebuilt": 0,
        "reference_hanging_indent": 0,
        "reference_line_spacing": 0,
        "reference_space_before": 0,
        "reference_space_after": 0,
        "headings_touched": 0,
        "text_normalized": 0,
    }

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        fixed_counts["margins"] += 1

    _remove_extra_empty_paragraphs(doc, fixed_counts)
    _remove_intro_leading_junk(doc, fixed_counts)
    _configure_title_page_vertical_center(doc, fixed_counts)

    title_page_indices = _get_title_page_indices(doc)

    ref_start_idx = _find_references_start(doc)
    ref_indices = set(_reference_indices(doc, ref_start_idx))

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        if idx in title_page_indices:
            _fix_title_page_paragraph(para, fixed_counts)
            continue

        if ref_start_idx is not None and idx == ref_start_idx:
            _fix_references_heading(para, fixed_counts)
            fixed_counts["headings_touched"] += 1
            continue

        if idx in ref_indices:
            _fix_reference_entry(para, fixed_counts)
            continue

        if _is_heading_like(para):
            _fix_heading_paragraph(para, fixed_counts)
            fixed_counts["headings_touched"] += 1
            continue

        if _is_instruction_like(para):
            _normalize_paragraph_text_in_place(para, fixed_counts)
            _apply_font_to_runs(para, fixed_counts)
            continue

        if _is_body_like(para):
            _fix_body_paragraph(para, fixed_counts)

    _force_cover_page_group_to_middle(doc, fixed_counts, title_page_indices)
    _fix_tables(doc, fixed_counts)
    _remove_intro_leading_junk(doc, fixed_counts)

    doc.save(output_path)
    return fixed_counts


def fix_apa(input_path: str, output_path: str) -> dict:
    return fix_apa_format(input_path, output_path)
