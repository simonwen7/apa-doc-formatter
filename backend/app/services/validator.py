import re
from collections import Counter
from typing import Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    from docx.enum.section import WD_ALIGN_VERTICAL
except Exception:
    WD_ALIGN_VERTICAL = None


APA_FONT_NAME = "Times New Roman"
APA_FONT_SIZE_PT = 12
APA_LINE_SPACING = 2
APA_FIRST_LINE_INDENT_IN = 0.5
MAJOR_HEADING_SPACE_PT = 24


def _safe_style_name(paragraph) -> str:
    try:
        return paragraph.style.name if paragraph.style else ""
    except Exception:
        return ""


def _has_text(paragraph) -> bool:
    return bool(paragraph.text and paragraph.text.strip())


def _non_empty_runs(paragraph):
    return [run for run in paragraph.runs if run.text and run.text.strip()]


def _all_non_empty_runs_bold(paragraph) -> bool:
    runs = _non_empty_runs(paragraph)
    if not runs:
        return False
    return all(bool(run.bold) for run in runs)


def _all_non_empty_runs_bold_italic(paragraph) -> bool:
    runs = _non_empty_runs(paragraph)
    if not runs:
        return False
    return all(bool(run.bold) and bool(run.italic) for run in runs)


def _line_spacing_ok(paragraph) -> bool:
    return paragraph.paragraph_format.line_spacing == APA_LINE_SPACING


def _first_line_indent_ok(paragraph) -> bool:
    indent = paragraph.paragraph_format.first_line_indent
    if indent is None:
        return False
    return abs(indent.inches - APA_FIRST_LINE_INDENT_IN) < 0.01


def _no_first_line_indent(paragraph) -> bool:
    indent = paragraph.paragraph_format.first_line_indent
    if indent is None:
        return True
    return abs(indent.inches) < 0.01


def _is_left_aligned_or_default(paragraph) -> bool:
    return paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)


def _is_centered(paragraph) -> bool:
    return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def _is_references_heading(paragraph) -> bool:
    return paragraph.text.strip().lower() == "references"


def _is_introduction_heading(paragraph) -> bool:
    return paragraph.text.strip().lower() == "introduction"


def _is_major_section_heading_text(text: str) -> bool:
    t = text.strip().lower()
    if not t:
        return False

    if t in {"introduction", "conclusion", "references", "appendix", "appendices"}:
        return True

    if t.startswith("section "):
        return True

    if t.startswith("appendix "):
        return True

    return False


def _find_references_start(doc):
    for idx, para in enumerate(doc.paragraphs):
        if _is_references_heading(para):
            return idx
    return None


def _reference_indices(doc, ref_start_idx):
    if ref_start_idx is None:
        return []

    indices = []
    for idx in range(ref_start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        if _has_text(para):
            indices.append(idx)
    return indices


def _is_hanging_indent(paragraph) -> bool:
    fmt = paragraph.paragraph_format
    left_indent = fmt.left_indent
    first_line_indent = fmt.first_line_indent

    if left_indent is None or first_line_indent is None:
        return False

    return abs(left_indent.inches - 0.5) < 0.01 and abs(first_line_indent.inches + 0.5) < 0.01


def _add_issue(issues, category: str, message: str):
    issues.append({"category": category, "message": message})


def _normalize_general_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\u00A0", " ")
    text = text.replace("\t", " ")
    text = text.replace("\r", " ")
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"([(\[])\s+", r"\1", text)
    text = re.sub(r"\s+([)\]])", r"\1", text)
    text = re.sub(r"\s*&\s*", " & ", text)
    text = re.sub(r"\s{2,}", " ", text)

    return text.strip()


def _normalize_reference_text(text: str) -> str:
    text = _normalize_general_text(text)
    text = re.sub(r"\.\s*\(\s*(\d{4}[a-z]?)\s*\)\s*\.", r". (\1).", text)
    text = re.sub(r"\s*,\s*", ", ", text)
    text = re.sub(r"\bdoi\s*:\s*", "https://doi.org/", text, flags=re.IGNORECASE)
    text = re.sub(r"(\d)\s+\(\s*([^)]+?)\s*\)", r"\1(\2)", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _parse_journal_reference(text: str) -> Optional[Dict[str, str]]:
    text = _normalize_reference_text(text)

    pattern = re.compile(
        r"""
        ^
        (?P<authors>.+?)
        \s*\(\s*(?P<year>\d{4}[a-z]?)\s*\)\.\s*
        (?P<title>.+?)\.\s*
        (?P<journal>[^,]+?)\s*,\s*
        (?P<volume>\d+)
        (?:\(\s*(?P<issue>[^)]+?)\s*\))?
        (?:\s*,\s*(?P<pages>[^.]+?))?
        \.\s*
        (?P<tail>.*)
        $
        """,
        re.VERBOSE
    )

    match = pattern.match(text)
    if not match:
        return None

    return {k: (v.strip() if v else "") for k, v in match.groupdict().items()}


def _paragraph_uses_non_apa_font(paragraph):
    for run in _non_empty_runs(paragraph):
        if run.font.name != APA_FONT_NAME:
            return run.font.name
    return None


def _paragraph_uses_wrong_font_size(paragraph):
    for run in _non_empty_runs(paragraph):
        size = run.font.size.pt if run.font.size else None
        if size != APA_FONT_SIZE_PT:
            return size
    return None


def _paragraph_has_page_break_before(paragraph) -> bool:
    try:
        return bool(paragraph.paragraph_format.page_break_before)
    except Exception:
        return False


def _previous_paragraph_has_manual_page_break(doc, idx: int) -> bool:
    if idx <= 0:
        return False

    prev_para = doc.paragraphs[idx - 1]
    try:
        for run in prev_para.runs:
            xml = run._element.xml
            if 'w:br w:type="page"' in xml or ("w:br" in xml and 'w:type="page"' in xml):
                return True
    except Exception:
        pass

    return False


def _heading_starts_on_new_page(doc, idx: int) -> bool:
    para = doc.paragraphs[idx]
    if _paragraph_has_page_break_before(para):
        return True
    if _previous_paragraph_has_manual_page_break(doc, idx):
        return True
    return False


def _title_page_is_vertically_centered(doc: Document) -> bool:
    if not doc.sections:
        return False
    try:
        first_section = doc.sections[0]
        value = first_section.vertical_alignment
        return value is not None and str(value).upper().endswith("CENTER")
    except Exception:
        return False


def _space_matches_target(paragraph, target_pt: float) -> bool:
    fmt = paragraph.paragraph_format
    before = fmt.space_before.pt if fmt.space_before else 0
    after = fmt.space_after.pt if fmt.space_after else 0
    return abs(before - target_pt) < 0.5 and abs(after - target_pt) < 0.5


def _validate_body_paragraph(paragraph, idx, issues):
    if not _line_spacing_ok(paragraph):
        _add_issue(issues, "Body", f"Paragraph {idx} line spacing is not explicitly set to double spacing")

    if not _first_line_indent_ok(paragraph):
        _add_issue(issues, "Body", f"Paragraph {idx} first-line indent is not 0.5 inch")

    bad_font = _paragraph_uses_non_apa_font(paragraph)
    if bad_font:
        _add_issue(issues, "Body", f"Paragraph {idx} uses non-APA font: {bad_font}")

    bad_size = _paragraph_uses_wrong_font_size(paragraph)
    if bad_size is not None:
        _add_issue(issues, "Body", f"Paragraph {idx} font size is not {APA_FONT_SIZE_PT} pt")


def _validate_heading_paragraph(paragraph, idx, issues):
    text = paragraph.text.strip().lower()

    if _is_major_section_heading_text(paragraph.text):
        if not _is_centered(paragraph):
            _add_issue(issues, "Heading", f"Paragraph {idx} major heading should be centered")
        if not _all_non_empty_runs_bold(paragraph):
            _add_issue(issues, "Heading", f"Paragraph {idx} major heading should be fully bold")
        if not _space_matches_target(paragraph, MAJOR_HEADING_SPACE_PT):
            _add_issue(issues, "Heading", f"Paragraph {idx} major heading spacing should be {MAJOR_HEADING_SPACE_PT} pt before and after")

    if not _line_spacing_ok(paragraph):
        _add_issue(issues, "Heading", f"Paragraph {idx} heading line spacing is not double")

    bad_font = _paragraph_uses_non_apa_font(paragraph)
    if bad_font:
        _add_issue(issues, "Heading", f"Paragraph {idx} uses non-APA font: {bad_font}")

    bad_size = _paragraph_uses_wrong_font_size(paragraph)
    if bad_size is not None:
        _add_issue(issues, "Heading", f"Paragraph {idx} font size is not 12 pt")


def _validate_reference_spacing_text(paragraph, idx, issues):
    raw_text = paragraph.text
    normalized = _normalize_reference_text(raw_text)

    if raw_text.strip() != normalized:
        _add_issue(issues, "References", f"Paragraph {idx} contains irregular spaces or punctuation spacing")


def _validate_reference_italics(paragraph, idx, issues):
    parsed = _parse_journal_reference(paragraph.text)
    if not parsed:
        return

    italic_text = "".join(run.text for run in paragraph.runs if run.italic)
    journal = parsed.get("journal", "")
    volume = parsed.get("volume", "")
    issue = parsed.get("issue", "")
    title = parsed.get("title", "")

    if journal and journal not in italic_text:
        _add_issue(issues, "References", f"Paragraph {idx} journal title should be italicized")

    if volume and volume not in italic_text:
        _add_issue(issues, "References", f"Paragraph {idx} volume number should be italicized")

    if issue and f"({issue})" in italic_text:
        _add_issue(issues, "References", f"Paragraph {idx} issue number should not be italicized")

    if title and title in italic_text:
        _add_issue(issues, "References", f"Paragraph {idx} article title should not be italicized")


def _validate_references(doc, ref_start_idx, issues):
    if ref_start_idx is None:
        return

    heading_para = doc.paragraphs[ref_start_idx]

    if not _heading_starts_on_new_page(doc, ref_start_idx):
        _add_issue(issues, "References", f"Paragraph {ref_start_idx} References heading should start on a new page")

    if not _is_centered(heading_para):
        _add_issue(issues, "References", f"Paragraph {ref_start_idx} References heading is not centered")

    if not _all_non_empty_runs_bold(heading_para):
        _add_issue(issues, "References", f"Paragraph {ref_start_idx} References heading is not fully bold")

    ref_indices = _reference_indices(doc, ref_start_idx)
    for idx in ref_indices:
        para = doc.paragraphs[idx]

        if not _line_spacing_ok(para):
            _add_issue(issues, "References", f"Paragraph {idx} reference entry is not double spaced")

        if not _is_hanging_indent(para):
            _add_issue(issues, "References", f"Paragraph {idx} reference entry does not use hanging indent")

        bad_font = _paragraph_uses_non_apa_font(para)
        if bad_font:
            _add_issue(issues, "References", f"Paragraph {idx} reference entry uses non-APA font: {bad_font}")

        bad_size = _paragraph_uses_wrong_font_size(para)
        if bad_size is not None:
            _add_issue(issues, "References", f"Paragraph {idx} reference entry font size is not 12 pt")

        _validate_reference_spacing_text(para, idx, issues)
        _validate_reference_italics(para, idx, issues)


def _validate_title_page_alignment(doc: Document, issues):
    intro_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == "introduction":
            intro_idx = i
            break

    if intro_idx is None:
        intro_idx = len(doc.paragraphs)

    for i in range(intro_idx):
        para = doc.paragraphs[i]
        if not para.text.strip():
            continue
        if not _is_centered(para):
            _add_issue(issues, "Title Page", f"Paragraph {i} on title page should be centered")


def _validate_no_excessive_empty_paragraphs(doc: Document, issues):
    paragraphs = doc.paragraphs
    for i in range(len(paragraphs) - 1):
        if not paragraphs[i].text.strip() and not paragraphs[i + 1].text.strip():
            _add_issue(issues, "Spacing", f"Paragraphs {i} and {i+1} contain consecutive empty lines")
            return


def _validate_tables(doc: Document, issues):
    for idx, table in enumerate(doc.tables):
        if table.alignment != WD_TABLE_ALIGNMENT.CENTER:
            _add_issue(issues, "Tables", f"Table {idx + 1} should be centered")

        if not table.rows:
            continue

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    bad_font = _paragraph_uses_non_apa_font(para)
                    if bad_font:
                        _add_issue(issues, "Tables", f"Table {idx + 1}, row {r_idx + 1}, cell {c_idx + 1} uses non-APA font: {bad_font}")
                        return

                    bad_size = _paragraph_uses_wrong_font_size(para)
                    if bad_size is not None:
                        _add_issue(issues, "Tables", f"Table {idx + 1}, row {r_idx + 1}, cell {c_idx + 1} font size is not 12 pt")
                        return


def validate_apa_docx(doc: Document):
    issues = []
    role_counts = Counter()

    for section_idx, section in enumerate(doc.sections):
        if abs(section.top_margin.inches - 1) > 0.01:
            _add_issue(issues, "Margins", f"Section {section_idx} top margin is not 1 inch")
        if abs(section.bottom_margin.inches - 1) > 0.01:
            _add_issue(issues, "Margins", f"Section {section_idx} bottom margin is not 1 inch")
        if abs(section.left_margin.inches - 1) > 0.01:
            _add_issue(issues, "Margins", f"Section {section_idx} left margin is not 1 inch")
        if abs(section.right_margin.inches - 1) > 0.01:
            _add_issue(issues, "Margins", f"Section {section_idx} right margin is not 1 inch")

    if not _title_page_is_vertically_centered(doc):
        _add_issue(issues, "Title Page", "Title page content should be vertically centered on the page")

    _validate_title_page_alignment(doc, issues)
    _validate_no_excessive_empty_paragraphs(doc, issues)
    _validate_tables(doc, issues)

    ref_start_idx = _find_references_start(doc)

    intro_idx = None
    for i, para in enumerate(doc.paragraphs):
        if _is_introduction_heading(para):
            intro_idx = i
            break

    if intro_idx is not None and not _heading_starts_on_new_page(doc, intro_idx):
        _add_issue(issues, "Introduction", f"Paragraph {intro_idx} Introduction heading should start on a new page")

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            role_counts["empty"] += 1
            continue

        if ref_start_idx is not None and idx == ref_start_idx:
            role_counts["references_heading"] += 1
            continue

        if ref_start_idx is not None and idx > ref_start_idx:
            role_counts["reference"] += 1
            continue

        if _is_major_section_heading_text(text):
            role_counts["heading"] += 1
            _validate_heading_paragraph(para, idx, issues)
            continue

        role_counts["body"] += 1
        _validate_body_paragraph(para, idx, issues)

    _validate_references(doc, ref_start_idx, issues)

    score = max(0, 100 - len(issues) * 2)

    return {
        "summary": {
            "score": score,
            "issue_count": len(issues),
            "role_counts": dict(role_counts),
            "table_count": len(doc.tables),
        },
        "issues": issues,
    }


def validate_apa_docx_path(path: str):
    doc = Document(path)
    return validate_apa_docx(doc)


def validate_apa_format(parsed):
    issues = []
    paragraphs = parsed.get("paragraphs", [])
    sections = parsed.get("sections", [])
    role_counts = Counter()

    def add_issue(category, message):
        issues.append({"category": category, "message": message})

    if sections:
        first_va = (sections[0].get("vertical_alignment") or "").lower()
        if "center" not in first_va:
            add_issue("Title Page", "Title page content should be vertically centered on the page")

    in_references = False
    for para in paragraphs:
        idx = para.get("index")
        text = (para.get("text") or "").strip()

        if not text:
            role_counts["empty"] += 1
            continue

        if text.lower() == "references":
            in_references = True
            role_counts["references_heading"] += 1
            continue

        if in_references:
            role_counts["reference"] += 1
            continue

        if _is_major_section_heading_text(text):
            role_counts["heading"] += 1
            continue

        role_counts["body"] += 1

    score = max(0, 100 - len(issues) * 2)

    return {
        "summary": {
            "score": score,
            "issue_count": len(issues),
            "role_counts": dict(role_counts),
            "table_count": parsed.get("table_count", 0),
        },
        "issues": issues,
    }


def validate_apa(parsed, template=None):
    return validate_apa_format(parsed)
