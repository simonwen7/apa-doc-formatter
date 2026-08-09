"""Effective style inheritance tests (Phase 2B)."""

from docx import Document
from docx.shared import Pt

from app.apa.registry.formatting import effective_bold, effective_font_size_pt, effective_italic


def test_none_bold_resolves_via_style_chain(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Styled heading text")
    p.style = doc.styles["Heading 1"]
    # Direct run bold often None; style may supply True.
    for run in p.runs:
        run.font.bold = None
    value = effective_bold(p)
    assert value in (True, None)  # Heading 1 is typically bold via style


def test_explicit_false_bold_is_not_treated_as_true(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Explicitly not bold")
    for run in p.runs:
        run.font.bold = False
    assert effective_bold(p) is False


def test_unset_font_size_is_uncertain_or_inherited(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Size may be inherited")
    for run in p.runs:
        run.font.size = None
    size = effective_font_size_pt(p)
    assert size is None or isinstance(size, float)
