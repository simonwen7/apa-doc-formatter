"""Synthetic golden-corpus DOCX builders for Phase 3B commercial QA.

Documents are generated in tests (tmp) so canonical committed fixtures stay immutable.
All content is synthetic — no private user data.
"""

from __future__ import annotations

import base64
import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.apa.registry.page_numbers import ensure_page_number_in_section, ensure_title_page_is_page_one
from tests.apa.fixture_builders import _body_para, _double, _ref_para, _set_margins

# Valid 1x1 PNG (synthetic; not user content).
_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


CORPUS_DIR = Path(__file__).resolve().parent


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles[f"Heading {level}"]
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        if level == 3:
            run.italic = True
    _double(p)


def _student_title_page(
    doc: Document,
    *,
    title: str,
    authors: list[str],
    affiliations: list[str],
    course: str = "PSYCH 101: Introduction to Psychology",
    instructor: str = "Dr. Ada Instructor",
    due: str = "January 15, 2026",
) -> None:
    t = doc.add_paragraph(title)
    t.style = doc.styles["Title"]
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
    _double(t)
    for author in authors:
        p = doc.add_paragraph(author)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _double(p)
    for aff in affiliations:
        p = doc.add_paragraph(aff)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _double(p)
    for text in (course, instructor, due):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _double(p)


def _body_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
    _double(p)
    p.paragraph_format.page_break_before = True


def _refs_start(doc: Document) -> None:
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    refs.paragraph_format.page_break_before = True


def build_a_clean_basic(path: Path) -> Path:
    """Clean formatting; may have zero SAFE issues."""
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    title = "Sleep Habits and Academic Focus"
    _student_title_page(
        doc,
        title=title,
        authors=["Jordan Example"],
        affiliations=["Department of Psychology, University of Example"],
    )
    _body_title(doc, title)
    _heading(doc, "Overview", 1)
    _body_para(
        doc,
        "University students often report that sleep quality shapes daytime attention "
        "and short-term memory for course material across a full academic term.",
    )
    _heading(doc, "Sleep Quality", 2)
    _body_para(
        doc,
        "Researchers continue to examine how nightly sleep duration relates to quiz "
        "performance and sustained attention during morning lecture periods.",
    )
    _heading(doc, "Measurement Approaches", 3)
    _body_para(
        doc,
        "Self-report diaries and actigraphy watches are commonly used together in "
        "undergraduate samples when laboratory polysomnography is unavailable.",
    )
    _refs_start(doc)
    _ref_para(
        doc,
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197. "
        "https://doi.org/10.1111/j.1749-6632.2009.04416.x",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_b_misformatted(path: Path) -> Path:
    """Many SAFE formatting-only violations; Fix should clear SAFE issues."""
    doc = Document()
    _set_margins(doc, 0.5)
    title = doc.add_paragraph("the effects of sleep on memory retrieval patterns")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Jordan Example")
    doc.add_paragraph("University of Example")
    heading = doc.add_paragraph("Overview")
    heading.style = doc.styles["Heading 1"]
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body = doc.add_paragraph(
        "Sleep plays an important role in memory consolidation for university students "
        "who balance coursework, employment, and extracurricular schedules each week."
    )
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    body.paragraph_format.line_spacing = 1.0
    body.paragraph_format.first_line_indent = Inches(0)
    body.paragraph_format.space_before = Pt(12)
    body.paragraph_format.space_after = Pt(12)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2 = doc.add_paragraph("Sleep Quality")
    h2.style = doc.styles["Heading 2"]
    body2 = doc.add_paragraph(
        "Quality metrics include latency, awakenings, and next-day alertness ratings."
    )
    body2.paragraph_format.first_line_indent = Inches(0)
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    ref = doc.add_paragraph(
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197."
    )
    ref.paragraph_format.left_indent = Inches(0)
    ref.paragraph_format.first_line_indent = Inches(0)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_c_author_only(path: Path) -> Path:
    """Formatting mostly correct; AUTHOR citation/reference issues present."""
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    title = "Attention During Morning Lectures"
    _student_title_page(
        doc,
        title=title,
        authors=["Jordan Example"],
        affiliations=["Department of Psychology, University of Example"],
    )
    _body_title(doc, title)
    _heading(doc, "Overview", 1)
    _body_para(
        doc,
        "Findings differ (Smith and Jones, 2020). A direct quotation needs review: "
        '"Attention faded after lunch" (Smith, 2020).',
    )
    _refs_start(doc)
    _ref_para(
        doc,
        "Smith, A., & Jones, B. (2020). Title. Journal of Example, 1(1), 1-2. "
        "doi:10.1234/example",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_d_mixed(path: Path) -> Path:
    """SAFE formatting problems + AUTHOR review items."""
    doc = Document()
    _set_margins(doc, 0.75)
    title = doc.add_paragraph("Mixed Formatting And Citation Review Paper")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.bold = False
    doc.add_paragraph("Jordan Example")
    doc.add_paragraph("Department of Psychology, University of Example")
    doc.add_paragraph("PSYCH 101: Introduction to Psychology")
    doc.add_paragraph("Dr. Ada Instructor")
    doc.add_paragraph("January 15, 2026")
    body_title = doc.add_paragraph("Mixed Formatting And Citation Review Paper")
    body_title.paragraph_format.page_break_before = True
    h1 = doc.add_paragraph("Overview")
    h1.style = doc.styles["Heading 1"]
    body = doc.add_paragraph(
        "Findings differ (Smith and Jones, 2020). Sleep and learning remain linked "
        "for many undergraduate students in controlled classroom studies."
    )
    body.paragraph_format.first_line_indent = Inches(0)
    body.paragraph_format.line_spacing = 1.0
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.paragraph_format.page_break_before = True
    ref = doc.add_paragraph(
        "Smith, A., & Jones, B. (2020). Title. Journal of Example, 1(1), 1-2. "
        "doi:10.1234/example"
    )
    ref.paragraph_format.first_line_indent = Inches(0)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_e_complex(path: Path) -> Path:
    """Complex DOCX objects + at least one SAFE formatting problem elsewhere."""
    doc = Document()
    _set_margins(doc, 0.85)  # SAFE margin issue
    ensure_page_number_in_section(doc.sections[0], 0)
    title = "Complex Document Preservation Paper"
    _student_title_page(
        doc,
        title=title,
        authors=["Jordan Example"],
        affiliations=["Department of Psychology, University of Example"],
    )
    # Header text
    header = doc.sections[0].header.paragraphs[0]
    run = header.add_run("Student research draft — confidential synthetic")
    run.font.size = Pt(10)

    _body_title(doc, title)
    _heading(doc, "Overview", 1)
    p = _body_para(
        doc,
        "Unicode check: café, naïve, São Paulo, em—dash, en–dash, and “curly quotes”. "
        "See https://example.org/research/sleep?year=2020&v=1.2 for related materials.",
    )
    # Hyperlink-ish run (URL text only; still preserved as text)
    bold = p.add_run(" Bold run.")
    bold.bold = True
    italic = p.add_run(" Italic run.")
    italic.italic = True

    # Bookmark
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "1")
    bookmark_start.set(qn("w:name"), "overviewAnchor")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "1")
    p._p.append(bookmark_start)
    p._p.append(bookmark_end)

    _body_para(doc, "As shown in Table 1, mean scores differed by condition.")
    num = doc.add_paragraph("Table 1")
    for run in num.runs:
        run.bold = True
    doc.add_paragraph("Mean Attention Scores by Condition")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Condition"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "12"

    try:
        doc.add_paragraph("Item one", style="List Number")
        doc.add_paragraph("Item two", style="List Number")
        doc.add_paragraph("Bullet one", style="List Bullet")
    except Exception:
        doc.add_paragraph("1. Item one")
        doc.add_paragraph("• Bullet one")

    _body_para(doc, "See Figure 1 for a schematic overview.")
    fnum = doc.add_paragraph("Figure 1")
    for run in fnum.runs:
        run.bold = True
    doc.add_paragraph("Schematic Overview Chart")
    img_p = doc.add_paragraph()
    img_p.add_run().add_picture(io.BytesIO(_TINY_PNG), width=Inches(1.0))

    _refs_start(doc)
    _ref_para(
        doc,
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197.",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_f_false_positive_stress(path: Path) -> Path:
    """Valid prose with citation-like traps; formatting clean."""
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    title = "False Positive Stress Corpus"
    _student_title_page(
        doc,
        title=title,
        authors=["Jordan Example"],
        affiliations=["Department of Psychology, University of Example"],
    )
    _body_title(doc, title)
    _heading(doc, "Overview", 1)
    _body_para(
        doc,
        "Results were significant (n = 40, p < .05). The year in prose only "
        "(2020 was difficult). We studied the American Psychological Association "
        "(APA) guidelines. Course enrollment used CS240. Model M-12 and version "
        "v2.1 were documented. See Table 1 for means. Ordinary parentheses "
        "(including notes) should not become citations. "
        "https://example.org/archive/2019/report remains a URL with a year path.",
    )
    # Long paragraph >40 words without being a block quote falsely
    _body_para(
        doc,
        "Students described routines that included early classes, laboratory "
        "sessions, part-time work, commuting, and evening study blocks that "
        "together reduced opportunities for consistent sleep across weekdays "
        "and weekends in this synthetic sample narrative for analyzer stress.",
    )
    _refs_start(doc)
    _ref_para(
        doc,
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197.",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_g_multi_author(path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    title = "Collaborative Student Title Page Paper"
    _student_title_page(
        doc,
        title=title,
        authors=["Jordan Example", "Alex Partner"],
        affiliations=["Department of Psychology, University of Example"],
        course="CS 240",
        instructor="Professor Jane Smith",
        due="15 January 2026",
    )
    _body_title(doc, title)
    _heading(doc, "Overview", 1)
    _body_para(
        doc,
        "Multiple authors shared one affiliation on this synthetic student title page.",
    )
    _refs_start(doc)
    _ref_para(
        doc,
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197.",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_h_reference_types(path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    title = "Reference Type Corpus Paper"
    _student_title_page(
        doc,
        title=title,
        authors=["Jordan Example"],
        affiliations=["Department of Psychology, University of Example"],
    )
    _body_title(doc, title)
    _heading(doc, "Overview", 1)
    _body_para(
        doc,
        "This corpus lists synthetic bibliographic examples for classification QA.",
    )
    _refs_start(doc)
    examples = [
        "Smith, A. (2020). Journal title here. Journal of Example, 12(3), 45-60. https://doi.org/10.1234/abcd",
        "Brown, B. (2019). Book title of learning. Academic Press.",
        "Chen, C. (2021). Ebook title online [e-book]. Open Academic Press. https://example.org/ebook",
        "Diaz, D. (2020, January 15). City news brief. The New York Times. https://nytimes.example/a",
        "Evans, E. (2020, March 1). Magazine feature. Time magazine. https://time.example/a",
        "Foster, F. (2021). Webpage guidance. Example Institute. https://example.org/guide",
        "U.S. Department of Health. (2020). Sleep report (Report No. 1). Government Printing Office.",
        "American Psychological Association. (2020). Style technical report. APA.",
        "Garcia, G. (2022). Preprint findings. PsyArXiv. https://doi.org/10.31234/osf.io/zzzz",
        "Harris, H. (2021, May 1). How sleep works [Video]. YouTube. https://www.youtube.com/watch?v=abc123",
        "APA [@APA]. (2020, January 1). Sleep tip [Tweet]. Twitter. https://twitter.com/APA/status/1",
        "Ibrahim, I. (2018). Dissertation title (Doctoral dissertation). University of Example.",
    ]
    for ex in examples:
        _ref_para(doc, ex)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


CORPUS_BUILDERS = {
    "A_clean_basic": {
        "builder": build_a_clean_basic,
        "expect_fix": False,
        "expect_safe_zero_before": True,
        "description": "Clean basic student paper",
    },
    "B_misformatted": {
        "builder": build_b_misformatted,
        "expect_fix": True,
        "expect_safe_zero_before": False,
        "description": "Heavily misformatted SAFE-only focus",
    },
    "C_author_only": {
        "builder": build_c_author_only,
        "expect_fix": False,
        "expect_safe_zero_before": True,
        "expect_author_min": 1,
        "description": "Author-review issues with clean formatting",
    },
    "D_mixed": {
        "builder": build_d_mixed,
        "expect_fix": True,
        "expect_safe_zero_before": False,
        "expect_author_min": 1,
        "description": "Mixed SAFE + AUTHOR",
    },
    "E_complex": {
        "builder": build_e_complex,
        "expect_fix": True,
        "expect_safe_zero_before": False,
        "preserve_objects": True,
        "description": "Complex DOCX preservation stress",
    },
    "F_false_positive_stress": {
        "builder": build_f_false_positive_stress,
        "expect_fix": False,
        "expect_safe_zero_before": True,
        "forbid_false_citation_rules": True,
        "description": "False-positive stress paper",
    },
    "G_multi_author": {
        "builder": build_g_multi_author,
        "expect_fix": False,
        "expect_safe_zero_before": True,
        "description": "Multi-author student title page",
    },
    "H_reference_types": {
        "builder": build_h_reference_types,
        "expect_fix": False,
        "expect_safe_zero_before": True,
        "description": "Common reference type examples",
    },
}
