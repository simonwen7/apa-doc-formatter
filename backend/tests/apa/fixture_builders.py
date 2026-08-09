"""Helpers to build deterministic DOCX fixtures for APA Phase 2A/2B tests."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from app.apa.registry.page_numbers import ensure_page_number_in_section, ensure_title_page_is_page_one


FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"


def _set_margins(doc: Document, inches: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _double(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _body_para(doc: Document, text: str):
    body = doc.add_paragraph(text)
    _double(body)
    body.paragraph_format.first_line_indent = Inches(0.5)
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return body


def _ref_para(doc: Document, text: str):
    ref = doc.add_paragraph(text)
    _double(ref)
    ref.paragraph_format.left_indent = Inches(0.5)
    ref.paragraph_format.first_line_indent = Inches(-0.5)
    return ref


def build_golden_student_paper(path: Path) -> Path:
    """Minimal correct APA 7 student paper for production SAFE rules."""
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])

    title = doc.add_paragraph("The Effects of Sleep on Memory")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)

    for text in (
        "Jane Student",
        "University of Example",
        "PSY 101: Introduction to Psychology",
        "Dr. Ada Instructor",
        "January 15, 2026",
    ):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _double(p)

    # First body page: repeated title
    body_title = doc.add_paragraph("The Effects of Sleep on Memory")
    body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in body_title.runs:
        run.bold = True
    _double(body_title)
    body_title.paragraph_format.page_break_before = True

    h1 = doc.add_paragraph("Overview")
    h1.style = doc.styles["Heading 1"]
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h1.runs:
        run.bold = True
    _double(h1)

    _body_para(
        doc,
        "Sleep plays an important role in memory consolidation for university students.",
    )

    h2 = doc.add_paragraph("Sleep Quality")
    h2.style = doc.styles["Heading 2"]
    h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h2.runs:
        run.bold = True
    _double(h2)

    _body_para(
        doc,
        "Researchers continue to examine how sleep quality affects academic performance.",
    )

    h3 = doc.add_paragraph("Measurement Approaches")
    h3.style = doc.styles["Heading 3"]
    h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h3.runs:
        run.bold = True
        run.italic = True
    _double(h3)

    _body_para(
        doc,
        "Self-report and actigraphy measures are commonly used in student samples.",
    )

    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    refs.paragraph_format.page_break_before = True

    _ref_para(
        doc,
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197.",
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_failing_phase2a_paper(path: Path) -> Path:
    """Phase 2A failing formatting fixture (still used by Phase 2A tests)."""
    doc = Document()
    _set_margins(doc, 0.5)

    title = doc.add_paragraph("the effects of sleep on memory retrieval patterns")
    title.style = doc.styles["Title"]
    doc.add_paragraph("Jane Student")

    heading = doc.add_paragraph("Overview")
    heading.style = doc.styles["Heading 1"]

    body = doc.add_paragraph(
        "Sleep plays an important role in memory consolidation for university students."
    )
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    body.paragraph_format.line_spacing = 1.0
    body.paragraph_format.first_line_indent = Inches(0)
    body.paragraph_format.space_before = Pt(12)
    body.paragraph_format.space_after = Pt(12)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]

    ref = doc.add_paragraph(
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197."
    )
    ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    ref.paragraph_format.line_spacing = 1.0
    ref.paragraph_format.left_indent = Inches(0)
    ref.paragraph_format.first_line_indent = Inches(0)
    ref.paragraph_format.space_before = Pt(8)
    ref.paragraph_format.space_after = Pt(8)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_malformed_title_page(path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    title = doc.add_paragraph("the effects of sleep on memory retrieval patterns")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.bold = False
    author = doc.add_paragraph("Jane Student")
    author.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1 = doc.add_paragraph("Overview")
    h1.style = doc.styles["Heading 1"]
    _body_para(doc, "Body text continues here with enough words for classification.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_no_introduction_paper(path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    title = doc.add_paragraph("Climate Anxiety in Students")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)
    for text in (
        "Alex Author",
        "Example University",
        "PSY 210: Environmental Psychology",
        "Dr. Green",
        "March 1, 2026",
    ):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _double(p)
    body_title = doc.add_paragraph("Climate Anxiety in Students")
    body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in body_title.runs:
        run.bold = True
    _double(body_title)
    body_title.paragraph_format.page_break_before = True
    _body_para(doc, "This paper begins without an Introduction heading by design.")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    _ref_para(doc, "Author, A. A. (2020). Sample reference title. Journal, 1(1), 1-2.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_abstract_paper(path: Path, *, well_formatted: bool = True) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    title = doc.add_paragraph("Attention and Learning")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)
    for text in (
        "Sam Student",
        "Example College",
        "EDU 100: Learning Science",
        "Professor Lee",
        "April 2, 2026",
    ):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _double(p)

    abstract = doc.add_paragraph("Abstract")
    abstract.style = doc.styles["Heading 1"]
    if well_formatted:
        abstract.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in abstract.runs:
            run.bold = True
        abstract.paragraph_format.page_break_before = True
    else:
        abstract.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in abstract.runs:
            run.bold = False

    abs_body = doc.add_paragraph(
        "This abstract summarizes attention and learning research for students."
    )
    if well_formatted:
        _double(abs_body)
        abs_body.paragraph_format.first_line_indent = Inches(0)
    else:
        abs_body.paragraph_format.line_spacing = 1.0
        abs_body.paragraph_format.first_line_indent = Inches(0.5)

    body_title = doc.add_paragraph("Attention and Learning")
    body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in body_title.runs:
        run.bold = True
    _double(body_title)
    body_title.paragraph_format.page_break_before = True
    _body_para(doc, "Body text follows the abstract when an abstract is included.")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    _ref_para(doc, "Author, A. A. (2021). Attention review. Journal, 2(2), 10-20.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_heading_fixture(path: Path, *, level: int, valid: bool) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    title = doc.add_paragraph("Heading Fixture Paper")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)
    p = doc.add_paragraph("Author Name")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _double(p)

    style_name = f"Heading {level}"
    heading = doc.add_paragraph(
        "Sample Heading Label" if level <= 3 else "Sample Heading Label. Continuing body text on the same line."
    )
    heading.style = doc.styles[style_name]
    heading.paragraph_format.page_break_before = True
    _double(heading)

    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if valid else WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.bold = True if valid else False
    elif level == 2:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT if valid else WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.bold = True if valid else False
    elif level == 3:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT if valid else WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.bold = True if valid else False
            run.italic = True if valid else False
    elif level in {4, 5}:
        if valid:
            heading.paragraph_format.first_line_indent = Inches(0.5)
            for run in heading.runs:
                run.bold = True
                run.italic = level == 5
        else:
            heading.paragraph_format.first_line_indent = Inches(0)
            for run in heading.runs:
                run.bold = False
                run.italic = False

    if level <= 3:
        _body_para(doc, "Following body paragraph after the heading.")

    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    _ref_para(doc, "Author, A. A. (2022). Heading sample. Journal, 3(3), 30-40.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_level4_split_unsafe(path: Path) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    ensure_page_number_in_section(doc.sections[0], 0)
    title = doc.add_paragraph("Split Heading Paper")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)
    h4 = doc.add_paragraph("Indented Topic")
    h4.style = doc.styles["Heading 4"]
    h4.paragraph_format.page_break_before = True
    _double(h4)
    _body_para(doc, "This body text is incorrectly separated from the Level 4 heading.")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    _ref_para(doc, "Author, A. A. (2023). Split heading. Journal, 4(4), 40-50.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_header_text_fixture(path: Path, *, with_page: bool) -> Path:
    doc = Document()
    _set_margins(doc, 1.0)
    header = doc.sections[0].header
    header.paragraphs[0].add_run("RUNNINGHEAD")
    if with_page:
        ensure_page_number_in_section(doc.sections[0], 0)
    title = doc.add_paragraph("Header Preservation Paper")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)
    _body_para(doc, "Body paragraph used to exercise page-number insertion safely.")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    _ref_para(doc, "Author, A. A. (2024). Header sample. Journal, 5(5), 50-60.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_multi_section_page_numbers(path: Path) -> Path:
    from docx.enum.section import WD_SECTION

    doc = Document()
    _set_margins(doc, 1.0)
    title = doc.add_paragraph("Multi Section Paper")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
    _double(title)
    _body_para(doc, "First section body text before the section break.")
    doc.add_section(WD_SECTION.NEW_PAGE)
    _set_margins(doc, 1.0)
    # Unlink second header so missing PAGE is detectable independently when desired.
    try:
        doc.sections[1].header.is_linked_to_previous = False
    except Exception:
        pass
    _body_para(doc, "Second section body text after the section break.")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    _double(refs)
    _ref_para(doc, "Author, A. A. (2025). Multi section. Journal, 6(6), 60-70.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path
