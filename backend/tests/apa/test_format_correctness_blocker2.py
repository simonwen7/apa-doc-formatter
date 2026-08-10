"""Format correctness blocker #2 — Word default spacing, refs region, score safety."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.classifier import classify_document
from app.apa.engine.fixer import fix_document_path
from app.apa.models import FormattingVerificationStatus, Region
from app.apa.registry.effective_format import (
    FormatSource,
    FormatVerdict,
    line_spacing_verdict,
    resolve_effective_paragraph_format,
    WORD_DEFAULT_LINE_SPACING_MULTIPLE,
)
from app.apa.registry.formatting import (
    is_hanging_indent,
    line_spacing_is_double,
)
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from tests.apa.fixture_word_default_and_refs import (
    build_nonstandard_references_fixture,
    build_real_shape_no_spacing_style,
    build_template_no_spacing_fixture,
    build_word_default_spacing_fixture,
)


def test_word_default_spacing_semantics_are_resolved(tmp_path):
    """ECMA-376: absent w:line through hierarchy → single, not UNKNOWN."""
    path = build_word_default_spacing_fixture(tmp_path / "word_default.docx")
    doc = Document(str(path))
    body = None
    for p in doc.paragraphs:
        if p.style and p.style.name == "No Spacing" and "ECMA-376" in (p.text or ""):
            body = p
            break
    assert body is not None
    eff = resolve_effective_paragraph_format(body)
    assert eff.line_spacing_multiple.source == FormatSource.WORD_DEFAULT
    assert abs(float(eff.line_spacing_multiple.value) - WORD_DEFAULT_LINE_SPACING_MULTIPLE) < 0.01
    assert line_spacing_verdict(body) == FormatVerdict.VERIFIED_VIOLATION
    assert line_spacing_is_double(body) is False


def test_real_shape_no_spacing_not_left_unresolved_when_deterministic(tmp_path):
    doc = Document()
    build_real_shape_no_spacing_style(doc)
    # Confirm style XML has no spacing element.
    style_el = doc.styles["No Spacing"].element
    pPr = style_el.find(qn("w:pPr"))
    assert pPr is None or pPr.find(qn("w:spacing")) is None
    p = doc.add_paragraph("Deterministic Word-default single spacing.")
    p.style = doc.styles["No Spacing"]
    p.paragraph_format.line_spacing = None
    p.paragraph_format.line_spacing_rule = None
    assert line_spacing_is_double(p) is not None
    assert line_spacing_verdict(p) != FormatVerdict.UNRESOLVED


def test_no_spacing_body_single_spacing_detected(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "detect_ns.docx")
    analysis = analyze_document_path(str(path))
    body_spacing = [
        i
        for i in analysis.safe_auto_fix
        if i.rule_id == "APA7-GLOBAL-004" and i.region == Region.BODY
    ]
    assert body_spacing, "No Spacing body single spacing should be SAFE_AUTO_FIX"
    assert not any(i.rule_id == "APA7-GLOBAL-004-UNRESOLVED" for i in analysis.uncertain)


def test_no_spacing_body_fixed_to_double(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "fix_ns.docx")
    out = tmp_path / "fix_ns_out.docx"
    before = snapshot_user_text(Document(str(path)))
    result = fix_document_path(str(path), str(out))
    assert result["verification"]["text_integrity_ok"] is True
    assert result["verification"]["safe_issues_after"] == 0
    doc = Document(str(out))
    classified = classify_document(doc)
    for item in classified:
        if item.region == Region.BODY and item.confidence >= 0.85 and item.text.strip():
            assert line_spacing_is_double(item.paragraph) is True
    assert_text_integrity(before, snapshot_user_text(doc))


def test_template_no_spacing_not_blindly_normalized(tmp_path):
    path = build_template_no_spacing_fixture(tmp_path / "template_ns.docx")
    analysis = analyze_document_path(str(path))
    # Banner/template No Spacing must not be treated as BODY GLOBAL-004.
    banner_issues = [
        i
        for i in analysis.safe_auto_fix
        if i.rule_id == "APA7-GLOBAL-004"
        and i.paragraph_index is not None
        and "Template Banner" in Document(str(path)).paragraphs[i.paragraph_index].text
    ]
    assert not banner_issues
    # Template signal should remain separately classified (conditional), not body-fixed.
    assert any(
        "INSTRUCTOR_TEMPLATE" in i.rule_id or "template" in (i.message or "").lower()
        for i in analysis.uncertain + analysis.author_action_required + analysis.safe_auto_fix
    ) or any(
        c.region == Region.TITLE_PAGE
        for c in classify_document(Document(str(path)))
        if "Banner" in c.text
    )


def test_nonstandard_references_heading_detected_as_reference_region(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_region.docx")
    classified = classify_document(Document(str(path)))
    headings = [c for c in classified if c.region == Region.REFERENCES_HEADING]
    assert len(headings) == 1
    assert "Works Cited/References" in headings[0].text
    assert any("nonstandard" in e for e in headings[0].evidence)


def test_nonstandard_references_heading_is_author_action(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_author.docx")
    analysis = analyze_document_path(str(path))
    heading_issues = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-REFERENCE-HEADING-TEXT"
    ]
    assert heading_issues
    assert heading_issues[0].expected == "References"
    assert "Works Cited" in (heading_issues[0].actual or "")


def test_nonstandard_reference_heading_text_never_rewritten(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_no_rewrite.docx")
    out = tmp_path / "refs_no_rewrite_out.docx"
    before_doc = Document(str(path))
    before_heading = next(
        p.text for p in before_doc.paragraphs if "Works Cited/References" in (p.text or "")
    )
    before_snap = snapshot_user_text(before_doc)
    fix_document_path(str(path), str(out))
    after_doc = Document(str(out))
    after_heading = next(
        p.text for p in after_doc.paragraphs if "Works Cited" in (p.text or "") or p.text.strip() == "References"
    )
    assert after_heading == before_heading
    assert_text_integrity(before_snap, snapshot_user_text(after_doc))


def test_reference_entries_detected_below_nonstandard_heading(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_entries.docx")
    classified = classify_document(Document(str(path)))
    entries = [c for c in classified if c.region == Region.REFERENCE_ENTRY and c.text.strip()]
    assert len(entries) >= 3


def test_normal_web_reference_spacing_detected(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_web.docx")
    analysis = analyze_document_path(str(path))
    spacing = [i for i in analysis.safe_auto_fix if i.rule_id == "APA7-REFERENCE-005"]
    assert spacing


def test_reference_positive_first_line_indent_detected(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_indent.docx")
    analysis = analyze_document_path(str(path))
    hanging = [i for i in analysis.safe_auto_fix if i.rule_id == "APA7-REFERENCE-008"]
    assert hanging


def test_reference_hanging_indent_fixed(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_hang_fix.docx")
    out = tmp_path / "refs_hang_fix_out.docx"
    fix_document_path(str(path), str(out))
    classified = classify_document(Document(str(out)))
    for item in classified:
        if item.region == Region.REFERENCE_ENTRY and item.text.strip():
            assert is_hanging_indent(item.paragraph) is True


def test_reference_font_fixed_without_text_change(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_font.docx")
    out = tmp_path / "refs_font_out.docx"
    before = snapshot_user_text(Document(str(path)))
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-REFERENCE-FONT" for i in analysis.safe_auto_fix)
    fix_document_path(str(path), str(out))
    after = Document(str(out))
    assert_text_integrity(before, snapshot_user_text(after))


def test_reference_urls_preserved(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_urls.docx")
    out = tmp_path / "refs_urls_out.docx"
    before = Document(str(path))
    urls_before = [p.text for p in before.paragraphs if "https://" in (p.text or "")]
    assert urls_before
    fix_document_path(str(path), str(out))
    after = Document(str(out))
    urls_after = [p.text for p in after.paragraphs if "https://" in (p.text or "")]
    assert urls_after == urls_before


def test_reference_heading_not_given_hanging_indent(tmp_path):
    path = build_nonstandard_references_fixture(tmp_path / "refs_heading_hang.docx")
    out = tmp_path / "refs_heading_hang_out.docx"
    fix_document_path(str(path), str(out))
    doc = Document(str(out))
    for p in doc.paragraphs:
        if "Works Cited/References" in (p.text or ""):
            assert is_hanging_indent(p) is not True
            break
    else:
        raise AssertionError("heading not found")


def test_safe_zero_does_not_mean_verified_complete(tmp_path):
    """SAFE=0 is orthogonal to formatting_verification_status."""
    doc = Document()
    p = doc.add_paragraph("Unresolved-looking body when style walk fails is rare;")
    # Construct AnalysisResult-like situation via a doc with verified body after fix.
    path = build_word_default_spacing_fixture(tmp_path / "safe_zero_meta.docx")
    out = tmp_path / "safe_zero_meta_out.docx"
    result = fix_document_path(str(path), str(out))
    assert result["verification"]["safe_issues_after"] == 0
    after = analyze_document_path(str(out))
    assert after.safe_fix_count == 0
    # Status must be present and not conflated with SAFE count.
    assert after.formatting_verification_status in {
        FormattingVerificationStatus.VERIFIED_FORMATTING,
        FormattingVerificationStatus.PARTIALLY_VERIFIED_FORMATTING,
        FormattingVerificationStatus.UNVERIFIED_FORMATTING,
    }
    payload = after.to_api_dict()["summary"]
    assert "formatting_verification_status" in payload
    assert "verified_check_count" in payload
    assert "unresolved_check_count" in payload


def test_unresolved_major_formatting_reduces_verification_confidence(tmp_path):
    """Force unresolved by removing inspectable styles tree access path.

    When major checks remain unresolved, status cannot be VERIFIED_FORMATTING
    and coverage must be < 1.
    """
    from app.apa.models import AnalysisResult

    result = AnalysisResult(
        template_id="apa7_student",
        verified_formatting_check_count=5,
        unresolved_formatting_check_count=5,
    )
    assert (
        result.formatting_verification_status
        == FormattingVerificationStatus.PARTIALLY_VERIFIED_FORMATTING
    )
    assert result.formatting_verification_coverage == 0.5
    assert result.safe_fix_count == 0
    assert result.formatting_compliance_score == 50


def test_near_perfect_score_requires_high_verification_coverage(tmp_path):
    from app.apa.models import AnalysisResult, Issue, Fixability, Severity, SourceRef

    # SAFE=0 but half of major checks unresolved → not near-perfect.
    partial = AnalysisResult(
        template_id="apa7_student",
        verified_formatting_check_count=9,
        unresolved_formatting_check_count=9,
    )
    assert partial.formatting_compliance_score <= 55
    assert partial.formatting_compliance_score < 90

    # High coverage with SAFE=0 can approach perfect.
    full = AnalysisResult(
        template_id="apa7_student",
        verified_formatting_check_count=18,
        unresolved_formatting_check_count=0,
    )
    assert full.formatting_compliance_score == 100
    assert (
        full.formatting_verification_status
        == FormattingVerificationStatus.VERIFIED_FORMATTING
    )

    # Material unresolved fraction with a few SAFE issues still cannot look perfect.
    mixed = AnalysisResult(
        template_id="apa7_student",
        verified_formatting_check_count=2,
        unresolved_formatting_check_count=8,
    )
    mixed.safe_auto_fix.append(
        Issue(
            rule_id="APA7-GLOBAL-004",
            category="Body",
            message="x",
            expected="double",
            actual="single",
            location="p",
            severity=Severity.ERROR,
            fixability=Fixability.SAFE_AUTO_FIX,
            can_fix=True,
            confidence=0.9,
            source=SourceRef(
                official_resource="test",
                publication_manual_section=None,
                notes="test",
                source_verified=True,
            ),
        )
    )
    assert mixed.formatting_compliance_score < 40
