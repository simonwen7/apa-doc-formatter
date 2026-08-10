"""Phase 2F P0 tests: fixture immutability + AUTHOR-analysis safety."""

from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from fastapi import HTTPException

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.models import Fixability
from app.apa.parsing.citations import extract_citations_from_paragraph
from app.apa.parsing.group_authors import (
    authors_alias_match,
    extract_introduced_abbreviation,
)
from app.apa.parsing.matching import MatchStatus, match_citations_to_references
from app.apa.parsing.reference_entry import parse_reference_entry
from app.apa.registry import all_rules, get_rule
from app.apa.registry.page_numbers import (
    ensure_page_number_in_section,
    ensure_title_page_is_page_one,
)
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from tests.apa.fixture_builders import FIXTURES_DIR


def _base_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    return doc


def _student_shell(doc: Document, *, title: str = "Phase 2F Fixture Paper") -> None:
    """Title page + body heading so later paragraphs classify as BODY."""
    t = doc.add_paragraph(title)
    t.style = doc.styles["Title"]
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
    author = doc.add_paragraph("Sample Author")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_paragraph("Overview")
    heading.style = doc.styles["Heading 1"]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.bold = True
    heading.paragraph_format.page_break_before = True


def _add_references_heading(doc: Document):
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    refs.paragraph_format.page_break_before = True
    return refs


def _fix_preserving_text(src: Path, out: Path) -> None:
    """Run Fix when possible; character content must remain unchanged either way."""
    before_text = [p.text for p in Document(str(src)).paragraphs]
    before_snap = snapshot_user_text(Document(str(src)))
    try:
        fix_document_path(str(src), str(out))
        after_doc = Document(str(out))
        assert_text_integrity(before_snap, snapshot_user_text(after_doc))
        assert [p.text for p in after_doc.paragraphs] == before_text
    except HTTPException:
        assert [p.text for p in Document(str(src)).paragraphs] == before_text


def test_test_suite_does_not_mutate_canonical_fixtures(canonical_fixture_fingerprints):
    before = dict(canonical_fixture_fingerprints)
    after = {
        path.name: hashlib.sha256(path.read_bytes()).hexdigest()
        for path in sorted(FIXTURES_DIR.glob("*.docx"))
    }
    assert after == before
    from tests.apa.fixture_builders import build_golden_student_paper

    tmp = Path("/tmp") / "forma_apa_fixture_immutability.docx"
    build_golden_student_paper(tmp)
    after2 = {
        path.name: hashlib.sha256(path.read_bytes()).hexdigest()
        for path in sorted(FIXTURES_DIR.glob("*.docx"))
    }
    assert after2 == before


def test_group_author_full_name_matches_reference():
    cite = extract_citations_from_paragraph(
        0, "Guidelines exist (World Health Organization, 2020)."
    )[0]
    ref = parse_reference_entry(
        1,
        "World Health Organization. (2020). Sleep guidance. https://doi.org/10.1234/abc",
    )
    matches = match_citations_to_references([cite], [ref])
    assert matches[0].status == MatchStatus.MATCHED


def test_group_author_abbreviation_alias_matching():
    intro = extract_introduced_abbreviation(
        "American Psychological Association (APA) published guidance."
    )
    assert intro is not None
    ok, conf, reason = authors_alias_match(
        "APA",
        "American Psychological Association",
        [intro],
    )
    assert ok and conf >= 0.85 and "alias" in reason


def test_unrelated_acronym_not_group_author():
    intro = extract_introduced_abbreviation("Students discussed ADHD (attention concerns).")
    assert intro is None or intro.confidence < 0.85


def test_same_author_same_year_detected(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Findings differ (Smith, 2020).")
    _add_references_heading(doc)
    doc.add_paragraph("Smith, A. (2020). First study. Journal, 1(1), 1-2.")
    doc.add_paragraph("Smith, A. (2020). Second study. Journal, 1(1), 3-4.")
    path = tmp_path / "same_year.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    ids = {i.rule_id for i in analysis.author_action_required}
    assert "APA7-REFCONTENT-SAME-YEAR-SUFFIX" in ids


def test_same_year_suffix_never_inserted(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text (Smith, 2020).")
    _add_references_heading(doc)
    doc.add_paragraph("Smith, A. (2020). First. Journal, 1(1), 1-2.")
    doc.add_paragraph("Smith, A. (2020). Second. Journal, 1(1), 3-4.")
    src = tmp_path / "sy.docx"
    out = tmp_path / "sy_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)
    target = out if out.exists() else src
    joined = "\n".join(p.text for p in Document(str(target)).paragraphs)
    assert "(2020a)" not in joined and "(2020b)" not in joined


def test_no_author_case_supported(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("See the report (Sleep and Learning, 2019).")
    _add_references_heading(doc)
    doc.add_paragraph("Sleep and learning in college. (2019). Example Press.")
    path = tmp_path / "noauth.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert all(not i.can_fix for i in analysis.all_issues() if "NO-AUTHOR" in i.rule_id)


def test_no_author_text_never_rewritten(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Body paragraph.")
    _add_references_heading(doc)
    doc.add_paragraph("Sleep and learning in college. (2019). Example Press.")
    src = tmp_path / "na.docx"
    out = tmp_path / "na_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)


def test_no_date_case_supported(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph(
        "Guidance exists (Centers for Disease Control and Prevention, n.d.)."
    )
    _add_references_heading(doc)
    doc.add_paragraph(
        "Centers for Disease Control and Prevention. (n.d.). Sleep tips. "
        "https://example.org/sleep"
    )
    path = tmp_path / "nd.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert analysis is not None


def test_no_date_marker_never_inserted(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Body cites Smith.")
    _add_references_heading(doc)
    doc.add_paragraph("Smith, A. Sleep paper. Journal of Example.")
    src = tmp_path / "ndmiss.docx"
    out = tmp_path / "ndmiss_out.docx"
    doc.save(src)
    before = [p.text for p in Document(str(src)).paragraphs]
    _fix_preserving_text(src, out)
    after = before if not out.exists() else [p.text for p in Document(str(out)).paragraphs]
    assert before == after
    assert not any("n.d." in t for t in after)


def test_personal_communication_does_not_require_reference(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph(
        "The finding was confirmed (J. Doe, personal communication, May 1, 2020)."
    )
    _add_references_heading(doc)
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal, 1(1), 1-2.")
    path = tmp_path / "pc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    missing = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-MATCH-CITATION-WITHOUT-REF"
    ]
    assert missing == []


def test_personal_communication_reference_entry_author_warning(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Confirmed (J. Doe, personal communication, May 1, 2020).")
    _add_references_heading(doc)
    doc.add_paragraph("Doe, J. (personal communication, 2020). Notes.")
    path = tmp_path / "pc_ref.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-MATCH-PERSONAL-COMM-IN-REFS"
        for i in analysis.author_action_required
    )


def test_personal_communication_reference_never_deleted(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Confirmed (J. Doe, personal communication, May 1, 2020).")
    _add_references_heading(doc)
    doc.add_paragraph("Doe, J. (personal communication, 2020). Notes.")
    src = tmp_path / "pc_del.docx"
    out = tmp_path / "pc_del_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)


def test_same_author_reference_order_analyzed(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text (Smith, 2021; Smith, 2019).")
    _add_references_heading(doc)
    doc.add_paragraph("Smith, A. (2021). Later. Journal, 1(1), 1-2.")
    doc.add_paragraph("Smith, A. (2019). Earlier. Journal, 1(1), 3-4.")
    path = tmp_path / "order.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-REFERENCE-SAME-AUTHOR-ORDER"
        for i in analysis.author_action_required
    )


def test_reference_order_never_auto_changed(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text.")
    _add_references_heading(doc)
    doc.add_paragraph("Smith, A. (2021). Later. Journal, 1(1), 1-2.")
    doc.add_paragraph("Smith, A. (2019). Earlier. Journal, 1(1), 3-4.")
    src = tmp_path / "ord.docx"
    out = tmp_path / "ord_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)


def test_doi_presentation_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text (Smith, 2020).")
    _add_references_heading(doc)
    doc.add_paragraph(
        "Smith, A. (2020). Title. Journal of Example, 1(1), 1-2. doi:10.1234/example"
    )
    path = tmp_path / "doi.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-REFCONTENT-DOI-PRESENTATION"
    ]
    assert hits
    assert all(i.can_fix is False for i in hits)


def test_doi_never_rewritten(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text (Smith, 2020).")
    _add_references_heading(doc)
    doi_line = (
        "Smith, A. (2020). Title. Journal of Example, 1(1), 1-2. doi:10.1234/example"
    )
    doc.add_paragraph(doi_line)
    src = tmp_path / "doi2.docx"
    out = tmp_path / "doi2_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)
    texts = [p.text for p in Document(str(out if out.exists() else src)).paragraphs]
    assert doi_line in texts


def test_url_presentation_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text (Smith, 2020).")
    _add_references_heading(doc)
    doc.add_paragraph(
        "Smith, A. (2020). Title. Example Site. https://example.org/page."
    )
    path = tmp_path / "url.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-REFCONTENT-URL-PRESENTATION"
        for i in analysis.author_action_required
    )


def test_url_never_rewritten(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Text.")
    _add_references_heading(doc)
    text = "Smith, A. (2020). Title. Example Site. https://example.org/page?x=1."
    doc.add_paragraph(text)
    src = tmp_path / "url2.docx"
    out = tmp_path / "url2_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)
    texts = [p.text for p in Document(str(out if out.exists() else src)).paragraphs]
    assert text in texts


def test_heading_hierarchy_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    h1 = doc.add_paragraph("Method")
    h1.style = doc.styles["Heading 1"]
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h1.runs:
        run.bold = True
    h3 = doc.add_paragraph("Participants Detail")
    h3.style = doc.styles["Heading 3"]
    path = tmp_path / "hier.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-HEADING-HIERARCHY"
    ]
    assert hits
    assert all(i.can_fix is False for i in hits)


def test_low_confidence_heading_hierarchy_not_definite(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Maybe heading")
    doc.add_paragraph("Another maybe")
    path = tmp_path / "lowhier.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert not any(
        i.rule_id == "APA7-HEADING-HIERARCHY" and i.confidence >= 0.9
        for i in analysis.author_action_required
    )


def test_heading_levels_never_auto_changed(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    h1 = doc.add_paragraph("Method")
    h1.style = doc.styles["Heading 1"]
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h1.runs:
        run.bold = True
    h3 = doc.add_paragraph("Participants Detail")
    h3.style = doc.styles["Heading 3"]
    src = tmp_path / "hl.docx"
    out = tmp_path / "hl_out.docx"
    doc.save(src)
    before_styles = [p.style.name for p in Document(str(src)).paragraphs]
    before_text = [p.text for p in Document(str(src)).paragraphs]
    try:
        fix_document_path(str(src), str(out))
        after = Document(str(out))
        assert [p.style.name for p in after.paragraphs] == before_styles
        assert [p.text for p in after.paragraphs] == before_text
    except HTTPException:
        assert [p.style.name for p in Document(str(src)).paragraphs] == before_styles


def test_heading4_title_case_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    h = doc.add_paragraph("this is mostly lowercase words for heading.")
    h.style = doc.styles["Heading 4"]
    path = tmp_path / "h4.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-HEADING4-TITLECASE" for i in analysis.author_action_required
    )


def test_heading5_title_case_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    h = doc.add_paragraph("this is mostly lowercase words for heading.")
    h.style = doc.styles["Heading 5"]
    path = tmp_path / "h5.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-HEADING5-TITLECASE" for i in analysis.author_action_required
    )


def test_table_title_case_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("As shown in Table 1, values differed.")
    num = doc.add_paragraph("Table 1")
    for run in num.runs:
        run.bold = True
    doc.add_paragraph("mean sleep hours by condition and night")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Condition"
    t.cell(0, 1).text = "Hours"
    t.cell(1, 0).text = "A"
    t.cell(1, 1).text = "1"
    path = tmp_path / "ttc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-TABLE-TITLE-CASE" for i in analysis.author_action_required
    )


def test_figure_title_case_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("See Figure 1.")
    num = doc.add_paragraph("Figure 1")
    for run in num.runs:
        run.bold = True
    doc.add_paragraph("sleep stage diagram overview chart")
    path = tmp_path / "ftc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert not any(
        i.rule_id == "APA7-FIGURE-TITLE-CASE" and i.can_fix
        for i in analysis.all_issues()
    )


def test_table_title_case_never_changed(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("As shown in Table 1, values differed.")
    doc.add_paragraph("Table 1")
    title = "mean sleep hours by condition and night"
    doc.add_paragraph(title)
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "2"
    src = tmp_path / "ttc2.docx"
    out = tmp_path / "ttc2_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)
    texts = [p.text for p in Document(str(out if out.exists() else src)).paragraphs]
    assert title in texts


def test_figure_title_case_never_changed(tmp_path: Path):
    title = "sleep stage diagram overview chart"
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("See Figure 1.")
    doc.add_paragraph("Figure 1")
    doc.add_paragraph(title)
    src = tmp_path / "ftc2.docx"
    out = tmp_path / "ftc2_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)
    texts = [p.text for p in Document(str(out if out.exists() else src)).paragraphs]
    assert title in texts


def test_missing_table_column_heading_never_generated(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("As shown in Table 1, values differed.")
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Mean Hours")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = ""
    t.cell(0, 1).text = "Hours"
    t.cell(1, 0).text = "A"
    t.cell(1, 1).text = "1"
    src = tmp_path / "col.docx"
    out = tmp_path / "col_out.docx"
    doc.save(src)
    analysis = analyze_document_path(str(src))
    assert any(
        i.rule_id == "APA7-TABLE-COLUMN-HEADING" for i in analysis.author_action_required
    )
    before = Document(str(src)).tables[0].cell(0, 0).text
    _fix_preserving_text(src, out)
    target = out if out.exists() else src
    assert Document(str(target)).tables[0].cell(0, 0).text == before


def test_appendix_table_numbering_analyzed(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("See Appendix A.")
    _add_references_heading(doc)
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal, 1(1), 1-2.")
    label = doc.add_paragraph("Appendix A")
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in label.runs:
        run.bold = True
    label.paragraph_format.page_break_before = True
    doc.add_paragraph("Extra Data")
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Extra Means")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "2"
    path = tmp_path / "apxt.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-APPENDIX-TABLE-NUMBERING"
        for i in analysis.author_action_required
    )


def test_appendix_figure_numbering_analyzed(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("See Appendix A.")
    _add_references_heading(doc)
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal, 1(1), 1-2.")
    label = doc.add_paragraph("Appendix A")
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.page_break_before = True
    doc.add_paragraph("Extra Figures")
    doc.add_paragraph("Figure 1")
    doc.add_paragraph("Diagram")
    path = tmp_path / "apxf.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert not any(
        i.rule_id == "APA7-APPENDIX-FIGURE-NUMBERING" and i.can_fix
        for i in analysis.all_issues()
    )


def test_appendix_identifiers_never_rewritten(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("See Appendix A.")
    _add_references_heading(doc)
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal, 1(1), 1-2.")
    label = doc.add_paragraph("Appendix A")
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in label.runs:
        run.bold = True
    label.paragraph_format.page_break_before = True
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Title")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "2"
    src = tmp_path / "apxid.docx"
    out = tmp_path / "apxid_out.docx"
    doc.save(src)
    _fix_preserving_text(src, out)


def test_et_al_uncertain_without_confident_reference_match(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    doc.add_paragraph("Findings differ (Smith et al., 2020).")
    _add_references_heading(doc)
    doc.add_paragraph("Jones, A. (2019). Other. Journal, 1(1), 1-2.")
    path = tmp_path / "etal.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    etal = [i for i in analysis.uncertain if i.rule_id == "APA7-CITATION-ETAL-REVIEW"]
    assert etal
    assert all(i.fixability == Fixability.CONDITIONAL for i in etal)


def test_ambiguous_match_not_reported_as_missing_reference():
    c1 = extract_citations_from_paragraph(0, "Results vary (Smith, 2020).")[0]
    r1 = parse_reference_entry(1, "Smith, A. (2020). First. Journal, 1(1), 1-2.")
    r2 = parse_reference_entry(2, "Smith, B. (2020). Second. Journal, 1(1), 3-4.")
    matches = match_citations_to_references([c1], [r1, r2])
    assert matches[0].status == MatchStatus.AMBIGUOUS
    assert matches[0].status != MatchStatus.CITATION_WITHOUT_REFERENCE


def test_policy_registry_rules_do_not_emit_document_issues(tmp_path: Path):
    doc = _base_doc()
    _student_shell(doc)
    try:
        doc.add_paragraph("Item one", style="List Number")
        doc.add_paragraph("Item two", style="List Number")
    except Exception:
        doc.add_paragraph("1. Item one")
    path = tmp_path / "lists.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    spam = [
        i
        for i in analysis.all_issues()
        if i.rule_id in {"APA7-LIST-INDENT", "APA7-LIST-NO-TEXT-REWRITE"}
    ]
    assert spam == []


def test_author_action_rules_have_no_fixer():
    for rule in all_rules():
        if rule.fixability == Fixability.AUTHOR_ACTION_REQUIRED:
            assert rule.fixer is None
            assert rule.fixer_implemented is False


def test_phase2f_post_fix_safe_zero(tmp_path: Path, failing_docx: Path):
    out1 = tmp_path / "f1.docx"
    out2 = tmp_path / "f2.docx"
    shutil.copy2(failing_docx, tmp_path / "in.docx")
    r1 = fix_document_path(str(tmp_path / "in.docx"), str(out1))
    assert r1["verification"]["verified"] is True
    assert r1["verification"]["safe_issues_after"] == 0
    assert r1["verification"]["text_integrity_ok"] is True
    assert r1["verification"]["document_preservation_ok"] is True
    r2 = fix_document_path(str(out1), str(out2))
    assert r2["fixed_counts"]["safe_rules_applied"] == 0


def test_reference_007_deprecated_not_safe_production():
    rule = get_rule("APA7-REFERENCE-007")
    assert rule is not None
    assert rule.production_supported is False
    assert rule.fixer is None
