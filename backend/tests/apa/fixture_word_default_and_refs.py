"""Synthetic fixtures for Word-default spacing + nonstandard reference regions.

Harmless placeholder text only — never real student content.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _ensure_style(doc: Document, name: str, *, based_on: str = "Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        try:
            style.base_style = doc.styles[based_on]
        except Exception:
            pass
        return style


def _strip_style_paragraph_spacing(style) -> None:
    """Match real-file No Spacing shape: no w:pPr spacing / no basedOn line."""
    el = style.element
    based = el.find(qn("w:basedOn"))
    if based is not None:
        el.remove(based)
    pPr = el.find(qn("w:pPr"))
    if pPr is not None:
        el.remove(pPr)


def _set_normal_web_like(style) -> None:
    """Reproduce Normal (Web)-like single spacing + Times 10."""
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "100")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)


def build_real_shape_no_spacing_style(doc: Document):
    """
    Build/mutate a No Spacing style matching the real paper OOXML shape:

    - no w:basedOn
    - no w:pPr / w:spacing
    - empty docDefaults pPrDefault interaction → Word single (ECMA §17.3.1.33)

    python-docx's built-in No Spacing usually has explicit line=240; strip that.
    """
    style = doc.styles["No Spacing"]
    _strip_style_paragraph_spacing(style)
    return style


def build_word_default_spacing_fixture(path: Path) -> Path:
    """Minimal fixture: real-shape No Spacing body paragraph, empty docDefaults."""
    doc = Document()
    # Ensure docDefaults pPrDefault has no spacing (typical).
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is not None:
        ppr_default = dd.find(qn("w:pPrDefault"))
        if ppr_default is not None:
            for child in list(ppr_default):
                ppr_default.remove(child)

    build_real_shape_no_spacing_style(doc)
    title = doc.add_paragraph("Synthetic Word Default Spacing Paper")
    for run in title.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    doc.add_paragraph("Alex Example")
    body = doc.add_paragraph(
        "This body paragraph uses real-shape No Spacing with no w:line anywhere "
        "in the hierarchy, so Word defaults to single spacing (ECMA-376)."
    )
    body.style = doc.styles["No Spacing"]
    body.paragraph_format.line_spacing = None
    body.paragraph_format.line_spacing_rule = None
    body.paragraph_format.first_line_indent = Inches(0.5)
    for run in body.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def build_template_no_spacing_fixture(path: Path) -> Path:
    """Instructor/template No Spacing material that must not be body-normalized."""
    doc = Document()
    build_real_shape_no_spacing_style(doc)

    banner = doc.add_paragraph("Course Template Banner — Do Not Auto-Normalize")
    banner.style = doc.styles["No Spacing"]
    for run in banner.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x80)
        run.font.bold = True
    pPr = banner._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "008080")
    pPr.append(shd)

    title = doc.add_paragraph("Synthetic Template Spacing Paper")
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = title._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "008080")
    pPr.append(shd)

    doc.add_paragraph("Alex Example")
    doc.add_paragraph("Example University")
    doc.add_paragraph("PSY 101")
    doc.add_paragraph("Dr. Example")
    doc.add_paragraph("January 1, 2026")

    heading = doc.add_paragraph("Method")
    heading.style = doc.styles["Heading 1"]
    heading.paragraph_format.page_break_before = True

    body = doc.add_paragraph(
        "This Normal body paragraph is explicitly double spaced as a control."
    )
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    body.paragraph_format.line_spacing = 2.0
    body.paragraph_format.first_line_indent = Inches(0.5)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def build_nonstandard_references_fixture(path: Path) -> Path:
    """
    Final section headed Works Cited/References with Normal (Web)-like entries,
    positive first-line indent, URLs, and body citations for matching evidence.
    """
    doc = Document()
    build_real_shape_no_spacing_style(doc)

    title = doc.add_paragraph("Synthetic Nonstandard References Paper")
    for run in title.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    doc.add_paragraph("Alex Example")
    doc.add_paragraph("Example University")
    doc.add_paragraph("PSY 101")
    doc.add_paragraph("Dr. Example")
    doc.add_paragraph("January 1, 2026")

    heading = doc.add_paragraph("Discussion")
    heading.style = doc.styles["Heading 1"]
    heading.paragraph_format.page_break_before = True

    body = doc.add_paragraph(
        "Prior work supports this claim (Example, 2020; Sample, 2021). "
        "Additional synthesis cites Demo et al. (2022) for related findings."
    )
    body.style = doc.styles["No Spacing"]
    body.paragraph_format.line_spacing = None
    body.paragraph_format.line_spacing_rule = None
    body.paragraph_format.first_line_indent = Inches(0.5)
    for run in body.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    refs_heading = doc.add_paragraph("Works Cited/References")
    refs_heading.style = doc.styles["Heading 1"]
    refs_heading.paragraph_format.page_break_before = True
    for run in refs_heading.runs:
        run.bold = True
    refs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    web = _ensure_style(doc, "Normal (Web)", based_on="Normal")
    _set_normal_web_like(web)

    entries = [
        (
            "Example, A. A. (2020). Synthetic title one. Journal of Examples, 1(1), "
            "1–10. https://doi.org/10.1234/example.2020.001"
        ),
        (
            "Sample, B. B. (2021). Synthetic title two. Example Press. "
            "https://www.example.org/sample-2021"
        ),
        (
            "Demo, C. C., Test, D. D., & Trial, E. E. (2022). Synthetic title three. "
            "Journal of Placeholders, 3(2), 20–35. https://doi.org/10.5678/demo.2022"
        ),
        (
            "Extra, F. F. (2019). Synthetic title four with URL wrap risk kept in one "
            "paragraph. https://www.example.org/path/to/resource"
        ),
    ]
    for text in entries:
        para = doc.add_paragraph(text)
        para.style = web
        para.paragraph_format.line_spacing = None
        para.paragraph_format.line_spacing_rule = None
        # Incorrect positive first-line indent (not hanging).
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.left_indent = Inches(0)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
