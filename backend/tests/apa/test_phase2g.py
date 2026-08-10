"""Phase 2G P1 tests: multi-author title page, reference types, citations, footnotes."""

from __future__ import annotations

import hashlib
import io
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from fastapi import HTTPException

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.classifier import classify_document
from app.apa.engine.fixer import fix_document_path
from app.apa.models import Fixability
from app.apa.parsing.citations import extract_citations_from_paragraph
from app.apa.parsing.docx_objects import (
    compare_document_preservation,
    snapshot_image_binaries,
    snapshot_note_ids,
)
from app.apa.parsing.reference_entry import (
    ReferenceType,
    classify_reference_type,
    parse_reference_entry,
)
from app.apa.registry import all_rules
from app.apa.registry.page_numbers import (
    ensure_page_number_in_section,
    ensure_title_page_is_page_one,
)
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from tests.apa.fixture_builders import FIXTURES_DIR


def _base_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    ensure_page_number_in_section(doc.sections[0], 0)
    ensure_title_page_is_page_one(doc.sections[0])
    return doc


def _add_centered(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _student_title(
    doc: Document,
    *,
    title: str = "Sleep and Learning",
    authors: list[str] | None = None,
    affiliations: list[str] | None = None,
    course: str = "PSYCH 101: Introduction to Psychology",
    instructor: str = "Dr. Ada Instructor",
    due: str = "January 15, 2026",
):
    t = doc.add_paragraph(title)
    t.style = doc.styles["Title"]
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
    for author in authors or ["Jane Student"]:
        _add_centered(doc, author)
    for aff in affiliations or ["Department of Psychology, University of Example"]:
        _add_centered(doc, aff)
    _add_centered(doc, course)
    _add_centered(doc, instructor)
    _add_centered(doc, due)
    body = doc.add_paragraph(title)
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in body.runs:
        run.bold = True
    body.paragraph_format.page_break_before = True


def _refs_heading(doc: Document):
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    refs.paragraph_format.page_break_before = True
    return refs


def _fix_preserving(src: Path, out: Path) -> None:
    before = [p.text for p in Document(str(src)).paragraphs]
    before_snap = snapshot_user_text(Document(str(src)))
    try:
        fix_document_path(str(src), str(out))
        after = Document(str(out))
        assert_text_integrity(before_snap, snapshot_user_text(after))
        assert [p.text for p in after.paragraphs] == before
    except HTTPException:
        assert [p.text for p in Document(str(src)).paragraphs] == before


def _inject_footnote_package(
    src: Path, dest: Path, note_text: str = "Supplementary detail."
) -> Path:
    """Inject minimal footnotes.xml + relationship + reference mark into a DOCX."""
    buf = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(buf, "w") as zout:
        namelist = zin.namelist()
        for item in namelist:
            data = zin.read(item)
            if item == "word/document.xml":
                xml = data.decode("utf-8")
                marker = (
                    '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                    '<w:footnoteReference w:id="1"/></w:r>'
                )
                xml = xml.replace("</w:p>", marker + "</w:p>", 1)
                data = xml.encode("utf-8")
            elif item == "word/_rels/document.xml.rels":
                xml = data.decode("utf-8")
                if "footnotes.xml" not in xml:
                    rel = (
                        '<Relationship Id="rIdFootnotes" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
                        'Target="footnotes.xml"/>'
                    )
                    xml = xml.replace("</Relationships>", rel + "</Relationships>")
                data = xml.encode("utf-8")
            zout.writestr(item, data)
        if "word/footnotes.xml" not in namelist:
            footnotes = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>{note_text}</w:t></w:r></w:p></w:footnote>
</w:footnotes>
"""
            zout.writestr("word/footnotes.xml", footnotes.encode("utf-8"))

    raw = buf.getvalue()
    out_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(raw), "r") as zin, zipfile.ZipFile(out_buf, "w") as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "[Content_Types].xml" and "footnotes" not in data.decode("utf-8"):
                xml = data.decode("utf-8")
                override = (
                    '<Override PartName="/word/footnotes.xml" '
                    'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                )
                xml = xml.replace("</Types>", override + "</Types>")
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    dest.write_bytes(out_buf.getvalue())
    return dest


def test_full_suite_leaves_working_tree_fixture_clean(canonical_fixture_fingerprints):
    before = dict(canonical_fixture_fingerprints)
    after = {
        p.name: hashlib.sha256(p.read_bytes()).hexdigest()
        for p in sorted(FIXTURES_DIR.glob("*.docx"))
    }
    assert after == before


def test_multiple_authors_title_page_supported(tmp_path: Path):
    doc = _base_doc()
    _student_title(
        doc,
        authors=["Jane Student", "John Scholar"],
        affiliations=["Department of Psychology, University of Example"],
    )
    path = tmp_path / "multi.docx"
    doc.save(path)
    classified = classify_document(Document(str(path)))
    authors = [c for c in classified if c.title_element == "author"]
    assert len(authors) >= 2
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-TITLE-MULTI-AUTHOR" for i in analysis.author_action_required
    )


def test_multiple_author_text_never_changed(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc, authors=["Jane Student", "John Scholar"])
    src = tmp_path / "ma.docx"
    out = tmp_path / "ma_out.docx"
    doc.save(src)
    _fix_preserving(src, out)


def test_multiple_affiliations_never_invented(tmp_path: Path):
    doc = _base_doc()
    _student_title(
        doc,
        authors=["Jane Student", "John Scholar"],
        affiliations=[
            "Department of Psychology, University of Example",
            "School of Education, University of Example",
        ],
    )
    src = tmp_path / "aff.docx"
    out = tmp_path / "aff_out.docx"
    doc.save(src)
    before = [p.text for p in Document(str(src)).paragraphs]
    analysis = analyze_document_path(str(src))
    assert any(
        i.rule_id == "APA7-TITLE-MULTI-AFFILIATION"
        for i in analysis.author_action_required
    )
    _fix_preserving(src, out)
    assert [p.text for p in Document(str(src)).paragraphs] == before


def test_affiliation_parser_handles_valid_variants(tmp_path: Path):
    doc = _base_doc()
    _student_title(
        doc,
        affiliations=["Psychology Program, Example Institute"],
        course="CS240",
        instructor="Professor Jane Smith",
        due="15 January 2026",
    )
    path = tmp_path / "affvar.docx"
    doc.save(path)
    classified = classify_document(Document(str(path)))
    elements = {c.title_element for c in classified if c.title_element}
    assert "affiliation" in elements
    assert "course" in elements
    analysis = analyze_document_path(str(path))
    missing = {
        i.rule_id
        for i in analysis.author_action_required
        if i.rule_id in {"APA7-TITLE-013", "APA7-TITLE-014", "APA7-TITLE-015"}
    }
    assert "APA7-TITLE-013" not in missing


def test_newspaper_reference_classified():
    t, c, _ = classify_reference_type(
        "Smith, A. (2020, January 15). City news. The New York Times. https://nytimes.com/a"
    )
    assert t == ReferenceType.NEWSPAPER_ARTICLE and c >= 0.85


def test_magazine_reference_classified():
    t, c, _ = classify_reference_type(
        "Smith, A. (2020, March 1). Feature story. Time magazine. https://time.com/a"
    )
    assert t == ReferenceType.MAGAZINE_ARTICLE and c >= 0.8


def test_ebook_reference_classified():
    t, c, _ = classify_reference_type(
        "Smith, A. (2020). Learning online [e-book]. Academic Press. https://example.org/ebook"
    )
    assert t == ReferenceType.EBOOK and c >= 0.8


def test_preprint_reference_classified():
    t, c, _ = classify_reference_type(
        "Smith, A. (2020). Sleep preprint. PsyArXiv. https://doi.org/10.31234/osf.io/abc"
    )
    assert t == ReferenceType.PREPRINT and c >= 0.85


def test_government_report_classified():
    t, c, _ = classify_reference_type(
        "U.S. Department of Health. (2020). Sleep report (Report No. 1). "
        "Government Printing Office."
    )
    assert t == ReferenceType.GOVERNMENT_REPORT and c >= 0.85


def test_organization_report_classified():
    t, c, _ = classify_reference_type(
        "American Psychological Association. (2020). Style technical report. APA."
    )
    assert t == ReferenceType.ORGANIZATION_REPORT and c >= 0.8


def test_online_video_reference_classified():
    t, c, _ = classify_reference_type(
        "Smith, A. (2020, January 1). How sleep works [Video]. YouTube. "
        "https://www.youtube.com/watch?v=abc"
    )
    assert t == ReferenceType.ONLINE_VIDEO and c >= 0.85


def test_social_media_reference_classified():
    t, c, _ = classify_reference_type(
        "APA [@APA]. (2020, January 1). Sleep tip [Tweet]. Twitter. "
        "https://twitter.com/APA/status/1"
    )
    assert t == ReferenceType.SOCIAL_MEDIA and c >= 0.85


def test_ambiguous_reference_stays_unknown():
    t, c, _ = classify_reference_type("Something unclear without markers.")
    assert t == ReferenceType.UNKNOWN
    assert c < 0.6


def test_journal_source_completeness_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Evidence exists (Smith, 2020).")
    _refs_heading(doc)
    doc.add_paragraph(
        "Smith, A. (2020). Title only here 1(1), 1-2. https://doi.org/10.1234/a"
    )
    path = tmp_path / "jsrc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id
        in {"APA7-REFCONTENT-JOURNAL-SOURCE", "APA7-REFCONTENT-JOURNAL-VOLUME"}
    ]
    assert all(i.can_fix is False for i in hits)


def test_book_source_completeness_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("See the book (Smith, 2020).")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Learning online [e-book].")
    path = tmp_path / "bsrc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-REFCONTENT-BOOK-SOURCE"
    ]
    assert hits
    assert all(not i.can_fix for i in hits)


def test_webpage_source_completeness_author_only(tmp_path: Path):
    a = parse_reference_entry(0, "Smith, A. (2020). Page title.")
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Online guidance (Smith, 2020).")
    _refs_heading(doc)
    doc.add_paragraph(
        "Smith, A. (2020). Page title. Retrieved January 1, 2021, from example"
    )
    path = tmp_path / "wsrc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-REFCONTENT-WEBPAGE-SOURCE"
    ]
    assert all(not i.can_fix for i in hits)
    _ = a


def test_report_source_completeness_author_only(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Agency guidance (World Health Organization, 2020).")
    _refs_heading(doc)
    doc.add_paragraph("World Health Organization. (2020). Sleep technical report.")
    path = tmp_path / "rsrc.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-REFCONTENT-REPORT-SOURCE"
    ]
    assert all(not i.can_fix for i in hits)


def test_one_author_parenthetical_analysis(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Findings differ (Smith, 2020).")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Title. Journal of Example, 1(1), 1-2.")
    path = tmp_path / "onep.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-CITATION-ONE-AUTHOR" for i in analysis.author_action_required
    )


def test_one_author_narrative_analysis(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Smith (2020) reported sleep effects.")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Title. Journal of Example, 1(1), 1-2.")
    path = tmp_path / "onen.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(
        i.rule_id == "APA7-CITATION-ONE-AUTHOR" for i in analysis.author_action_required
    )


def test_two_author_parenthetical_analysis(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Findings differ (Smith and Jones, 2020).")
    _refs_heading(doc)
    doc.add_paragraph(
        "Smith, A., & Jones, B. (2020). Title. Journal of Example, 1(1), 1-2."
    )
    path = tmp_path / "twop.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-CITATION-TWO-AUTHOR"
    ]
    assert hits
    assert all(not i.can_fix for i in hits)


def test_two_author_narrative_analysis(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Smith & Jones (2020) reported sleep effects.")
    _refs_heading(doc)
    doc.add_paragraph(
        "Smith, A., & Jones, B. (2020). Title. Journal of Example, 1(1), 1-2."
    )
    path = tmp_path / "twon.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    hits = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-CITATION-TWO-AUTHOR"
    ]
    assert hits


def test_citation_characters_never_changed(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    cite = "Findings differ (Smith and Jones, 2020)."
    doc.add_paragraph(cite)
    _refs_heading(doc)
    doc.add_paragraph(
        "Smith, A., & Jones, B. (2020). Title. Journal of Example, 1(1), 1-2."
    )
    src = tmp_path / "cite.docx"
    out = tmp_path / "cite_out.docx"
    doc.save(src)
    _fix_preserving(src, out)
    texts = [p.text for p in Document(str(out if out.exists() else src)).paragraphs]
    assert cite in texts


def test_reference_characters_never_changed(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Text (Smith, 2020).")
    _refs_heading(doc)
    ref = "Smith, A. (2020). Title. New York, NY: Academic Press."
    doc.add_paragraph(ref)
    src = tmp_path / "ref.docx"
    out = tmp_path / "ref_out.docx"
    doc.save(src)
    _fix_preserving(src, out)
    texts = [p.text for p in Document(str(out if out.exists() else src)).paragraphs]
    assert ref in texts


def test_et_al_definite_only_with_confident_author_count(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Findings differ (Smith et al., 2020).")
    _refs_heading(doc)
    doc.add_paragraph(
        "Smith, A., Jones, B., & Lee, C. (2020). Title. Journal of Example, 1(1), 1-2."
    )
    path = tmp_path / "etal_ok.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    etal = [i for i in analysis.uncertain if i.rule_id == "APA7-CITATION-ETAL-REVIEW"]
    assert etal == []


def test_reference_missing_component_not_duplicated(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Text (Smith, 2020).")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Learning science.")
    path = tmp_path / "dedupe.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    book = [
        i
        for i in analysis.author_action_required
        if i.rule_id == "APA7-REFCONTENT-BOOK-SOURCE" and i.paragraph_index is not None
    ]
    assert len(book) <= 1


def test_citation_false_positive_statistics():
    assert (
        extract_citations_from_paragraph(
            0, "Results were significant (n = 120, p < .05)."
        )
        == []
    )


def test_citation_false_positive_parenthetical_prose():
    assert (
        extract_citations_from_paragraph(0, "The year in prose only (2020 was difficult).")
        == []
    )


def test_citation_false_positive_acronym_definition():
    cites = extract_citations_from_paragraph(
        0, "We studied the American Psychological Association (APA) guidelines."
    )
    assert cites == []


def test_footnote_text_preserved(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Body with a note.")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Title. Journal of Example, 1(1), 1-2.")
    base = tmp_path / "fn_base.docx"
    doc.save(base)
    src = tmp_path / "fn.docx"
    _inject_footnote_package(base, src, note_text="Footnote alpha text.")
    before = snapshot_user_text(Document(str(src)))
    out = tmp_path / "fn_out.docx"
    fix_document_path(str(src), str(out))
    after = snapshot_user_text(Document(str(out)))
    assert_text_integrity(before, after)
    assert any("Footnote alpha text." in text for _, text in after.segments)


def test_footnote_relationships_preserved(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Body with a note.")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Title. Journal of Example, 1(1), 1-2.")
    base = tmp_path / "fnr_base.docx"
    doc.save(base)
    src = tmp_path / "fnr.docx"
    _inject_footnote_package(base, src)
    out = tmp_path / "fnr_out.docx"
    before_images = snapshot_image_binaries(Document(str(src)))
    fix_document_path(str(src), str(out))
    after_images = snapshot_image_binaries(Document(str(out)))
    report = compare_document_preservation(
        before_path=str(src),
        after_path=str(out),
        before_images=before_images,
        after_images=after_images,
        before_doc=Document(str(src)),
        after_doc=Document(str(out)),
        text_integrity_ok=True,
    )
    assert report.footnotes_preserved is True
    assert report.document_preservation_ok is True


def test_footnote_ids_preserved(tmp_path: Path):
    doc = _base_doc()
    _student_title(doc)
    doc.add_paragraph("Body with a note.")
    _refs_heading(doc)
    doc.add_paragraph("Smith, A. (2020). Title. Journal of Example, 1(1), 1-2.")
    base = tmp_path / "fni_base.docx"
    doc.save(base)
    src = tmp_path / "fni.docx"
    _inject_footnote_package(base, src)
    out = tmp_path / "fni_out.docx"
    before_ids = snapshot_note_ids(str(src))
    fix_document_path(str(src), str(out))
    after_ids = snapshot_note_ids(str(out))
    assert before_ids["footnotes"] == after_ids["footnotes"]
    assert "1" in before_ids["footnotes"]


def test_fix_elsewhere_does_not_modify_footnotes(tmp_path: Path, failing_docx: Path):
    src = tmp_path / "fail_fn.docx"
    _inject_footnote_package(failing_docx, src, note_text="Must survive formatting.")
    out = tmp_path / "fail_fn_out.docx"
    before = snapshot_user_text(Document(str(src)))
    result = fix_document_path(str(src), str(out))
    assert result["verification"]["text_integrity_ok"] is True
    assert result["verification"]["document_preservation_ok"] is True
    after = snapshot_user_text(Document(str(out)))
    assert_text_integrity(before, after)
    assert any("Must survive formatting." in text for _, text in after.segments)


def test_phase2g_author_rules_have_no_fixer():
    for rule in all_rules():
        if rule.rule_id.startswith(
            (
                "APA7-TITLE-MULTI",
                "APA7-TITLE-AFFILIATION",
                "APA7-REFCONTENT-JOURNAL",
                "APA7-REFCONTENT-BOOK",
                "APA7-REFCONTENT-WEBPAGE",
                "APA7-REFCONTENT-REPORT",
                "APA7-REFCONTENT-PREPRINT",
                "APA7-REFCONTENT-VIDEO",
                "APA7-REFCONTENT-SOCIAL",
                "APA7-REFCONTENT-PERIODICAL",
                "APA7-CITATION-ONE-AUTHOR",
                "APA7-CITATION-TWO-AUTHOR",
            )
        ):
            assert rule.fixer is None
            assert rule.fixability == Fixability.AUTHOR_ACTION_REQUIRED


def test_phase2g_post_fix_safe_zero(tmp_path: Path, failing_docx: Path):
    out1 = tmp_path / "g1.docx"
    out2 = tmp_path / "g2.docx"
    shutil.copy2(failing_docx, tmp_path / "in.docx")
    r1 = fix_document_path(str(tmp_path / "in.docx"), str(out1))
    assert r1["verification"]["verified"] is True
    assert r1["verification"]["safe_issues_after"] == 0
    assert r1["verification"]["text_integrity_ok"] is True
    assert r1["verification"]["document_preservation_ok"] is True
    r2 = fix_document_path(str(out1), str(out2))
    assert r2["fixed_counts"]["safe_rules_applied"] == 0
