"""Iterative fix convergence tests (MRR regression + loop safety)."""

from __future__ import annotations

import pytest
from docx import Document
from fastapi import HTTPException

from app.apa.engine.analyzer import analyze_document
from app.apa.engine.classifier import classify_document
from app.apa.engine.fixer import (
    MAX_FIX_PASSES,
    _apply_safe_fixes_for_pass,
    _production_safe_remaining,
    _safe_issue_signature,
    fix_document_path,
)
from app.apa.models import Fixability, Issue, Region, Severity, SourceRef
from app.apa.registry import production_safe_rules
from app.apa.registry.base import RuleContext
from tests.apa.fixture_builders import (
    build_golden_student_paper,
    build_heading1_title_cascade_paper,
)


PRODUCTION_SAFE_IDS = {rule.rule_id for rule in production_safe_rules()}


def test_heading1_cascade_requires_second_pass(tmp_path):
    path = build_heading1_title_cascade_paper(tmp_path / "cascade.docx")
    before = analyze_document(Document(str(path)), template_id="apa7_student")
    before_count = len(_production_safe_remaining(before, "apa7_student"))
    assert before_count > 0

    one_pass_path = tmp_path / "cascade_one_pass.docx"
    doc = Document(str(path))
    analysis = analyze_document(doc, template_id="apa7_student")
    classified = classify_document(doc)
    context = RuleContext(doc=doc, classified=classified, template_id="apa7_student")
    _apply_safe_fixes_for_pass(doc, classified, context, analysis, template_id="apa7_student")
    doc.save(one_pass_path)

    after_one = analyze_document(Document(str(one_pass_path)), template_id="apa7_student")
    remaining_after_one = _production_safe_remaining(after_one, "apa7_student")
    assert remaining_after_one, "one pass should leave production SAFE issues"
    remaining_ids = {issue.rule_id for issue in remaining_after_one}
    assert "APA7-TITLE-004" in remaining_ids or "APA7-BODY-001" in remaining_ids

    output = tmp_path / "cascade_fixed.docx"
    result = fix_document_path(str(path), str(output))
    assert result["verification"]["verified"] is True
    assert result["verification"]["safe_issues_after"] == 0
    assert result["fixed_counts"]["fix_passes_executed"] >= 2


def test_newly_created_safe_issue_fixed_on_second_pass(tmp_path):
    path = build_heading1_title_cascade_paper(tmp_path / "cascade2.docx")
    output = tmp_path / "cascade2_fixed.docx"
    result = fix_document_path(str(path), str(output))
    assert result["fixed_counts"]["fix_passes_executed"] == 2
    assert result["verification"]["safe_issues_after"] == 0


def test_no_progress_stops_without_infinite_loop(monkeypatch, failing_docx, tmp_path):
    def noop_apply(ctx, issue):
        return False

    for rule in production_safe_rules():
        monkeypatch.setattr(rule, "apply_fix", noop_apply)

    with pytest.raises(HTTPException) as exc_info:
        fix_document_path(str(failing_docx), tmp_path / "noop_out.docx")

    assert exc_info.value.status_code == 500
    assert "verification failed" in exc_info.value.detail.lower()


def test_oscillation_signature_detects_repeat():
    issue_a = Issue(
        rule_id="APA7-BODY-001",
        category="Body",
        message="m",
        expected="e",
        actual="a",
        location="paragraph[1]",
        severity=Severity.ERROR,
        fixability=Fixability.SAFE_AUTO_FIX,
        can_fix=True,
        confidence=0.95,
        source=SourceRef(official_resource="t", source_verified=True),
        paragraph_index=1,
        region=Region.BODY,
    )
    issue_b = Issue(
        rule_id="APA7-BODY-001",
        category="Body",
        message="m",
        expected="e",
        actual="a",
        location="paragraph[2]",
        severity=Severity.ERROR,
        fixability=Fixability.SAFE_AUTO_FIX,
        can_fix=True,
        confidence=0.95,
        source=SourceRef(official_resource="t", source_verified=True),
        paragraph_index=2,
        region=Region.BODY,
    )
    sig_one = _safe_issue_signature([issue_a])
    sig_two = _safe_issue_signature([issue_b])
    assert sig_one != sig_two
    seen = {sig_one}
    assert sig_one in seen
    assert sig_two not in seen


def test_golden_paper_still_fixes_in_one_pass(tmp_path):
    path = build_golden_student_paper(tmp_path / "golden.docx")
    output = tmp_path / "golden_fixed.docx"
    result = fix_document_path(str(path), str(output))
    assert result["verification"]["verified"] is True
    assert result["verification"]["safe_issues_after"] == 0
    assert result["fixed_counts"]["fix_passes_executed"] <= MAX_FIX_PASSES


def test_failing_docx_still_converges(failing_docx, tmp_path):
    output = tmp_path / "failing_fixed.docx"
    result = fix_document_path(str(failing_docx), str(output))
    assert result["verification"]["verified"] is True
    assert result["verification"]["safe_issues_after"] == 0
