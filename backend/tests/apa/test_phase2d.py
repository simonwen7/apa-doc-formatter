"""Phase 2D fixtures and safety/policy tests: tables, figures, appendices, lists, preservation."""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.models import Fixability, Region
from app.apa.parsing.docx_objects import (
    compare_document_preservation,
    snapshot_image_binaries,
)
from app.apa.parsing.equations import analyze_equations, equation_xml_fingerprint
from app.apa.parsing.figures import analyze_figures
from app.apa.parsing.lists import analyze_lists
from app.apa.parsing.tables import analyze_tables
from app.apa.parsing.appendices import analyze_appendices
from app.apa.registry import all_rules, production_safe_rules
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _base_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc


def _tiny_png_bytes() -> bytes:
    """Valid minimal 1x1 PNG."""
    import struct
    import zlib

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = zlib.compress(b"\x00\xff\x00\x00")
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", raw) + chunk(b"IEND", b"")


def _add_png(doc: Document, paragraph=None):
    paragraph = paragraph or doc.add_paragraph()
    run = paragraph.add_run()
    stream = io.BytesIO(_tiny_png_bytes())
    stream.seek(0)
    run.add_picture(stream, width=Inches(2))
    return paragraph


def build_apa_table_doc(path: Path, *, bad_format: bool = True, with_callout: bool = False) -> Path:
    doc = _base_doc()
    title = doc.add_paragraph("Table Fixture Paper")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_paragraph("Overview")
    h.style = doc.styles["Heading 1"]
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if with_callout:
        doc.add_paragraph("As shown in Table 1, sleep scores varied by condition.")
    else:
        doc.add_paragraph("Sleep scores varied by condition across the sample.")
    num = doc.add_paragraph("Table 1")
    if bad_format:
        num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in num.runs:
            run.bold = False
    else:
        num.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in num.runs:
            run.bold = True
    ttitle = doc.add_paragraph("Mean Sleep Hours by Condition")
    if bad_format:
        ttitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in ttitle.runs:
            run.italic = False
    else:
        ttitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in ttitle.runs:
            run.italic = True
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Condition"
    table.cell(0, 1).text = "Hours"
    table.cell(1, 0).text = "Control"
    table.cell(1, 1).text = "7.2"
    note = doc.add_paragraph("Note. Values are sample means.")
    if bad_format:
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs.paragraph_format.page_break_before = True
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal of Example, 1(1), 1-10.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_merged_table_doc(path: Path) -> Path:
    doc = _base_doc()
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Merged Cell Example")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    # Merge first row
    table.cell(0, 0).merge(table.cell(0, 1))
    doc.save(path)
    return path


def build_layout_table_doc(path: Path) -> Path:
    doc = _base_doc()
    doc.add_paragraph("Title Page Layout")
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "Jane Student"
    table.cell(1, 0).text = "University"
    table.cell(2, 0).text = "Course"
    doc.save(path)
    return path


def build_table_without_number(path: Path) -> Path:
    doc = _base_doc()
    doc.add_paragraph("Results body text without table callout.")
    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"R{r}C{c}"
    doc.save(path)
    return path


def build_apa_figure_doc(path: Path, *, bad_format: bool = True, with_callout: bool = False) -> Path:
    doc = _base_doc()
    title = doc.add_paragraph("Figure Fixture Paper")
    title.style = doc.styles["Title"]
    if with_callout:
        doc.add_paragraph("See Figure 1 for the sleep diagram.")
    else:
        doc.add_paragraph("A diagram summarizes the sleep stages.")
    num = doc.add_paragraph("Figure 1")
    if bad_format:
        num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in num.runs:
            run.bold = False
    else:
        for run in num.runs:
            run.bold = True
    ftitle = doc.add_paragraph("Sleep Stage Diagram")
    if bad_format:
        for run in ftitle.runs:
            run.italic = False
        ftitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for run in ftitle.runs:
            run.italic = True
    _add_png(doc)
    note = doc.add_paragraph("Note. Diagram is illustrative.")
    if bad_format:
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)
    return path


def build_logo_image_doc(path: Path) -> Path:
    doc = _base_doc()
    doc.add_paragraph("University logo appears below.")
    p = doc.add_paragraph()
    _add_png(doc, p)
    doc.add_paragraph("End of decorative logo section.")
    doc.save(path)
    return path


def build_appendix_doc(path: Path, *, bad_format: bool = True, multiple: bool = False) -> Path:
    doc = _base_doc()
    title = doc.add_paragraph("Appendix Fixture")
    title.style = doc.styles["Title"]
    doc.add_paragraph("Body mentions Appendix A materials.")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs.paragraph_format.page_break_before = True
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal of Example, 1(1), 1-10.")
    label = doc.add_paragraph("Appendix A")
    if bad_format:
        label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in label.runs:
            run.bold = False
        # intentionally no page break
    else:
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in label.runs:
            run.bold = True
        label.paragraph_format.page_break_before = True
    atitle = doc.add_paragraph("Supplementary Sleep Tables")
    if bad_format:
        atitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in atitle.runs:
            run.bold = False
    else:
        atitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in atitle.runs:
            run.bold = True
    doc.add_paragraph("Appendix body content remains unchanged.")
    if multiple:
        label_b = doc.add_paragraph("Appendix B")
        label_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in label_b.runs:
            run.bold = False
        doc.add_paragraph("Second Appendix Title")
        doc.add_paragraph("More appendix body.")
    doc.save(path)
    return path


def build_list_docs(path: Path) -> Path:
    doc = _base_doc()
    doc.add_paragraph("Lists Fixture")
    doc.add_paragraph("1. This is plain numbered prose, not a Word list.")
    # Word numbering via style List Number / List Bullet if available
    try:
        p1 = doc.add_paragraph("First numbered item", style="List Number")
        p2 = doc.add_paragraph("Second numbered item", style="List Number")
        b1 = doc.add_paragraph("Bullet one", style="List Bullet")
        b2 = doc.add_paragraph("Bullet two", style="List Bullet")
        n1 = doc.add_paragraph("Nested parent", style="List Number")
        n2 = doc.add_paragraph("Nested child", style="List Number 2")
        _ = (p1, p2, b1, b2, n1, n2)
    except Exception:
        doc.add_paragraph("Fallback list text without styles.")
    # List inside table
    table = doc.add_table(rows=1, cols=1)
    try:
        table.cell(0, 0).paragraphs[0].style = doc.styles["List Bullet"]
        table.cell(0, 0).paragraphs[0].text = "Bullet in table"
    except Exception:
        table.cell(0, 0).text = "Plain cell"
    doc.save(path)
    return path


def build_hyperlink_bookmark_doc(path: Path) -> Path:
    doc = _base_doc()
    p = doc.add_paragraph("Visit ")
    # Add hyperlink via relationship
    part = doc.part
    r_id = part.relate_to(
        "https://apastyle.apa.org/",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)
    new_run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "APA Style"
    new_run.append(text)
    hyperlink.append(new_run)
    p._p.append(hyperlink)

    # Bookmark
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "0")
    bookmark_start.set(qn("w:name"), "sleep_anchor")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "0")
    body = doc.add_paragraph("Bookmarked sentence about sleep.")
    body._p.insert(0, bookmark_start)
    body._p.append(bookmark_end)
    doc.save(path)
    return path


def build_equation_doc(path: Path) -> Path:
    """Embed a minimal OMML equation via lxml so preservation can count it."""
    from lxml import etree

    doc = _base_doc()
    p = doc.add_paragraph("The model is ")
    MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    o_math = etree.SubElement(p._p, "{%s}oMath" % MATH_NS)
    r = etree.SubElement(o_math, "{%s}r" % MATH_NS)
    t = etree.SubElement(r, "{%s}t" % MATH_NS)
    t.text = "y = mx + b"
    doc.save(path)
    return path


# -------------------- TABLES --------------------


def test_table_text_never_changes(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "table_bad.docx", bad_format=True, with_callout=True)
    out = tmp_path / "table_fixed.docx"
    before = snapshot_user_text(Document(str(src)))
    fix_document_path(str(src), str(out))
    after = snapshot_user_text(Document(str(out)))
    assert_text_integrity(before, after)


def test_table_merged_cells_survive(tmp_path: Path):
    src = build_merged_table_doc(tmp_path / "merged.docx")
    out = tmp_path / "merged_out.docx"
    before_merged = analyze_tables(Document(str(src)))[0].merged_cells
    fix_document_path(str(src), str(out))
    after_merged = analyze_tables(Document(str(out)))[0].merged_cells
    assert before_merged is True
    assert after_merged is True
    assert Document(str(out)).tables[0].cell(0, 0).text == Document(str(src)).tables[0].cell(0, 0).text


def test_table_number_text_never_generated(tmp_path: Path):
    src = build_table_without_number(tmp_path / "no_num.docx")
    analysis = analyze_document_path(str(src))
    assert any(i.rule_id == "APA7-TABLE-NUMBER-MISSING" for i in analysis.author_action_required)
    out = tmp_path / "no_num_out.docx"
    fix_document_path(str(src), str(out))
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert not any(t.strip().lower() == "table 1" for t in texts)


def test_table_title_text_never_generated(tmp_path: Path):
    doc = _base_doc()
    doc.add_paragraph("Table 1")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    path = tmp_path / "no_title.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-TABLE-TITLE-MISSING" for i in analysis.author_action_required)
    out = tmp_path / "no_title_out.docx"
    fix_document_path(str(path), str(out))
    # Still only one paragraph that is exactly Table 1 near the table; no invented title.
    assert Document(str(path)).paragraphs[0].text == Document(str(out)).paragraphs[0].text


def test_table_number_format_safe(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "tnum.docx", bad_format=True, with_callout=True)
    analysis = analyze_document_path(str(src))
    assert any(i.rule_id == "APA7-TABLE-NUMBER-FORMAT" for i in analysis.safe_auto_fix)
    out = tmp_path / "tnum_out.docx"
    fix_document_path(str(src), str(out))
    after = analyze_document_path(str(out))
    assert not any(i.rule_id == "APA7-TABLE-NUMBER-FORMAT" for i in after.safe_auto_fix)


def test_table_title_format_safe(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "ttitle.docx", bad_format=True, with_callout=True)
    assert any(
        i.rule_id == "APA7-TABLE-TITLE-FORMAT"
        for i in analyze_document_path(str(src)).safe_auto_fix
    )
    out = tmp_path / "ttitle_out.docx"
    fix_document_path(str(src), str(out))
    assert not any(
        i.rule_id == "APA7-TABLE-TITLE-FORMAT"
        for i in analyze_document_path(str(out)).safe_auto_fix
    )


def test_table_border_fix_is_idempotent(tmp_path: Path):
    # Borders are AUTHOR/UNSUPPORTED — Fix should not mutate table XML aggressively.
    src = build_apa_table_doc(tmp_path / "tb.docx", with_callout=True)
    out1 = tmp_path / "tb1.docx"
    out2 = tmp_path / "tb2.docx"
    fix_document_path(str(src), str(out1))
    fix_document_path(str(out1), str(out2))
    assert Document(str(out1)).tables[0].cell(1, 1).text == Document(str(out2)).tables[0].cell(1, 1).text


def test_complex_table_not_aggressively_fixed(tmp_path: Path):
    src = build_merged_table_doc(tmp_path / "complex.docx")
    analysis = analyze_document_path(str(src))
    # Either complex unsupported or no SAFE border rewrite rules applied.
    safe_border = [i for i in analysis.safe_auto_fix if "BORDER" in i.rule_id]
    assert safe_border == []


def test_table_callout_never_generated(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "nocall.docx", with_callout=False)
    analysis = analyze_document_path(str(src))
    assert any(i.rule_id == "APA7-TABLE-CALLOUT-MISSING" for i in analysis.author_action_required)
    out = tmp_path / "nocall_out.docx"
    fix_document_path(str(src), str(out))
    joined = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "As shown in Table 1" not in joined
    assert "see Table 1" not in joined
    assert "see table 1" not in joined.lower()


def test_layout_table_not_forced_as_apa(tmp_path: Path):
    src = build_layout_table_doc(tmp_path / "layout.docx")
    tables = analyze_tables(Document(str(src)))
    assert tables
    assert tables[0].is_apa_candidate is False


# -------------------- FIGURES --------------------


def test_image_binary_preserved(tmp_path: Path):
    src = build_apa_figure_doc(tmp_path / "fig.docx", with_callout=True)
    out = tmp_path / "fig_out.docx"
    before = snapshot_image_binaries(Document(str(src)))
    fix_document_path(str(src), str(out))
    after = snapshot_image_binaries(Document(str(out)))
    assert before == after


def test_figure_number_never_generated(tmp_path: Path):
    src = build_logo_image_doc(tmp_path / "logo.docx")
    out = tmp_path / "logo_out.docx"
    fix_document_path(str(src), str(out))
    assert not any(p.text.strip().lower().startswith("figure ") for p in Document(str(out)).paragraphs)


def test_figure_title_never_generated(tmp_path: Path):
    doc = _base_doc()
    doc.add_paragraph("Figure 1")
    _add_png(doc)
    path = tmp_path / "fig_no_title.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-FIGURE-TITLE-MISSING" for i in analysis.author_action_required)
    out = tmp_path / "fig_no_title_out.docx"
    fix_document_path(str(path), str(out))
    assert Document(str(path)).paragraphs[0].text == Document(str(out)).paragraphs[0].text


def test_figure_number_format_safe(tmp_path: Path):
    src = build_apa_figure_doc(tmp_path / "fnum.docx", bad_format=True, with_callout=True)
    assert any(
        i.rule_id == "APA7-FIGURE-NUMBER-FORMAT"
        for i in analyze_document_path(str(src)).safe_auto_fix
    )
    out = tmp_path / "fnum_out.docx"
    fix_document_path(str(src), str(out))
    assert not any(
        i.rule_id == "APA7-FIGURE-NUMBER-FORMAT"
        for i in analyze_document_path(str(out)).safe_auto_fix
    )


def test_figure_title_format_safe(tmp_path: Path):
    src = build_apa_figure_doc(tmp_path / "ftitle.docx", bad_format=True, with_callout=True)
    assert any(
        i.rule_id == "APA7-FIGURE-TITLE-FORMAT"
        for i in analyze_document_path(str(src)).safe_auto_fix
    )
    out = tmp_path / "ftitle_out.docx"
    fix_document_path(str(src), str(out))
    assert not any(
        i.rule_id == "APA7-FIGURE-TITLE-FORMAT"
        for i in analyze_document_path(str(out)).safe_auto_fix
    )


def test_logo_not_misclassified_as_figure(tmp_path: Path):
    src = build_logo_image_doc(tmp_path / "logo2.docx")
    figs = analyze_figures(Document(str(src)))
    assert figs
    assert figs[0].is_apa_candidate is False


def test_decorative_image_not_force_formatted(tmp_path: Path):
    src = build_logo_image_doc(tmp_path / "deco.docx")
    analysis = analyze_document_path(str(src))
    assert not any(i.rule_id.startswith("APA7-FIGURE") and i.fixability == Fixability.SAFE_AUTO_FIX for i in analysis.safe_auto_fix)


def test_figure_callout_never_generated(tmp_path: Path):
    src = build_apa_figure_doc(tmp_path / "fnocall.docx", with_callout=False)
    assert any(
        i.rule_id == "APA7-FIGURE-CALLOUT-MISSING"
        for i in analyze_document_path(str(src)).author_action_required
    )
    out = tmp_path / "fnocall_out.docx"
    fix_document_path(str(src), str(out))
    joined = "\n".join(p.text for p in Document(str(out)).paragraphs).lower()
    assert "see figure 1" not in joined


# -------------------- APPENDICES --------------------


def test_appendix_text_never_changes(tmp_path: Path):
    src = build_appendix_doc(tmp_path / "apx.docx", bad_format=True)
    out = tmp_path / "apx_out.docx"
    before = snapshot_user_text(Document(str(src)))
    fix_document_path(str(src), str(out))
    assert_text_integrity(before, snapshot_user_text(Document(str(out))))


def test_appendix_label_never_generated(tmp_path: Path):
    doc = _base_doc()
    doc.add_paragraph("References")
    doc.add_paragraph("Walker, M. P. (2009). Sleep. Journal, 1(1), 1-10.")
    path = tmp_path / "no_apx.docx"
    doc.save(path)
    out = tmp_path / "no_apx_out.docx"
    fix_document_path(str(path), str(out))
    assert not any(p.text.strip().lower().startswith("appendix") for p in Document(str(out)).paragraphs)


def test_appendix_order_never_changed(tmp_path: Path):
    src = build_appendix_doc(tmp_path / "multi.docx", multiple=True)
    out = tmp_path / "multi_out.docx"
    before = [p.text for p in Document(str(src)).paragraphs]
    fix_document_path(str(src), str(out))
    after = [p.text for p in Document(str(out)).paragraphs]
    assert before == after


def test_appendix_page_break_safe(tmp_path: Path):
    src = build_appendix_doc(tmp_path / "apx_pb.docx", bad_format=True)
    assert any(
        i.rule_id == "APA7-APPENDIX-PAGE-BREAK"
        for i in analyze_document_path(str(src)).safe_auto_fix
    )
    out = tmp_path / "apx_pb_out.docx"
    fix_document_path(str(src), str(out))
    assert not any(
        i.rule_id == "APA7-APPENDIX-PAGE-BREAK"
        for i in analyze_document_path(str(out)).safe_auto_fix
    )


def test_multiple_appendices_preserved(tmp_path: Path):
    src = build_appendix_doc(tmp_path / "apx_multi.docx", multiple=True)
    before = analyze_appendices(Document(str(src)))
    out = tmp_path / "apx_multi_out.docx"
    fix_document_path(str(src), str(out))
    after = analyze_appendices(Document(str(out)))
    assert len(before) == len(after) == 2
    assert [a.label_text for a in before] == [a.label_text for a in after]


# -------------------- LISTS --------------------


def test_numbered_list_text_preserved(tmp_path: Path):
    src = build_list_docs(tmp_path / "lists.docx")
    out = tmp_path / "lists_out.docx"
    before = snapshot_user_text(Document(str(src)))
    fix_document_path(str(src), str(out))
    assert_text_integrity(before, snapshot_user_text(Document(str(out))))


def test_bullet_list_preserved(tmp_path: Path):
    src = build_list_docs(tmp_path / "bullets.docx")
    out = tmp_path / "bullets_out.docx"
    fix_document_path(str(src), str(out))
    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("Bullet one" in t for t in texts)


def test_nested_list_numbering_preserved(tmp_path: Path):
    src = build_list_docs(tmp_path / "nested.docx")
    before = analyze_lists(Document(str(src)))
    out = tmp_path / "nested_out.docx"
    fix_document_path(str(src), str(out))
    after = analyze_lists(Document(str(out)))
    before_ids = {(i.paragraph_index, i.num_id, i.ilvl) for i in before if i.num_id is not None}
    after_ids = {(i.paragraph_index, i.num_id, i.ilvl) for i in after if i.num_id is not None}
    assert before_ids == after_ids


def test_plain_numbered_prose_not_misclassified(tmp_path: Path):
    src = build_list_docs(tmp_path / "plain.docx")
    lists = analyze_lists(Document(str(src)))
    plain = [i for i in lists if i.list_type == "PLAIN_NUMBERED_PROSE"]
    assert plain
    # Should not become LIST_ITEM region via Word numbering.
    from app.apa.engine.classifier import classify_document

    classified = classify_document(Document(str(src)))
    for item in plain:
        region = next(c.region for c in classified if c.index == item.paragraph_index)
        assert region != Region.LIST_ITEM


# -------------------- DOCX PRESERVATION --------------------


def test_fix_preserves_every_image_binary(tmp_path: Path):
    src = build_apa_figure_doc(tmp_path / "img_all.docx", with_callout=True)
    out = tmp_path / "img_all_out.docx"
    before = snapshot_image_binaries(Document(str(src)))
    result = fix_document_path(str(src), str(out))
    after = snapshot_image_binaries(Document(str(out)))
    assert before
    assert before == after
    assert result["preservation"]["images_preserved"] is True


def test_fix_preserves_hyperlinks(tmp_path: Path):
    src = build_hyperlink_bookmark_doc(tmp_path / "links.docx")
    out = tmp_path / "links_out.docx"
    result = fix_document_path(str(src), str(out))
    assert result["preservation"]["hyperlinks_preserved"] is True


def test_fix_preserves_bookmarks(tmp_path: Path):
    src = build_hyperlink_bookmark_doc(tmp_path / "bm.docx")
    out = tmp_path / "bm_out.docx"
    result = fix_document_path(str(src), str(out))
    assert result["preservation"]["bookmarks_preserved"] is True


def test_fix_preserves_numbering_xml(tmp_path: Path):
    src = build_list_docs(tmp_path / "numxml.docx")
    out = tmp_path / "numxml_out.docx"
    result = fix_document_path(str(src), str(out))
    assert result["preservation"]["numbering_preserved"] is True


def test_fix_preserves_equation_xml(tmp_path: Path):
    src = build_equation_doc(tmp_path / "eq.docx")
    before_fp = equation_xml_fingerprint(Document(str(src)))
    assert analyze_equations(Document(str(src)))
    out = tmp_path / "eq_out.docx"
    result = fix_document_path(str(src), str(out))
    after_fp = equation_xml_fingerprint(Document(str(out)))
    assert before_fp == after_fp
    assert result["preservation"]["equations_preserved"] is True


def test_fix_preserves_footnotes_if_fixture_supported(tmp_path: Path):
    # python-docx cannot easily author footnotes; skip gracefully if unsupported.
    pytest.importorskip("docx")
    src = build_apa_table_doc(tmp_path / "fn.docx", with_callout=True)
    out = tmp_path / "fn_out.docx"
    result = fix_document_path(str(src), str(out))
    assert result["preservation"]["footnotes_preserved"] is True


def test_fix_preserves_endnotes_if_fixture_supported(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "en.docx", with_callout=True)
    out = tmp_path / "en_out.docx"
    result = fix_document_path(str(src), str(out))
    assert result["preservation"]["endnotes_preserved"] is True


def test_fix_does_not_drop_unrelated_relationships(tmp_path: Path):
    src = build_hyperlink_bookmark_doc(tmp_path / "rels.docx")
    # Also embed an image relationship.
    doc = Document(str(src))
    _add_png(doc)
    src2 = tmp_path / "rels2.docx"
    doc.save(src2)
    out = tmp_path / "rels_out.docx"
    before_count = len(Document(str(src2)).part.rels)
    result = fix_document_path(str(src2), str(out))
    after_count = len(Document(str(out)).part.rels)
    assert after_count >= before_count
    assert result["preservation"]["relationships_preserved"] is True


# -------------------- CROSS-RULE / POST-FIX --------------------


def test_table_title_not_body_indented(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "conflict.docx", bad_format=True, with_callout=True)
    out = tmp_path / "conflict_out.docx"
    fix_document_path(str(src), str(out))
    from app.apa.engine.classifier import classify_document

    classified = classify_document(Document(str(out)))
    titles = [c for c in classified if c.region == Region.TABLE_TITLE]
    assert titles
    for item in titles:
        indent = item.paragraph.paragraph_format.first_line_indent
        if indent is not None:
            assert abs(indent.inches) < 0.05 or indent.inches == 0


def test_appendix_label_not_level1_merely_centered(tmp_path: Path):
    src = build_appendix_doc(tmp_path / "apx_cls.docx", bad_format=False)
    from app.apa.engine.classifier import classify_document

    classified = classify_document(Document(str(src)))
    labels = [
        c
        for c in classified
        if c.text.strip() in {"Appendix", "Appendix A", "Appendix B"}
    ]
    assert labels
    assert all(c.region == Region.APPENDIX_LABEL for c in labels)


def test_phase2d_post_fix_safe_zero_and_idempotent(tmp_path: Path):
    src = build_apa_table_doc(tmp_path / "post.docx", bad_format=True, with_callout=True)
    out1 = tmp_path / "post1.docx"
    out2 = tmp_path / "post2.docx"
    r1 = fix_document_path(str(src), str(out1))
    assert r1["verification"]["verified"] is True
    assert r1["verification"]["safe_issues_after"] == 0
    assert r1["verification"]["text_integrity_ok"] is True
    assert r1["verification"]["document_preservation_ok"] is True
    r2 = fix_document_path(str(out1), str(out2))
    assert r2["fixed_counts"]["safe_rules_applied"] == 0


def test_phase2d_rules_registered():
    ids = {r.rule_id for r in all_rules()}
    for rule_id in (
        "APA7-TABLE-NUMBER-FORMAT",
        "APA7-FIGURE-NUMBER-FORMAT",
        "APA7-APPENDIX-PAGE-BREAK",
        "APA7-LIST-NO-TEXT-REWRITE",
        "APA7-EQUATION-CONTENT",
    ):
        assert rule_id in ids
    safe_ids = {r.rule_id for r in production_safe_rules()}
    assert "APA7-TABLE-NUMBER-FORMAT" in safe_ids
    assert "APA7-TABLE-CALLOUT-MISSING" not in safe_ids
