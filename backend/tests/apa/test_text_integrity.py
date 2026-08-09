from docx import Document

from app.apa.engine.fixer import fix_document_path
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text


def test_fix_preserves_every_user_character(failing_docx, tmp_path):
    before = snapshot_user_text(Document(str(failing_docx)))
    output = tmp_path / "fixed.docx"
    fix_document_path(str(failing_docx), str(output))
    after = snapshot_user_text(Document(str(output)))
    assert_text_integrity(before, after)


def test_golden_snapshot_stable(golden_docx):
    first = snapshot_user_text(Document(str(golden_docx)))
    second = snapshot_user_text(Document(str(golden_docx)))
    assert first.fingerprint == second.fingerprint
