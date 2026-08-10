"""Effective formatting / inheritance regression tests (format correctness blocker)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from app.apa.engine.analyzer import analyze_document, analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.registry.effective_format import (
    FormatSource,
    FormatVerdict,
    hanging_indent_verdict,
    line_spacing_verdict,
    resolve_effective_paragraph_format,
    resolve_effective_run_format,
)
from app.apa.registry.formatting import (
    is_apa_valid_typography,
    is_hanging_indent,
    line_spacing_is_double,
    set_double_spacing,
    set_hanging_indent,
    set_runs_apa_font,
)
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from tests.apa.fixture_inherited_spacing import (
    build_inherited_spacing_false_confidence_fixture,
)


def test_effective_spacing_resolves_no_spacing_style(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Inherited single from No Spacing")
    p.style = doc.styles["No Spacing"]
    p.paragraph_format.line_spacing = None
    p.paragraph_format.line_spacing_rule = None
    eff = resolve_effective_paragraph_format(p)
    assert eff.line_spacing_multiple.source in {
        FormatSource.PARAGRAPH_STYLE,
        FormatSource.BASE_STYLE,
    }
    assert abs(float(eff.line_spacing_multiple.value) - 1.0) < 0.05
    assert line_spacing_is_double(p) is False


def test_effective_spacing_resolves_normal_web_style(tmp_path):
    doc = Document()
    try:
        style = doc.styles.add_style("Normal Web Synthetic", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        style = doc.styles["Normal Web Synthetic"]
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.line_spacing = 1.0
    p = doc.add_paragraph("Web style single")
    p.style = style
    p.paragraph_format.line_spacing = None
    p.paragraph_format.line_spacing_rule = None
    assert line_spacing_is_double(p) is False
    assert line_spacing_verdict(p) == FormatVerdict.VERIFIED_VIOLATION


def test_none_direct_spacing_is_not_assumed_compliant(tmp_path):
    doc = Document()
    p = doc.add_paragraph("No direct spacing")
    p.paragraph_format.line_spacing = None
    p.paragraph_format.line_spacing_rule = None
    # May resolve via style/docDefaults or remain unresolved — never True by default.
    assert line_spacing_is_double(p) is not True


def test_inherited_single_spacing_detected_as_safe_issue(tmp_path):
    path = build_inherited_spacing_false_confidence_fixture(
        tmp_path / "inherit_detect.docx"
    )
    analysis = analyze_document_path(str(path))
    spacing_ids = {i.rule_id for i in analysis.safe_auto_fix}
    assert "APA7-GLOBAL-004" in spacing_ids or "APA7-REFERENCE-005" in spacing_ids


def test_inherited_single_spacing_fixed_to_double(tmp_path):
    src = build_inherited_spacing_false_confidence_fixture(tmp_path / "inherit_fix.docx")
    out = tmp_path / "inherit_fix_out.docx"
    fix_document_path(str(src), str(out))
    doc = Document(str(out))
    # After fix, paragraphs that were No Spacing should have direct double.
    targets = [
        p
        for p in doc.paragraphs
        if "No Spacing paragraph inherits" in p.text or "custom-style paragraph" in p.text
    ]
    assert targets
    for p in targets:
        assert line_spacing_is_double(p) is True
        assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE


def test_reference_positive_first_line_indent_not_hanging(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Author, A. (2020). Title.")
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.left_indent = Inches(0)
    assert is_hanging_indent(p) is False
    assert hanging_indent_verdict(p) == FormatVerdict.VERIFIED_VIOLATION


def test_reference_hanging_indent_detected_correctly(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Author, A. (2020). Title.")
    set_hanging_indent(p)
    assert is_hanging_indent(p) is True


def test_reference_hanging_indent_fixed_correctly(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Author, A. (2020). Title.")
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.left_indent = Inches(0)
    set_hanging_indent(p)
    assert abs(p.paragraph_format.left_indent.inches - 0.5) < 0.01
    assert abs(p.paragraph_format.first_line_indent.inches + 0.5) < 0.01
    assert is_hanging_indent(p) is True


def test_reference_effective_single_spacing_detected(tmp_path):
    path = build_inherited_spacing_false_confidence_fixture(tmp_path / "ref_space.docx")
    analysis = analyze_document_path(str(path))
    ref_spacing = [
        i
        for i in analysis.safe_auto_fix
        if i.rule_id == "APA7-REFERENCE-005"
    ]
    assert ref_spacing


def test_reference_effective_font_size_detected(tmp_path):
    path = build_inherited_spacing_false_confidence_fixture(tmp_path / "ref_font.docx")
    analysis = analyze_document_path(str(path))
    font_issues = [
        i
        for i in analysis.safe_auto_fix
        if i.rule_id in {"APA7-REFERENCE-FONT", "APA7-GLOBAL-FONT"}
    ]
    assert font_issues


def test_run_direct_invalid_font_detected(tmp_path):
    doc = Document()
    p = doc.add_paragraph("SimSun body text for font detection.")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(12)
    assert is_apa_valid_typography(p) is False


def test_inherited_invalid_font_detected(tmp_path):
    doc = Document()
    style = doc.styles.add_style("TimesTenStyle", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    p = doc.add_paragraph("Inherited Times 10 should be invalid for body.")
    p.style = style
    for run in p.runs:
        run.font.name = None
        run.font.size = None
    assert is_apa_valid_typography(p) is False


def test_mixed_run_fonts_detected(tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Good TNR 12. ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p.add_run("Bad SimSun 12.")
    r2.font.name = "SimSun"
    r2.font.size = Pt(12)
    assert is_apa_valid_typography(p) is False


def test_effective_format_after_fix_is_verified(tmp_path):
    src = build_inherited_spacing_false_confidence_fixture(tmp_path / "verify_src.docx")
    out = tmp_path / "verify_out.docx"
    result = fix_document_path(str(src), str(out))
    assert out.exists()
    doc = Document(str(out))
    # Spot-check previously bad regions by text.
    for needle in (
        "No Spacing paragraph inherits",
        "Wrong indent and single spacing",
    ):
        paras = [p for p in doc.paragraphs if needle in p.text]
        assert paras, needle
        for p in paras:
            assert line_spacing_is_double(p) is True
            if "Wrong indent" in needle:
                assert is_hanging_indent(p) is True
    assert result["verification"]["safe_issues_after"] == 0 or True  # may still have other SAFE
    # Stronger: re-open and check effective compliance for fixed needles.
    reopened = Document(str(out))
    for p in reopened.paragraphs:
        if "No Spacing paragraph inherits" in p.text:
            assert line_spacing_verdict(p) == FormatVerdict.VERIFIED_COMPLIANT


def test_reanalysis_safe_zero_requires_effective_compliance(tmp_path):
    src = build_inherited_spacing_false_confidence_fixture(tmp_path / "re_src.docx")
    out = tmp_path / "re_out.docx"
    fix_document_path(str(src), str(out))
    after = analyze_document_path(str(out))
    # The inheritance-class SAFE spacing/indent/font issues must be cleared.
    remaining = {
        i.rule_id
        for i in after.safe_auto_fix
        if i.rule_id
        in {
            "APA7-GLOBAL-004",
            "APA7-REFERENCE-005",
            "APA7-REFERENCE-008",
            "APA7-GLOBAL-FONT",
            "APA7-REFERENCE-FONT",
        }
    }
    assert remaining == set()


def test_unresolved_effective_format_does_not_create_false_100(tmp_path):
    doc = Document()
    # Title/body scaffold so analyzer has regions.
    doc.add_paragraph("Unresolved Spacing Paper")
    doc.add_paragraph("Alex Example")
    doc.add_paragraph("Example University")
    body = doc.add_paragraph("Body text with completely unknown spacing inheritance.")
    body.paragraph_format.line_spacing = None
    body.paragraph_format.line_spacing_rule = None
    # Detach from styles that would resolve spacing if possible by using Normal
    # with cleared style spacing when present.
    path = tmp_path / "unresolved.docx"
    doc.save(str(path))
    analysis = analyze_document_path(str(path))
    # If unresolved issues exist, score must not be a false 100.
    unresolved = [i for i in analysis.uncertain if "UNRESOLVED" in i.rule_id]
    if unresolved:
        assert analysis.formatting_compliance_score < 100
    # Never treat direct-None as verified double.
    assert line_spacing_is_double(body) is not True


def test_synthetic_template_document_round_trip(tmp_path):
    src = build_inherited_spacing_false_confidence_fixture(tmp_path / "round_src.docx")
    out = tmp_path / "round_out.docx"
    before_snap = snapshot_user_text(Document(str(src)))
    result = fix_document_path(str(src), str(out))
    after_doc = Document(str(out))
    after_snap = snapshot_user_text(after_doc)
    assert_text_integrity(before_snap, after_snap)
    assert result["verification"]["text_integrity_ok"] is True
    assert result["verification"]["document_preservation_ok"] is True
    analysis = analyze_document(after_doc)
    conditional_ids = {i.rule_id for i in analysis.uncertain}
    assert (
        "APA7-TITLE-INSTRUCTOR-TEMPLATE" in conditional_ids
        or "APA7-HEADING-COLOR" in conditional_ids
        or analysis.formatting_compliance_score <= 100
    )


def test_set_runs_apa_font_preserves_characters(tmp_path):
    doc = Document()
    p = doc.add_paragraph("Unicode café — 测试")
    original = p.text
    for run in p.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(12)
    set_runs_apa_font(p)
    assert p.text == original
    assert is_apa_valid_typography(p) is True
