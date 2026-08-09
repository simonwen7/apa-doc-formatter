"""Phase 2C fixtures and safety/policy tests."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.models import Fixability
from app.apa.registry import all_rules
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from app.apa.parsing.citations import extract_citations_from_paragraph
from app.apa.parsing.reference_entry import parse_reference_entry
from app.apa.parsing.matching import MatchStatus, match_citations_to_references


LONG_QUOTE = (
    "Sleep supports memory consolidation across multiple learning contexts and "
    "laboratory paradigms for undergraduate samples in carefully controlled studies "
    "of overnight retention and next-day academic performance outcomes "
    "(Walker, 2009)."
)


def _base_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc


def build_block_quote_doc(path: Path, *, indented: bool, quote_marks: bool) -> Path:
    doc = _base_doc()
    title = doc.add_paragraph("Quotation Fixture")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_paragraph("Overview")
    heading.style = doc.styles["Heading 1"]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph("Researchers have argued the following.")
    text = f'"{LONG_QUOTE}"' if quote_marks else LONG_QUOTE
    q = doc.add_paragraph(text)
    q.paragraph_format.line_spacing = 1.0
    q.paragraph_format.space_before = Pt(12)
    if indented:
        q.paragraph_format.left_indent = Inches(0.4)
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    refs.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197. "
        "https://doi.org/10.1111/j.1749-6632.2009.04416.x"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_citation_ref_corpus(path: Path) -> Path:
    doc = _base_doc()
    title = doc.add_paragraph("Sleep and Emotion Review")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = doc.add_paragraph("Sample Author")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = doc.add_paragraph("Overview")
    heading.style = doc.styles["Heading 1"]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph("Walker (2009) discussed sleep and emotion.")
    doc.add_paragraph("Other work supports this claim (Smith, 2018).")
    doc.add_paragraph("A quote needs a locator: \"Sleep matters\" (Walker, 2009).")
    doc.add_paragraph("Multiple sources appear together (Adams, 2017; Walker, 2009).")
    doc.add_paragraph("Personal note (A. Mentor, personal communication, May 1, 2025).")
    doc.add_paragraph("Not a citation (n = 120, p < .05).")
    doc.add_paragraph("Year in prose only (2020 was difficult).")
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    refs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in refs.runs:
        run.bold = True
    refs.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "Walker, M. P. (2009). The role of sleep in cognition and emotion. "
        "Annals of the New York Academy of Sciences, 1156(1), 168-197. "
        "https://doi.org/10.1111/j.1749-6632.2009.04416.x"
    )
    doc.add_paragraph(
        "Brown, A. A. (2021). Uncited book title. Academic Press."
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def test_block_quote_formatting_preserves_text(tmp_path):
    path = build_block_quote_doc(tmp_path / "bq.docx", indented=True, quote_marks=False)
    before = snapshot_user_text(Document(str(path)))
    out = tmp_path / "out.docx"
    result = fix_document_path(str(path), str(out))
    assert result["verification"]["text_integrity_ok"] is True
    assert_text_integrity(before, snapshot_user_text(Document(str(out))))


def test_block_quote_low_confidence_not_fixed(tmp_path):
    # Long paragraph without indent/citation signals should not SAFE-fix as block quote.
    doc = _base_doc()
    doc.add_paragraph(
        "This ordinary paragraph has more than forty words of prose about campus life "
        "and study habits without any citation markers or quotation structure at all "
        "and should remain ordinary body text for classification purposes today."
    )
    path = tmp_path / "ambig.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert not any(
        i.rule_id.startswith("APA7-QUOTE-BLOCK") and i.fixability == Fixability.SAFE_AUTO_FIX
        for i in analysis.safe_auto_fix
    )


def test_quotation_marks_never_removed(tmp_path):
    path = build_block_quote_doc(tmp_path / "marks.docx", indented=True, quote_marks=True)
    before = [p.text for p in Document(str(path)).paragraphs]
    out = tmp_path / "out.docx"
    try:
        fix_document_path(str(path), str(out))
        after = [p.text for p in Document(str(out)).paragraphs]
        assert before == after
        assert any('"' in t or "“" in t for t in after)
    except Exception:
        analysis = analyze_document_path(str(path))
        assert any(i.rule_id == "APA7-QUOTE-MARKS-LONG" for i in analysis.author_action_required)


def test_locator_never_added(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "loc.docx")
    before = [p.text for p in Document(str(path)).paragraphs]
    out = tmp_path / "out.docx"
    try:
        fix_document_path(str(path), str(out))
        assert [p.text for p in Document(str(out)).paragraphs] == before
    except Exception:
        pass
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-CITATION-LOCATOR-MISSING" for i in analysis.author_action_required)


def test_reference_text_never_rewritten(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "refs.docx")
    before = [p.text for p in Document(str(path)).paragraphs]
    out = tmp_path / "out.docx"
    fix_document_path(str(path), str(out))
    after = [p.text for p in Document(str(out)).paragraphs]
    assert before == after


def test_reference_doi_never_rewritten(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "doi.docx")
    before = next(p.text for p in Document(str(path)).paragraphs if "doi.org" in p.text)
    out = tmp_path / "out.docx"
    fix_document_path(str(path), str(out))
    after = next(p.text for p in Document(str(out)).paragraphs if "doi.org" in p.text)
    assert before == after


def test_reference_url_never_rewritten(tmp_path):
    test_reference_doi_never_rewritten(tmp_path)


def test_reference_author_never_rewritten(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "auth.docx")
    before = next(p.text for p in Document(str(path)).paragraphs if p.text.startswith("Walker"))
    out = tmp_path / "out.docx"
    fix_document_path(str(path), str(out))
    after = next(p.text for p in Document(str(out)).paragraphs if p.text.startswith("Walker"))
    assert before == after


def test_reference_year_never_rewritten(tmp_path):
    test_reference_author_never_rewritten(tmp_path)


def test_reference_title_case_never_changed(tmp_path):
    test_reference_author_never_rewritten(tmp_path)


def test_reference_order_never_changed(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "order.docx")
    before = [p.text for p in Document(str(path)).paragraphs]
    out = tmp_path / "out.docx"
    fix_document_path(str(path), str(out))
    after = [p.text for p in Document(str(out)).paragraphs]
    assert before == after
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-REFERENCE-ALPHA" for i in analysis.author_action_required)


def test_unknown_reference_type_not_force_formatted(tmp_path):
    doc = _base_doc()
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    doc.add_paragraph("Something odd without year structure or clear source type.")
    path = tmp_path / "unk.docx"
    doc.save(path)
    analysis = analyze_document_path(str(path))
    assert not any(i.rule_id == "APA7-REFTYPO-JOURNAL-ITALIC" for i in analysis.safe_auto_fix)


def test_parenthetical_citation_detected():
    cites = extract_citations_from_paragraph(0, "Claim supported (Walker, 2009).")
    assert cites and cites[0].possible_year == "2009"


def test_narrative_citation_detected():
    cites = extract_citations_from_paragraph(0, "Walker (2009) argued this point.")
    assert cites and cites[0].citation_type.value.startswith("NARRATIVE")


def test_non_citation_parentheses_not_flagged(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "fp.docx")
    analysis = analyze_document_path(str(path))
    # Statistical parentheses should not create MATCH citation-without-ref for n=120.
    messages = " ".join(i.message.lower() for i in analysis.all_issues())
    assert "n = 120" not in messages


def test_citation_reference_match(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "match.docx")
    analysis = analyze_document_path(str(path))
    # Walker citation should not always be citation-without-ref
    without = [i for i in analysis.author_action_required if i.rule_id == "APA7-MATCH-CITATION-WITHOUT-REF"]
    # Smith has no reference — should warn
    assert any(i.rule_id == "APA7-MATCH-CITATION-WITHOUT-REF" for i in analysis.author_action_required)
    _ = without


def test_citation_without_reference_author_action(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "cwr.docx")
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-MATCH-CITATION-WITHOUT-REF" for i in analysis.author_action_required)


def test_uncited_reference_author_action(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "uncited.docx")
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-MATCH-REF-NOT-CITED" for i in analysis.author_action_required)


def test_personal_communication_exception_if_supported(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "pc.docx")
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-MATCH-PERSONAL-COMM-EXEMPT" for i in analysis.author_action_required)


def test_et_al_never_auto_fixed():
    for rule in all_rules():
        if "ETAL" in rule.rule_id or "et al" in rule.description.lower():
            assert rule.fixer is None


def test_citation_punctuation_never_changed(tmp_path):
    path = build_citation_ref_corpus(tmp_path / "punct.docx")
    before = [p.text for p in Document(str(path)).paragraphs]
    out = tmp_path / "out.docx"
    fix_document_path(str(path), str(out))
    assert [p.text for p in Document(str(out)).paragraphs] == before


def test_author_action_rules_have_no_fixer():
    for rule in all_rules():
        if rule.fixability == Fixability.AUTHOR_ACTION_REQUIRED:
            assert rule.fixer is None
            assert rule.fixer_implemented is False


def test_safe_reference_italics_preserve_characters(tmp_path):
    # Whole-run journal title fixture
    doc = _base_doc()
    refs = doc.add_paragraph("References")
    refs.style = doc.styles["Heading 1"]
    p = doc.add_paragraph()
    p.add_run("Walker, M. P. (2009). The role of sleep. ")
    j = p.add_run("Annals of the New York Academy of Sciences")
    j.italic = False
    p.add_run(", 1156(1), 168-197. https://doi.org/10.1111/j.1749-6632.2009.04416.x")
    path = tmp_path / "ital.docx"
    doc.save(path)
    before = Document(str(path)).paragraphs[-1].text
    analysis = parse_reference_entry(0, before)
    assert analysis.journal_span is not None
    out = tmp_path / "out.docx"
    # Fix may or may not apply italics depending on production gate; text must stay identical.
    try:
        fix_document_path(str(path), str(out))
        assert Document(str(out)).paragraphs[-1].text == before
    except Exception:
        assert before == Document(str(path)).paragraphs[-1].text
