from pathlib import Path

from docx import Document

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.models import Fixability
from app.apa.registry import production_safe_rules
from app.apa.text_integrity import snapshot_user_text


PRODUCTION_SAFE_IDS = {rule.rule_id for rule in production_safe_rules()}


def test_fix_clears_production_safe_issues(failing_docx, tmp_path):
    output = tmp_path / "fixed.docx"
    result = fix_document_path(str(failing_docx), str(output))

    assert result["verification"]["verified"] is True
    assert result["verification"]["safe_issues_after"] == 0

    after = analyze_document_path(str(output))
    remaining = [
        issue.rule_id
        for issue in after.safe_auto_fix
        if issue.rule_id in PRODUCTION_SAFE_IDS
    ]
    assert remaining == []


def test_author_action_rules_are_never_fixed(failing_docx, tmp_path):
    before = analyze_document_path(str(failing_docx))
    author_before = {issue.rule_id for issue in before.author_action_required}
    assert author_before

    output = tmp_path / "fixed.docx"
    result = fix_document_path(str(failing_docx), str(output))

    # Author-action issues remain author-action; never become applied fixes.
    assert all(
        issue["fixability"] == Fixability.AUTHOR_ACTION_REQUIRED.value
        for issue in result["author_action_required"]
    )
    applied = result["fixed_counts"]
    assert "text_normalized" not in applied or applied.get("text_normalized", 0) == 0

    after_doc = Document(str(output))
    before_doc = Document(str(failing_docx))
    # Title text characters must remain unchanged.
    assert before_doc.paragraphs[0].text == after_doc.paragraphs[0].text


def test_legacy_fixer_disabled():
    from app.services.fixer import fix_apa_format
    import pytest

    with pytest.raises(RuntimeError):
        fix_apa_format("in.docx", "out.docx")
