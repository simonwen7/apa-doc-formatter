"""Phase 3B commercial golden-corpus runner.

Documents are generated into tmp; canonical fixtures under tests/apa/fixtures
are never mutated.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.parsing.docx_objects import snapshot_image_binaries, snapshot_package_parts
from app.apa.registry import production_safe_rules
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from tests.apa.golden_corpus.builders import CORPUS_BUILDERS

PRODUCTION_SAFE_IDS = {rule.rule_id for rule in production_safe_rules()}

# Citation-like false positives we do not want on stress paper F.
FORBIDDEN_FALSE_POSITIVE_RULES = {
    "APA7-BLOCKQUOTE-001",
    "APA7-CITATION-ONE-AUTHOR",
    "APA7-CITATION-TWO-AUTHOR",
    "APA7-CITATION-ETAL",
}


def _safe_ids(analysis) -> list[str]:
    return [
        issue.rule_id
        for issue in analysis.safe_auto_fix
        if issue.rule_id in PRODUCTION_SAFE_IDS
    ]


def _count_bucket(analysis, name: str) -> int:
    return len(getattr(analysis, name) or [])


def _round_trip_analyze(path: Path) -> list[str]:
    """Serialize → reopen from disk → analyze SAFE ids."""
    reopened = Document(str(path))
    # Force a fresh save/load cycle for the reopen path under test.
    scratch = path.with_name(path.stem + "_reopen.docx")
    reopened.save(scratch)
    analysis = analyze_document_path(str(scratch))
    return _safe_ids(analysis)


def test_golden_corpus_inventory_complete():
    expected = {
        "A_clean_basic",
        "B_misformatted",
        "C_author_only",
        "D_mixed",
        "E_complex",
        "F_false_positive_stress",
        "G_multi_author",
        "H_reference_types",
    }
    assert set(CORPUS_BUILDERS) == expected


def test_golden_corpus_commercial_qa(tmp_path: Path):
    rows: list[dict] = []
    fixtures_before = {
        p.name: p.stat().st_mtime_ns
        for p in (Path(__file__).parent / "fixtures").glob("*.docx")
    }

    for name, spec in CORPUS_BUILDERS.items():
        src = tmp_path / f"{name}.docx"
        out = tmp_path / f"{name}_fixed.docx"
        out2 = tmp_path / f"{name}_fixed2.docx"
        spec["builder"](src)

        before_analysis = analyze_document_path(str(src))
        safe_before = _safe_ids(before_analysis)
        author_before = _count_bucket(before_analysis, "author_action_required")
        uncertain_before = _count_bucket(before_analysis, "uncertain")

        unexpected: list[str] = []
        if spec.get("expect_safe_zero_before") and safe_before:
            unexpected.append(f"expected_zero_safe_got={safe_before}")
        if spec.get("expect_author_min") and author_before < spec["expect_author_min"]:
            unexpected.append("expected_author_issues")
        if spec.get("forbid_false_citation_rules"):
            bad = [
                i.rule_id
                for i in before_analysis.author_action_required
                if i.rule_id in FORBIDDEN_FALSE_POSITIVE_RULES
            ]
            if bad:
                unexpected.append(f"false_positive_rules={bad}")

        fix_attempted = False
        verified = None
        safe_after: list[str] = safe_before
        text_ok = True
        preserv_ok = True
        idempotent = True

        if spec.get("expect_fix"):
            assert safe_before, f"{name}: expected SAFE issues before Fix"
            fix_attempted = True
            before_snap = snapshot_user_text(Document(str(src)))
            before_parts = snapshot_package_parts(str(src))
            before_images = snapshot_image_binaries(Document(str(src)))

            result = fix_document_path(str(src), str(out))
            verified = result["verification"]["verified"]
            assert verified is True, f"{name}: Fix not verified"
            assert result["verification"]["safe_issues_after"] == 0
            assert result["verification"]["text_integrity_ok"] is True
            assert result["verification"]["document_preservation_ok"] is True

            after_snap = snapshot_user_text(Document(str(out)))
            assert_text_integrity(before_snap, after_snap)

            after_parts = snapshot_package_parts(str(out))
            after_images = snapshot_image_binaries(Document(str(out)))
            for part, digest in before_parts.items():
                if part.startswith("word/media/") or part.startswith("word/embeddings/"):
                    assert part in after_parts, f"{name}: lost package part {part}"
                    assert after_parts[part] == digest, f"{name}: changed {part}"
            for img_key, digest in before_images.items():
                assert img_key in after_images, f"{name}: lost image {img_key}"
                assert after_images[img_key] == digest

            # Round-trip: reopen serialized DOCX from disk.
            safe_after = _round_trip_analyze(out)
            assert safe_after == [], f"{name}: SAFE remained after reopen: {safe_after}"

            # Download-equivalent copy + analyze again.
            download_copy = tmp_path / f"{name}_download.docx"
            shutil.copyfile(out, download_copy)
            assert _round_trip_analyze(download_copy) == []

            # Idempotency.
            second = fix_document_path(str(out), str(out2))
            assert second["verification"]["verified"] is True
            assert second["verification"]["safe_issues_before"] == 0
            assert second["fixed_counts"]["safe_rules_applied"] == 0

            # AUTHOR issues must still be present when expected (mixed).
            if spec.get("expect_author_min"):
                after_analysis = analyze_document_path(str(out))
                assert (
                    _count_bucket(after_analysis, "author_action_required")
                    >= spec["expect_author_min"]
                )
        else:
            # Untouched documents: Fix not required; SAFE should stay zero.
            assert not safe_before, f"{name}: unexpected SAFE {safe_before}"
            # Optional explicit Fix should apply 0 rules if attempted.
            result = fix_document_path(str(src), str(out))
            verified = result["verification"]["verified"]
            assert result["fixed_counts"]["safe_rules_applied"] == 0
            assert result["verification"]["safe_issues_after"] == 0
            safe_after = _safe_ids(analyze_document_path(str(out)))

        rows.append(
            {
                "filename": name,
                "safe_count_before": len(safe_before),
                "author_count_before": author_before,
                "uncertain_count_before": uncertain_before,
                "fix_attempted": fix_attempted,
                "verified": verified,
                "safe_count_after": len(safe_after),
                "text_integrity_ok": text_ok,
                "document_preservation_ok": preserv_ok,
                "idempotent": idempotent,
                "unexpected_warnings": unexpected,
            }
        )
        assert not unexpected, f"{name}: {unexpected}"

    # Fixture immutability.
    fixtures_after = {
        p.name: p.stat().st_mtime_ns
        for p in (Path(__file__).parent / "fixtures").glob("*.docx")
    }
    assert fixtures_after == fixtures_before

    # Ensure we exercised fix + no-fix paths.
    assert any(r["fix_attempted"] for r in rows)
    assert any(not r["fix_attempted"] for r in rows)


def test_unicode_fingerprint_not_normalized(tmp_path: Path):
    """Fix must not normalize curly quotes / accents / dashes."""
    from tests.apa.golden_corpus.builders import build_e_complex

    src = tmp_path / "unicode.docx"
    out = tmp_path / "unicode_fixed.docx"
    build_e_complex(src)
    before = snapshot_user_text(Document(str(src)))
    # Confirm markers exist in snapshot.
    joined = "\n".join(text for _, text in before.segments)
    assert "café" in joined
    assert ("—" in joined) or ("–" in joined) or ("“" in joined)
    result = fix_document_path(str(src), str(out))
    assert result["verification"]["verified"] is True
    after = snapshot_user_text(Document(str(out)))
    assert_text_integrity(before, after)
