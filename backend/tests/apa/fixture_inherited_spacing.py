"""Synthetic DOCX reproducing inherited-style APA false-confidence failures.

Uses only harmless placeholder text — never real student content.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _ensure_style(doc: Document, name: str, *, based_on: str = "Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        try:
            style.base_style = doc.styles[based_on]
        except Exception:
            pass
        return style


def _set_run_east_asian_font(run, name: str, size_pt: float) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def build_inherited_spacing_false_confidence_fixture(path: Path) -> Path:
    """
    Reproduce the production false SAFE=0 class of failures:

    - Normal paragraph with explicit double spacing
    - No Spacing paragraph inheriting single spacing
    - Normal (Web)-like reference paragraph at single spacing
    - Custom style
    - Incorrect positive first-line reference indent
    - Valid hanging-indent reference
    - Run-level SimSun/宋体 12
    - Inherited Times 10 on a style
    - Colored heading
    - Mixed direct/style formatting
    """
    doc = Document()

    # Title-like decorative banner (instructor template signals).
    banner = doc.add_paragraph("Course Template Banner")
    run = banner.runs[0]
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x80)
    run.font.bold = True

    title = doc.add_paragraph("Synthetic Student Paper Title")
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Shading on title for template detection.
    pPr = title._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "008080")
    pPr.append(shd)

    doc.add_paragraph("Alex Example")
    doc.add_paragraph("Example University")
    doc.add_paragraph("PSY 101")
    doc.add_paragraph("Dr. Example")
    doc.add_paragraph("January 1, 2026")

    # Body title / colored heading.
    heading = doc.add_paragraph("Method")
    heading.style = doc.styles["Heading 1"]
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x80)
        run.font.bold = True
    heading.paragraph_format.page_break_before = True

    # Explicitly double-spaced Normal body paragraph.
    body_ok = doc.add_paragraph(
        "This Normal paragraph is explicitly double spaced for contrast."
    )
    body_ok.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    body_ok.paragraph_format.line_spacing = 2.0
    body_ok.paragraph_format.first_line_indent = Inches(0.5)
    body_ok.paragraph_format.space_before = Pt(0)
    body_ok.paragraph_format.space_after = Pt(0)
    for run in body_ok.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # No Spacing paragraph: no direct line spacing → inherits single.
    no_spacing = doc.add_paragraph(
        "This No Spacing paragraph inherits single spacing from its style."
    )
    no_spacing.style = doc.styles["No Spacing"]
    # Clear any accidental direct spacing so inheritance is the source of truth.
    no_spacing.paragraph_format.line_spacing = None
    no_spacing.paragraph_format.line_spacing_rule = None
    no_spacing.paragraph_format.first_line_indent = Inches(0.5)
    _set_run_east_asian_font(no_spacing.runs[0], "SimSun", 12)

    # Custom style based on No Spacing.
    custom = _ensure_style(doc, "Custom Single Body", based_on="No Spacing")
    custom.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    custom.paragraph_format.line_spacing = 1.0
    custom_para = doc.add_paragraph(
        "This custom-style paragraph inherits single spacing via basedOn."
    )
    custom_para.style = custom
    custom_para.paragraph_format.line_spacing = None
    custom_para.paragraph_format.line_spacing_rule = None
    custom_para.paragraph_format.first_line_indent = Inches(0.5)

    # Mixed fonts in one paragraph.
    mixed = doc.add_paragraph()
    mixed.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    mixed.paragraph_format.line_spacing = 2.0
    mixed.paragraph_format.first_line_indent = Inches(0.5)
    r1 = mixed.add_run("Valid Times New Roman twelve. ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = mixed.add_run("Invalid SimSun twelve.")
    _set_run_east_asian_font(r2, "宋体", 12)

    # References
    refs = doc.add_paragraph("References")
    refs.paragraph_format.page_break_before = True
    for run in refs.runs:
        run.bold = True

    # Normal (Web)-like single-spaced reference with positive first-line indent.
    web_style = _ensure_style(doc, "Normal Web Synthetic", based_on="Normal")
    web_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    web_style.paragraph_format.line_spacing = 1.0
    web_style.font.name = "Times New Roman"
    web_style.font.size = Pt(10)

    bad_ref = doc.add_paragraph(
        "Bad, A. A. (2020). Wrong indent and single spacing. Journal of Examples, 1(1), 1–2."
    )
    bad_ref.style = web_style
    bad_ref.paragraph_format.line_spacing = None
    bad_ref.paragraph_format.line_spacing_rule = None
    bad_ref.paragraph_format.first_line_indent = Inches(0.5)  # NOT hanging
    bad_ref.paragraph_format.left_indent = Inches(0)

    # Valid hanging indent reference (control).
    good_ref = doc.add_paragraph(
        "Good, B. B. (2021). Correct hanging indent example. Journal of Examples, 2(2), 3–4."
    )
    good_ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    good_ref.paragraph_format.line_spacing = 2.0
    good_ref.paragraph_format.left_indent = Inches(0.5)
    good_ref.paragraph_format.first_line_indent = Inches(-0.5)
    for run in good_ref.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
