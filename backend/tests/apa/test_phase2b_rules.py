"""Phase 2B rule detect/fix/integrity tests."""

from __future__ import annotations

from docx import Document

from app.apa.engine.analyzer import analyze_document_path
from app.apa.engine.fixer import fix_document_path
from app.apa.registry import production_safe_rules
from app.apa.text_integrity import assert_text_integrity, snapshot_user_text
from tests.apa.fixture_builders import (
    build_abstract_paper,
    build_golden_student_paper,
    build_heading_fixture,
    build_malformed_title_page,
)


PRODUCTION_SAFE_IDS = {rule.rule_id for rule in production_safe_rules()}


def test_golden_phase2b_zero_safe(tmp_path):
    path = build_golden_student_paper(tmp_path / "golden2b.docx")
    analysis = analyze_document_path(str(path))
    remaining = [i.rule_id for i in analysis.safe_auto_fix if i.rule_id in PRODUCTION_SAFE_IDS]
    assert remaining == [], remaining


def test_page_number_detect_and_fix(tmp_path):
    path = build_malformed_title_page(tmp_path / "no_page.docx")
    analysis = analyze_document_path(str(path))
    assert any(i.rule_id == "APA7-GLOBAL-011" for i in analysis.safe_auto_fix)


def test_title_safe_rules_detect(tmp_path):
    path = build_malformed_title_page(tmp_path / "title_bad.docx")
    ids = {i.rule_id for i in analyze_document_path(str(path)).safe_auto_fix}
    assert "APA7-TITLE-005" in ids or "APA7-TITLE-007" in ids


def test_heading_level_fixtures(tmp_path):
    for level in (1, 2, 3):
        good = build_heading_fixture(tmp_path / f"h{level}_ok.docx", level=level, valid=True)
        bad = build_heading_fixture(tmp_path / f"h{level}_bad.docx", level=level, valid=False)
        good_ids = {i.rule_id for i in analyze_document_path(str(good)).safe_auto_fix if "HEADING" in i.rule_id}
        bad_ids = {i.rule_id for i in analyze_document_path(str(bad)).safe_auto_fix if "HEADING" in i.rule_id}
        # Valid fixture should not retain level-specific SAFE issues for that level
        assert not any(f"HEADING{level}" in r for r in good_ids)
        assert any(f"HEADING{level}" in r for r in bad_ids)


def test_heading_fix_integrity_and_idempotent(tmp_path):
    path = build_heading_fixture(tmp_path / "h2_fix.docx", level=2, valid=False)
    before = snapshot_user_text(Document(str(path)))
    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"
    result = fix_document_path(str(path), str(out1))
    assert result["verification"]["verified"] is True
    assert result["verification"]["safe_issues_after"] == 0
    assert_text_integrity(before, snapshot_user_text(Document(str(out1))))
    second = fix_document_path(str(out1), str(out2))
    assert second["verification"]["safe_issues_before"] == 0
    assert second["fixed_counts"]["safe_rules_applied"] == 0


def test_abstract_fix_path(tmp_path):
    path = build_abstract_paper(tmp_path / "abs_fix.docx", well_formatted=False)
    before = snapshot_user_text(Document(str(path)))
    out = tmp_path / "abs_out.docx"
    result = fix_document_path(str(path), str(out))
    assert result["verification"]["verified"] is True
    assert_text_integrity(before, snapshot_user_text(Document(str(out))))
    after = analyze_document_path(str(out))
    assert not [
        i.rule_id
        for i in after.safe_auto_fix
        if i.rule_id in PRODUCTION_SAFE_IDS and i.rule_id.startswith("APA7-ABSTRACT")
    ]
